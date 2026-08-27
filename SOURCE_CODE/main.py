"""
main.py  - AI Automation Tool (version: see VERSION below)
Supports six workflow modes: coding, writing, rct_search, appraisal, search, sr.
Supports six AI providers: deepseek (default), openai, anthropic, ollama, groq, qwen.
RAG layer: per-session, mode-specific input/ folder indexing via rag.py.
Coding mode: Builder (pipeline), Reviewer (standalone), Tester (standalone).
"""

from __future__ import annotations

# --- Supported Python versions ---------------------------------------------
# This gate must run BEFORE any third-party import (dotenv, numpy, chromadb).
# Otherwise an unsupported interpreter produces a ModuleNotFoundError or a
# pip build failure instead of an actionable message.
#
# 3.13+ is excluded because several pinned dependencies have no wheels for it.
# Docker supplies its own interpreter and is unaffected.
import sys

MIN_PYTHON = (3, 11)
MAX_PYTHON_EXCLUSIVE = (3, 13)

if not (MIN_PYTHON <= sys.version_info[:2] < MAX_PYTHON_EXCLUSIVE):
    _found = ".".join(str(part) for part in sys.version_info[:3])
    sys.stderr.write(
        "\n"
        "  This project requires Python 3.11 or 3.12.\n"
        f"  You are running Python {_found}\n"
        f"  ({sys.executable})\n"
        "\n"
        "  Options:\n"
        "\n"
        "  1. Install Python 3.12 and rebuild the virtual environment:\n"
        "       https://www.python.org/downloads/release/python-3129/\n"
        "       py -3.12 -m venv .venv          (Windows)\n"
        "       python3.12 -m venv .venv        (macOS/Linux)\n"
        "\n"
        "  2. Use Docker, which supplies its own Python:\n"
        "       cd docker && docker compose run --rm cli\n"
        "\n"
        "  See Readme/Setup_Instructions_for_Users.txt for details.\n"
        "\n"
    )
    raise SystemExit(1)


import argparse
import json
import os
import re
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
from urllib.request import urlopen
import uuid
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv


# Add the SOURCE_CODE directory to Python path
PROJECT_ROOT = Path(__file__).resolve().parent.parent
SOURCE_CODE_DIR = PROJECT_ROOT / "SOURCE_CODE"
sys.path.insert(0, str(SOURCE_CODE_DIR))

# --- Heavy imports -----------------------------------------------------
# document_reader (-> pytesseract -> pandas) and the pipeline modules below
# are the slow part of startup (~7s cold). If Ctrl+C lands while one of
# these is still loading, Python raises KeyboardInterrupt from wherever the
# import chain happens to be (e.g. deep inside pandas/_libs), producing an
# alarming traceback even though nothing actually failed. The real
# entry-point handler at the bottom of this file only covers code that runs
# after imports finish, so it can't catch this. Catch it here too and print
# the same clean message instead.
try:
    # Now imports work from SOURCE_CODE/
    from utils.path_utils import PATH_MANAGER, get_input_dir, get_output_dir
    from utils.document_reader import DocumentReader
    # from utils.rag import RAGUtils (functions imported directly)

    # NOTE: run_coding, run_writing, run_search, and run_sr used to be imported
    # here unconditionally, but nothing in this file ever calls them - every
    # handler (handle_coding_mode, handle_writing_mode, handle_search_mode,
    # run_sr_launcher) does its own local import from a different submodule
    # path when it actually needs one. Importing pipelines.sr alone dragged in
    # scipy/matplotlib/pymupdf (~2.8s) on every run, including coding mode,
    # for a name that was never used (Known Issue #18). Only the two pipeline
    # imports actually referenced later in this file are kept below.
    from pipelines.appraisal import run_appraisal
    from pipelines.rct_search import run_rct_search_pipeline

    # Ensure project root is on sys.path BEFORE any imports
    _ROOT = Path(__file__).resolve().parent.parent
    if str(_ROOT) not in sys.path:
        sys.path.insert(0, str(_ROOT))

    # Import coding mode functions
    try:
        from pipelines.coding.coding import (
            run_builder,
            run_reviewer,
            run_tester,
            parse_direct_instructions,
        )
    except ModuleNotFoundError:
        from pipelines.coding.coding import (
            run_builder,
            run_reviewer,
            run_tester,
            parse_direct_instructions,
        )

    # Import RCT Search with fallback
    try:
        from pipelines.rct_search.rct_search import run_rct_search_pipeline
    except ModuleNotFoundError:
        try:
            from modes.rct_search import run_rct_search_pipeline
        except ModuleNotFoundError:
            # Define fallback if module doesn't exist
            def run_rct_search_pipeline(provider="deepseek", model=None, dry_run=False, reports_dir=None):
                print("[RCT Search] Module not available. Please ensure src/modes/rct_search.py exists.")
                return None

    load_dotenv()

    # -----------------------------------------------------------------------
    # Session 4 modules
    # -----------------------------------------------------------------------
    from streaming import stream_to_console, tee_stream
    from checkpoint import PipelineCheckpoint, find_resumable_checkpoint, prompt_resume
except KeyboardInterrupt:
    print("\n\nSession stopped. Returning to menu...\n")
    raise SystemExit(0)






# ---------------------------------------------------------------------------
# Version
# ---------------------------------------------------------------------------
VERSION = "2.4.13"

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE_DIR            = Path(__file__).resolve().parent.parent   # D:\AI_kcMedicalResearch

# Docs folders (guidance .md files per mode)
DOCS_DIR            = BASE_DIR / "docs"
DOCS_CODING         = DOCS_DIR / "coding"
DOCS_WRITING        = DOCS_DIR / "writing"
DOCS_APPRAISAL      = DOCS_DIR / "appraisal"
DOCS_RCT_SEARCH     = DOCS_DIR / "rct_search"
DOCS_SEARCH         = DOCS_DIR / "search"
DOCS_SR             = DOCS_DIR / "sr"

# AI prompt files
AI_DIR = BASE_DIR / "prompts"

# Input folders (auto-loaded at startup per mode)
INPUT_DIR           = BASE_DIR / "input"
INPUT_CODING        = INPUT_DIR / "coding"
INPUT_WRITING       = INPUT_DIR / "writing"
INPUT_APPRAISAL     = INPUT_DIR / "appraisal"
INPUT_RCT_SEARCH    = INPUT_DIR / "rct_search"
INPUT_SEARCH        = INPUT_DIR / "search"
INPUT_SR            = INPUT_DIR / "sr"

# Output folders (deliverables)
OUTPUT_DIR          = BASE_DIR / "output"
OUTPUT_CODING       = OUTPUT_DIR / "coding"
OUTPUT_WRITING      = OUTPUT_DIR / "writing"
OUTPUT_APPRAISAL    = OUTPUT_DIR / "appraisal"
OUTPUT_RCT_SEARCH   = OUTPUT_DIR / "rct_search"
OUTPUT_SEARCH       = OUTPUT_DIR / "search"
OUTPUT_SR           = OUTPUT_DIR / "sr"

# Reports folder (session transcripts / operation logs)
REPORTS_DIR         = BASE_DIR / "reports"

# Ensure all directories exist at startup
for _d in [
    DOCS_CODING, DOCS_WRITING, DOCS_APPRAISAL,
    DOCS_RCT_SEARCH, DOCS_SEARCH, DOCS_SR,
    INPUT_CODING, INPUT_WRITING, INPUT_APPRAISAL,
    INPUT_RCT_SEARCH, INPUT_SEARCH, INPUT_SR,
    OUTPUT_CODING, OUTPUT_WRITING, OUTPUT_APPRAISAL,
    OUTPUT_RCT_SEARCH, OUTPUT_SEARCH, OUTPUT_SR,
    REPORTS_DIR,
    # reports/ subfolders  - each mode writes here
    REPORTS_DIR / "coding",
    REPORTS_DIR / "writing",
    REPORTS_DIR / "appraisal",
    REPORTS_DIR / "rct_search",
    REPORTS_DIR / "search",
    REPORTS_DIR / "transcripts",
]:
    _d.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Auto-load input files
# ---------------------------------------------------------------------------
import mimetypes  # noqa: E402  (after path setup)

_MODE_EXTENSIONS = {
    "coding":     {".py", ".js", ".ts", ".html", ".css", ".java", ".c",
                   ".cpp", ".cs", ".rb", ".go", ".rs", ".txt", ".md",
                   ".php", ".swift", ".kt", ".r", ".sh", ".sql", ".svg"},
    "writing":    {".txt", ".md", ".docx", ".pdf"},
    "appraisal":  {".pdf", ".txt", ".md", ".docx"},
    "rct_search": {".txt", ".md", ".pdf", ".docx"},
    "search":     {".txt", ".md"},
    "sr":         {".pdf"},
}


def auto_load_input_files(mode: str) -> list[Path]:
    """
    Scan input/<mode>/ and return a sorted list of accepted files.
    Prints what was found. Returns empty list if folder is empty.
    """
    folder_map = {
        "coding":     INPUT_CODING,
        "writing":    INPUT_WRITING,
        "appraisal":  INPUT_APPRAISAL,
        "rct_search": INPUT_RCT_SEARCH,
        "search":     INPUT_SEARCH,
        "sr":         INPUT_SR,
    }
    folder  = folder_map.get(mode, INPUT_DIR / mode)
    allowed = _MODE_EXTENSIONS.get(mode, set())
    files   = sorted(
        p for p in folder.iterdir()
        if p.is_file() and p.suffix.lower() in allowed
    )
    if files:
        print(f"[auto-load] Found {len(files)} file(s) in input/{mode}/:")
        for f in files:
            print(f"   - {f.name}")
    else:
        print(f"[auto-load] No input files in input/{mode}/  - proceeding without pre-loaded context.")
    return files


# ---------------------------------------------------------------------------
# Environment / provider config -- sourced from providers.py + local overrides
# ---------------------------------------------------------------------------
from providers import (
    OLLAMA_HOST, OLLAMA_MODEL, OPENAI_API_KEY, OPENAI_MODEL,
    ANTHROPIC_API_KEY, ANTHROPIC_MODEL, DEEPSEEK_API_KEY, DEEPSEEK_MODEL,
    GROQ_API_KEY, GROQ_MODEL, DASHSCOPE_API_KEY, DASHSCOPE_BASE_URL,
    QWEN_MODEL,QWEN_VISION_MODEL,DEFAULT_PROVIDER,
)
EMBEDDING_PROVIDER = os.getenv("EMBEDDING_PROVIDER", "ollama")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "nomic-embed-text")



# ---------------------------------------------------------------------------
# Role / mode definitions
# ---------------------------------------------------------------------------
ROLE_FILES_CODING = {
    "Builder":  {"prompt": AI_DIR / "builder-prompt.md",  "report": "builder-report.md"},
    "Reviewer": {"prompt": AI_DIR / "reviewer-prompt.md", "report": "reviewer-report.md"},
    "Tester":   {"prompt": AI_DIR / "tester-prompt.md",   "report": "tester-report.md"},
}

ROLE_FILES_WRITING = {
    "Writer": {"prompt": AI_DIR / "writer-prompt.md", "report": "writer-report.md"},
    "Editor": {"prompt": AI_DIR / "editor-prompt.md", "report": "editor-report.md"},
    "QA":     {"prompt": AI_DIR / "qa-prompt.md",     "report": "qa-report.md"},
}

ROLE_FILES_RCT_SEARCH = {
    "Formulator": {"prompt": AI_DIR / "formulator-prompt.md", "report": "formulator-report.md"},
    "Searcher":   {"prompt": AI_DIR / "searcher-prompt.md",   "report": "searcher-report.md"},
    "Validator":  {"prompt": AI_DIR / "validator-prompt.md",  "report": "validator-report.md"},
}

ROLE_FILES_APPRAISAL = {
    "Appraiser": {
        "prompt": AI_DIR / "appraisal-prompt.md",
        "report": "appraisal-report.md",
    },
    "Methodologist": {
        "prompt": AI_DIR / "methodologist-prompt.md",
        "report": "methodologist-report.md",
    },
    "Summariser": {
        "prompt": AI_DIR / "summariser-prompt.md",
        "report": "summariser-report.md",
    },
}

ROLE_FILES_SEARCH = {
    "Researcher": {
        "prompt": AI_DIR / "researcher-prompt.md",
        "report": "researcher-report.md",
    },
}

ALL_MODES: dict[str, dict] = {
    "coding":     ROLE_FILES_CODING,
    "writing":    ROLE_FILES_WRITING,
    "rct_search": ROLE_FILES_RCT_SEARCH,
    "appraisal":  ROLE_FILES_APPRAISAL,
    "search":     ROLE_FILES_SEARCH,
}


# ---------------------------------------------------------------------------
# Role-specific documentation injection (least-privilege)
# ---------------------------------------------------------------------------
DOC_FILES_BY_ROLE: dict[str, list[Path]] = {
    "Builder":  [DOCS_CODING / "PRD.md",
                 DOCS_CODING / "architecture.md",
                 DOCS_CODING / "coding-standards.md"],
    "Reviewer": [DOCS_CODING / "PRD.md",
                 DOCS_CODING / "architecture.md",
                 DOCS_CODING / "decision-log.md"],
    "Tester":   [DOCS_CODING / "PRD.md",
                 DOCS_CODING / "architecture.md",
                 DOCS_CODING / "test-strategy.md"],
    "Writer":   [DOCS_WRITING / "project-brief.md",
                 DOCS_WRITING / "style-guide.md"],
    "Editor":   [DOCS_WRITING / "project-brief.md",
                 DOCS_WRITING / "editorial-standards.md"],
    "QA":       [DOCS_WRITING / "project-brief.md",
                 DOCS_WRITING / "qa-checklist.md"],
    "Formulator": [DOCS_RCT_SEARCH / "pico-framework.md"],
    "Searcher":   [DOCS_RCT_SEARCH / "pico-framework.md",
                   DOCS_RCT_SEARCH / "database-guide.md"],
    "Validator":  [DOCS_RCT_SEARCH / "pico-framework.md",
                   DOCS_RCT_SEARCH / "validation-criteria.md"],
    "Appraiser":     [],
    "Methodologist": [],
    "Summariser":    [],
    "Researcher":    [],
}


# ---------------------------------------------------------------------------
# Provider registry
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Provider functions -- delegated to providers.py (backward-compatible names)
# ---------------------------------------------------------------------------
from providers import (
    call_openai_provider,
    call_anthropic_provider,
    call_ollama_provider,
    call_deepseek_provider,
    call_groq_provider,
    call_qwen_provider,
    call_ai_with_fallback,
)


# PROVIDERS dict -- re-exported from providers.py
from providers import PROVIDERS



def call_ai(
    prompt: str,
    provider: str = "deepseek",
    model: str | None = None,
    stream: bool = False,
) -> str:
    """
    Dispatch an AI call to the correct provider function.
    If stream=True and running in a terminal, streams tokens to console.
    Uses fallback chain from FALLBACK_PROVIDERS env var on transient errors.
    """
    fallback_raw = os.getenv("FALLBACK_PROVIDERS", "deepseek,qwen,groq")
    fallback_chain = [p.strip() for p in fallback_raw.split(",") if p.strip()]

    if stream and sys.stdout.isatty():
        try:
            return stream_to_console(
                prompt=prompt,
                provider=provider,
                model=model,
            )
        except Exception:
            pass  # Fall through to non-streaming

    if fallback_chain:
        return call_ai_with_fallback(
            prompt=prompt,
            provider=provider,
            model=model,
            fallback_chain=fallback_chain,
        )

    fn = PROVIDERS.get(provider, call_ollama_provider)
    return fn(prompt, model=model)


# ---------------------------------------------------------------------------
# Context builder (with optional RAG)
# ---------------------------------------------------------------------------
def build_project_context(
    role_name: str,
    query: str = "",
    mode: str = "",
    session_id: str = "",
) -> str:
    """
    Build the project context string injected into every AI prompt.
    Reads small docs/ files for the role, then optionally appends RAG chunks.
    """
    doc_files = DOC_FILES_BY_ROLE.get(role_name, [])
    sections: list[str] = []

    for doc_path in doc_files:
        if doc_path.exists():
            content = doc_path.read_text(encoding="utf-8", errors="replace").strip()
            if content:
                sections.append(f"### {doc_path.name}\n{content}")

    if query and mode and session_id:
        try:
            from utils import rag
            rag_context = rag.retrieve(query, mode, session_id)
            if rag_context:
                sections.append(rag_context)
        except Exception as exc:  # noqa: BLE001
            print(f"[RAG] Retrieval warning: {exc}")

    return "\n\n".join(sections)


def truncate_context(text: str, max_chars: int = 2000) -> str:
    """Return text unchanged if within max_chars, otherwise truncate with ellipsis."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\u2026"


# ---------------------------------------------------------------------------
# File utilities
# ---------------------------------------------------------------------------
def read_text_file(path: Path) -> str:
    """Read a text file, strip whitespace. Returns '' if missing."""
    try:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    except (FileNotFoundError, OSError):
        return ""


def save_report(
    path: Path,
    role_name: str,
    model: str,
    task: str,
    response: str,
) -> None:
    """Append a role interaction to a markdown report file."""
    path.parent.mkdir(parents=True, exist_ok=True)
    entry = (
        f"\n## {role_name}\n"
        f"**Model:** {model}\n\n"
        f"**Task:** {task}\n\n"
        f"**Response:**\n{response}\n"
    )
    with path.open("a", encoding="utf-8") as fh:
        fh.write(entry)


def start_session_transcript(reports_dir: Path) -> Path:
    """Create a new timestamped session transcript file. Returns its Path."""
    reports_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path      = reports_dir / f"session_{timestamp}.md"
    path.write_text(
        f"# Session Transcript\nStarted: {timestamp}\n",
        encoding="utf-8",
    )
    return path


def append_to_transcript(
    path: Path,
    role_name: str,
    step: int,
    task: str,
    response: str,
) -> None:
    """Append one interaction step to an existing transcript file."""
    entry = (
        f"\n## Step {step}  - {role_name}\n"
        f"**Task:** {task}\n\n"
        f"**Response:**\n{response}\n"
    )
    with path.open("a", encoding="utf-8") as fh:
        fh.write(entry)


def print_session_summary(
    transcript_path: Path,
    step_count: int,
    role_counts: dict[str, int],
) -> None:
    """Print a summary of the completed session to stdout."""
    print(f"\n{'='*55}")
    print(f"  Session complete")
    print(f"  Transcript : {transcript_path.name}")
    print(f"  Total steps: {step_count}")
    if step_count == 0:
        print("  No steps recorded.")
    else:
        print("\n  Role usage:")
        for role, count in role_counts.items():
            if count > 0:
                print(f"    {role}: {count} step(s)")
    print(f"{'='*55}\n")


# ---------------------------------------------------------------------------
# Colour helpers
# ---------------------------------------------------------------------------
COLOURS = {
    "Builder":       "\033[94m",
    "Reviewer":      "\033[93m",
    "Tester":        "\033[92m",
    "Writer":        "\033[95m",
    "Editor":        "\033[96m",
    "QA":            "\033[91m",
    "Formulator":    "\033[94m",
    "Searcher":      "\033[92m",
    "Validator":     "\033[93m",
    "Appraiser":     "\033[95m",
    "Methodologist": "\033[96m",
    "Summariser":    "\033[92m",
    "Researcher":    "\033[94m",
}
RESET = "\033[0m"


def role_color(role_name: str) -> str:
    return COLOURS.get(role_name, RESET)


def choose_role(mode: str = "coding") -> tuple[str, dict]:
    """Prompt the user to pick a role and return (role_name, role_config)."""
    roles      = ALL_MODES[mode]
    role_names = list(roles.keys())

    print(f"\nSelect a role for {mode} mode:")
    for i, name in enumerate(role_names, start=1):
        colour = role_color(name)
        print(f"  {colour}{i}. {name}{RESET}")

    while True:
        try:
            choice = input("Enter number: ").strip()
        except (EOFError, KeyboardInterrupt):
            choice = "1"
            print("[no input detected - defaulting to role 1]")
        if choice.isdigit() and 1 <= int(choice) <= len(role_names):
            name = role_names[int(choice) - 1]
            return name, roles[name]
        print(f"Please enter a number between 1 and {len(role_names)}.")


# ---------------------------------------------------------------------------
# Session management
# ---------------------------------------------------------------------------
def list_sessions(reports_dir: str = str(REPORTS_DIR)) -> None:
    """Print all saved session transcript files, newest first."""
    path = Path(reports_dir)
    if not path.exists():
        print("No reports folder found.")
        return
    files = sorted(path.glob("session_*.md"), reverse=True)
    if not files:
        print("No session transcripts found.")
        return
    print(f"\n{'#':<4} {'Filename':<45} {'Size':>8}")
    print("-" * 60)
    for i, f in enumerate(files, start=1):
        print(f"{i:<4} {f.name:<45} {f.stat().st_size:>6} B")


def read_session(filename: str, reports_dir: str = str(REPORTS_DIR)) -> None:
    """Print the contents of a saved session transcript."""
    path = Path(reports_dir) / filename
    if not path.exists():
        print(f"File not found: {filename}. Use --list-sessions to see available files.")
        return
    print(f"--- {filename} ---")
    print(path.read_text(encoding="utf-8"))


def delete_session(filename: str, reports_dir: str = str(REPORTS_DIR)) -> None:
    """Delete a saved session transcript after confirmation."""
    path = Path(reports_dir) / filename
    if not path.exists():
        print(f"File not found: {filename}. Use --list-sessions to see available files.")
        return
    try:
        confirm = input(f"Delete '{filename}'? [y/N]: ").strip().lower()
    except (EOFError, KeyboardInterrupt):
        print()
        confirm = "n"
    if confirm == "y":
        path.unlink()
        print(f"Deleted: {filename}")
    else:
        print("Cancelled.")


def export_session(filename: str, reports_dir: str = str(REPORTS_DIR)) -> None:
    """Export a session transcript as a plain-text .txt file."""
    src = Path(reports_dir) / filename
    if not src.exists():
        print(f"File not found: {filename}. Use --list-sessions to see available files.")
        return
    raw     = src.read_text(encoding="utf-8")
    cleaned = re.sub(r"#{1,6}\s*", "", raw)
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*(.+?)\*",    r"\1", cleaned)
    txt_name = Path(filename).stem + ".txt"
    dest     = Path(reports_dir) / txt_name
    dest.write_text(cleaned, encoding="utf-8")
    print(f"Exported to: {dest}")


def rename_session(filename: str, reports_dir: str = str(REPORTS_DIR)) -> None:
    """Rename a session transcript file, appending .md automatically."""
    src = Path(reports_dir) / filename
    if not src.exists():
        print(f"File not found: {filename}. Use --list-sessions to see available files.")
        return
    try:
        raw_name = input("New filename (without extension): ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        raw_name = ""
    if not raw_name:
        print("Name cannot be empty. Cancelled.")
        return
    new_name = raw_name if raw_name.endswith(".md") else raw_name + ".md"
    dest = Path(reports_dir) / new_name
    if dest.exists():
        print(f"A file named '{new_name}' already exists. Cancelled.")
        return
    src.rename(dest)
    print(f"Renamed to: {new_name}")


def show_stats(reports_dir: str = str(REPORTS_DIR)) -> None:
    """Print session statistics for all saved transcripts."""
    path = Path(reports_dir)
    if not path.exists():
        print("No reports folder found.")
        return
    files = sorted(path.glob("session_*.md"))
    if not files:
        print("No session transcripts found.")
        return

    total_size  = 0
    role_counts: dict[str, int] = {}
    all_role_names = [r for mode_roles in ALL_MODES.values() for r in mode_roles]
    for role in all_role_names:
        role_counts[role] = 0

    for f in files:
        total_size += f.stat().st_size
        content     = f.read_text(encoding="utf-8", errors="replace")
        for role in all_role_names:
            role_counts[role] += content.count(f"## {role}")

    print(f"\nTotal sessions    : {len(files)}")
    print(f"Total size        : {total_size} bytes")
    print("\nRole usage across all sessions:")
    for role, count in role_counts.items():
        colour = role_color(role)
        print(f"  {colour}{role:<14}{RESET}: {count} interaction(s)")


# ---------------------------------------------------------------------------
# File readers
# ---------------------------------------------------------------------------
def _read_docx(path: Path) -> str:
    """Extract plain text from a .docx file using python-docx."""
    try:
        import docx as _docx
        doc = _docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:  # noqa: BLE001
        return f"[Could not read DOCX: {exc}]"


def _read_pdf_pymupdf(path: Path, max_chars: int = 30_000) -> str:
    """Extract plain text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz
        doc  = fitz.open(str(path))
        text = "\n\n".join(page.get_text() for page in doc)
        doc.close()
        return text[:max_chars]
    except Exception as exc:  # noqa: BLE001
        return f"[Could not read PDF: {exc}]"


def _md_to_docx(md_text: str, title: str, out_path: Path) -> None:
    """Convert a markdown report string to a .docx file using python-docx."""
    import docx as _docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document   = _docx.Document()
    title_para = document.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
        elif stripped.startswith("#### "):
            document.add_heading(stripped[5:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            document.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        elif stripped.startswith("**") and stripped.endswith("**"):
            p   = document.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        else:
            document.add_paragraph(stripped)

    document.save(str(out_path))

def _add_hyperlink(paragraph, text: str, url: str):
    """Add a clickable hyperlink run to an existing paragraph."""
    import docx as _docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    new_t = OxmlElement("w:t")
    new_t.text = text
    new_run.append(new_t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _ranked_articles_to_docx(
    ranked: list[dict],
    title: str,
    out_path: Path,
    topic: str = "",
) -> None:
    """
    Write a ranked article list as a proper Word table with clickable
    PubMed hyperlinks. Each row: Rank | Score | Title | PMID | Link.
    """
    import docx as _docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn

    doc = _docx.Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if topic:
        p = doc.add_paragraph()
        run = p.add_run(f"Topic: {topic}")
        run.italic = True

    doc.add_paragraph("")

    if not ranked:
        doc.add_paragraph("No articles retrieved from PubMed.")
        doc.save(str(out_path))
        return

    # Caption
    cap = doc.add_paragraph()
    cap.add_run(
        f"All {len(ranked)} RCT articles retrieved from PubMed, "
        "ordered by PICO relevance score (10 = most relevant)."
    ).italic = True

    doc.add_paragraph("")

    # Table: Rank | Score | Title | PMID | Link
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for i, label in enumerate(["Rank", "Score", "Title", "PMID", "PubMed Link"]):
        hdr[i].text = label
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True

    # Set column widths
    widths = [Cm(1.2), Cm(1.5), Cm(9.0), Cm(2.2), Cm(2.8)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    # Data rows
    for r in ranked:
        row = table.add_row()
        row.cells[0].text = str(r["rank"])
        row.cells[1].text = f"{r['score']}/10"
        row.cells[2].text = r["title"]
        row.cells[3].text = r["pmid"]
        # Clickable hyperlink in Link cell
        cell_para = row.cells[4].paragraphs[0]
        _add_hyperlink(cell_para, "PubMed", r["url"])

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.add_run(
        "Select your top articles, download PDFs and place them in "
        "input/sr/ to run the SR pipeline."
    ).italic = True

    doc.add_paragraph("")
    ref = doc.add_paragraph()
    ref.add_run(
        "For explanation on ranking, please refer to the full report "
        "in the reports folder."
    ).italic = True

    doc.save(str(out_path))

# ---------------------------------------------------------------------------
# Writing report
# ---------------------------------------------------------------------------

def generate_writing_report(
    docs_dir: Path = DOCS_WRITING,
    reports_dir: Path = REPORTS_DIR,
    provider: str = "deepseek",
    model: str | None = None,
    input_dir: Path | None = None,
) -> Path:

    """
    Read input files, send to AI with writing-report prompt, save as
    output/writing/writing_report_{ts}.md and .docx.
    Returns the path of the saved .md report.
    """
    prompt_path = AI_DIR / "writing-report-prompt.md"
    if not prompt_path.exists():
        raise FileNotFoundError(f"Writing report prompt not found: {prompt_path}")

    if input_dir is not None:
        _SUPPORTED = {".txt", ".md", ".pdf", ".docx"}
        files = sorted(
            p for p in input_dir.iterdir()
            if p.is_file() and p.suffix.lower() in _SUPPORTED
        ) if input_dir.exists() else []
    else:
        files = auto_load_input_files("writing")

    if not files and docs_dir.exists():
        SUPPORTED = {".txt", ".md", ".pdf", ".docx"}
        files = sorted(
            f for f in docs_dir.rglob("*")
            if f.is_file() and f.suffix.lower() in SUPPORTED
        )

    if not files:
        print(f"No files found in input/writing/ or {docs_dir}.")
        print("Add .txt, .md, .pdf, or .docx files to: input/writing/")
        return OUTPUT_WRITING / "writing_report_empty.md"

    system_prompt = prompt_path.read_text(encoding="utf-8", errors="replace")

    sections: list[str] = []
    for f in files:
        suffix = f.suffix.lower()
        if suffix == ".pdf":
            content = _read_pdf_pymupdf(f)
        elif suffix == ".docx":
            content = _read_docx(f)
        else:
            content = read_text_file(f)
        if content.strip():
            sections.append(f"### {f.name}\n{content.strip()}")
        else:
            print(f"  Warning: no readable content in {f.name}  - skipped.")

    if not sections:
        print("No readable content found in any file. Exiting.")
        return OUTPUT_WRITING / "writing_report_empty.md"

    combined    = "\n\n".join(sections)
    full_prompt = f"{system_prompt}\n\n## Documents to Summarise\n\n{combined}"

    print(f"Generating writing report from {len(sections)} file(s)...")
    try:
        response = call_ai(prompt=full_prompt, provider=provider, model=model)
    except RuntimeError as exc:
        response = f"[ERROR generating report: {exc}]"

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    OUTPUT_WRITING.mkdir(parents=True, exist_ok=True)

    md_path    = OUTPUT_WRITING / f"writing_report_{timestamp}.md"
    md_content = f"# Writing Report\nGenerated: {timestamp}\n\n{response}\n"
    md_path.write_text(md_content, encoding="utf-8")
    print(f"Markdown report saved : output\\writing\\{md_path.name}")

    docx_path = OUTPUT_WRITING / f"writing_report_{timestamp}.docx"
    try:
        _md_to_docx(response, f"Writing Report  - {timestamp}", docx_path)
        print(f"Word document saved   : output\\writing\\{docx_path.name}")
    except Exception as exc:  # noqa: BLE001
        print(f"  Warning: could not generate .docx  - {exc}")

    return md_path


# ---------------------------------------------------------------------------
# Code revision pipeline (--revise flag)
# ---------------------------------------------------------------------------
CODE_EXTENSIONS = {".py", ".js", ".ts", ".cs", ".java", ".sql", ".html", ".css"}
GUIDANCE_DOCS   = {
    "PRD.md", "architecture.md", "coding-standards.md",
    "decision-log.md", "test-strategy.md",
}


def _read_code_files(docs_coding: Path) -> list[dict]:
    """Return list of {name, content} for code files in docs/coding/."""
    results = []
    if not docs_coding.exists():
        return results
    for f in sorted(docs_coding.iterdir()):
        if (f.is_file()
                and f.suffix.lower() in CODE_EXTENSIONS
                and f.name not in GUIDANCE_DOCS):
            content = f.read_text(encoding="utf-8", errors="replace")
            if content.strip():
                results.append({"name": f.name, "content": content})
    return results


def generate_code_revision(
    start_role: str = "Builder",
    docs_dir: Path = DOCS_CODING,
    reports_dir: Path = REPORTS_DIR,
    provider: str = "deepseek",
    model: str | None = None,
    dry_run: bool = False,
) -> Path:
    """
    Single-pass code revision pipeline (legacy --revise flag).
    Builder -> Build -> Review -> Test
    Reviewer -> Review -> Test
    Tester -> Test only
    Saves results as reports/code_revision_{ts}.md and .docx.
    Returns path of the .md report.
    """
    all_stages = ["Builder", "Reviewer", "Tester"]
    start_idx  = all_stages.index(start_role) if start_role in all_stages else 0
    stages     = all_stages[start_idx:]

    code_files = _read_code_files(docs_dir)
    if not code_files:
        print(f"No code files found in {docs_dir}.")
        print(f"Add .py/.js/.ts/.cs/.java/.sql/.html/.css files to: {docs_dir}")
        return reports_dir / "code_revision_empty.md"

    print(f"\n{'='*55}")
    print(f"  CODE REVISION PIPELINE (--revise)")
    print(f"  Starting role : {start_role}")
    print(f"  Pipeline      : {' -> '.join(stages)}")
    print(f"  Code files    : {len(code_files)}")
    print(f"{'='*55}\n")

    code_context = "\n\n".join(
        f"### {cf['name']}\n`\n{cf['content']}\n`"
        for cf in code_files
    )

    guidance_parts = []
    for doc_name in GUIDANCE_DOCS:
        doc_path = docs_dir / doc_name
        if doc_path.exists():
            guidance_parts.append(
                f"### {doc_name}\n"
                f"{doc_path.read_text(encoding='utf-8', errors='replace').strip()}"
            )
    guidance_context = "\n\n".join(guidance_parts)

    try:
        task = input("Describe the revision task (or press Enter for general review): ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        task = ""
    if not task:
        task = "Review and improve the code quality, readability, and correctness."

    reports_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    md_path   = reports_dir / f"code_revision_{timestamp}.md"
    docx_path = reports_dir / f"code_revision_{timestamp}.docx"

    full_report_parts = [
        "# Code Revision Report",
        f"Generated: {timestamp}",
        f"Pipeline: {' -> '.join(stages)}",
        f"Task: {task}",
        "",
    ]

    previous_response = ""

    for stage in stages:
        print(f"Running {stage}...")
        role_cfg    = ALL_MODES["coding"][stage]
        prompt_path = Path(role_cfg["prompt"])
        role_prompt = prompt_path.read_text(encoding="utf-8", errors="replace")

        parts = [role_prompt]
        if guidance_context:
            parts.append(f"## Project Guidance\n{guidance_context}")
        parts.append(f"## Code Files\n{code_context}")
        if previous_response:
            parts.append(f"## Previous Stage Output\n{previous_response}")
        parts.append(f"## Revision Task\n{task}")
        full_prompt = "\n\n".join(parts)

        if dry_run:
            response = f"[DRY RUN] {stage} would respond here."
        else:
            try:
                response = call_ai(prompt=full_prompt, provider=provider, model=model)
            except RuntimeError as exc:
                response = f"[ERROR in {stage}: {exc}]"

        previous_response = response
        full_report_parts.append(f"## {stage} Output\n\n{response}\n")
        print(f"{stage} complete.\n")

    md_content = "\n".join(full_report_parts)
    md_path.write_text(md_content, encoding="utf-8")
    print(f"Markdown report saved : reports\\{md_path.name}")

    try:
        _md_to_docx(md_content, f"Code Revision Report  - {timestamp}", docx_path)
        print(f"Word document saved   : reports\\{docx_path.name}")
    except Exception as exc:  # noqa: BLE001
        print(f"  Warning: could not generate .docx  - {exc}")

    return md_path


# ---------------------------------------------------------------------------
# Article helpers (appraisal mode)
# ---------------------------------------------------------------------------
ARTICLE_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
ARTICLE_SIZE_LIMIT = 8000


def _read_article_files(input_appraisal: Path) -> list[dict]:
    """Read appraisal article files from input/appraisal/."""
    results = []
    if not input_appraisal.exists():
        return results
    for f in sorted(input_appraisal.iterdir()):
        if f.suffix.lower() not in ARTICLE_EXTENSIONS:
            continue
        try:
            if f.suffix.lower() == ".pdf":
                text = _read_pdf_pymupdf(f)
            elif f.suffix.lower() == ".docx":
                text = _read_docx(f)
            else:
                text = f.read_text(encoding="utf-8", errors="replace")
            if not text.strip():
                continue
            if len(text) <= ARTICLE_SIZE_LIMIT:
                results.append({"name": f.name, "content": text})
            else:
                print(f"[Appraisal] {f.name} exceeds {ARTICLE_SIZE_LIMIT} chars  - passed to RAG.")
        except Exception as exc:  # noqa: BLE001
            print(f"[Appraisal] Could not read {f.name}: {exc}")
    return results


def _read_topic_file(topic_path: Path) -> str:
    if not topic_path.exists():
        return ""
    return topic_path.read_text(encoding="utf-8", errors="replace").strip()
    

def rct_search_reminder() -> None:
    """Print a reminder to edit the PICO framework before starting a search."""
    print("\n" + "=" * 55)
    print("  RCT SEARCH MODE")
    print("=" * 55)
    print("  IMPORTANT: Before proceeding, ensure you have edited")
    print("  your PICO framework file:")
    print(f"  {DOCS_RCT_SEARCH / 'pico-framework.md'}")
    print()
    print("  Population  : who are the patients?")
    print("  Intervention: what is being tested?")
    print("  Comparison  : what is the control/comparator?")
    print("  Outcome     : what are you measuring?")
    print()
    print("  Search links will be saved to the output/rct_search/ folder.")
    print("=" * 55 + "\n")
    try:
        confirm = input("Have you edited your PICO file? [y/N]: ").strip().lower()
    except (EOFError, KeyboardInterrupt):
        print()
        confirm = "n"
    if confirm != "y":
        print("Please edit the PICO file first, then re-run rct_search mode.")
        sys.exit(0)


def save_rct_search_links(
    response: str,
    reports_dir: Path = REPORTS_DIR,
) -> Path:
    """
    Extract URLs from an AI response and save them as a markdown link list.
    Returns the path of the saved file.
    """
    urls = re.findall(r'https?://\S+', response)
    urls = [u.rstrip(".,);\"'") for u in urls]

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path  = reports_dir / f"rct_search_{timestamp}.md"
    reports_dir.mkdir(parents=True, exist_ok=True)

    lines = [
        "# RCT Search Links",
        f"Generated: {timestamp}",
        "",
        "Download this file from the `reports/` folder.",
        "",
        "## Search Result Links",
        "",
    ]
    if urls:
        for url in urls:
            lines.append(f"- [{url}]({url})")
    else:
        lines.append("_No URLs found in the AI response._")

    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"\n[RCT Search] Links saved to: {out_path.name}")
    print(f"[RCT Search] Open reports\\{out_path.name} to download your links.\n")
    return out_path


# ---------------------------------------------------------------------------
# PubMed fetch
# ---------------------------------------------------------------------------
def fetch_pubmed_articles(
    query: str,
    max_results: int = 10,
) -> list[dict]:
    """
    Search PubMed via NCBI E-utilities (no API key required).
    Returns list of dicts: pmid, title, abstract, url.
    Returns empty list on any error.
    """
    base       = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    search_url = (
        f"{base}/esearch.fcgi?db=pubmed&term={urllib.parse.quote(query)}"
        f"&retmax={max_results}&retmode=json"
    )
    try:
        with urlopen(search_url, timeout=15) as resp:
            search_data = json.loads(resp.read())
        pmids = search_data["esearchresult"]["idlist"]
    except Exception:  # noqa: BLE001
        return []

    if not pmids:
        return []

    ids_str   = ",".join(pmids)
    fetch_url = (
        f"{base}/efetch.fcgi?db=pubmed&id={ids_str}"
        f"&rettype=abstract&retmode=xml"
    )
    try:
        with urlopen(fetch_url, timeout=15) as resp:
            xml_data = resp.read().decode("utf-8", errors="replace")
    except Exception:  # noqa: BLE001
        return []

    import xml.etree.ElementTree as ET
    try:
        root = ET.fromstring(xml_data)
    except ET.ParseError:
        return []

    articles = []
    for article in root.findall(".//PubmedArticle"):
        pmid_el     = article.find(".//PMID")
        title_el    = article.find(".//ArticleTitle")
        abstract_el = article.find(".//AbstractText")

        pmid     = pmid_el.text.strip()     if pmid_el     is not None else "unknown"
        title    = title_el.text.strip()    if title_el    is not None else "No title"
        abstract = abstract_el.text.strip() if abstract_el is not None else "No abstract available."

        title    = re.sub(r"\s+", " ", title)
        abstract = re.sub(r"\s+", " ", abstract)

        articles.append({
            "pmid":     pmid,
            "title":    title,
            "abstract": abstract,
            "url":      f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
        })

    return articles


# ---------------------------------------------------------------------------
# SR launcher
# ---------------------------------------------------------------------------

def run_sr_launcher(provider: str = "", model: str = "") -> None:
    """Run the SR automation pipeline (CLI mode  - no Streamlit)."""
    import shutil
    import subprocess as _sp
    import json as _json
    import yaml as _yaml
    from datetime import datetime

    print("\n" + "=" * 58)
    print("  SR Automation Pipeline")
    print("  PRISMA 2020  |  Cochrane Handbook v6.5")
    print("=" * 58 + "\n")

    # -- Step 1: find PDFs in input/sr/ ---------------------------------------
    pdf_files = sorted(INPUT_SR.glob("*.pdf"))
    if not pdf_files:
        print("[SR] No PDF files found in input/sr/. Aborting.")
        return

    print(f"[SR] Found {len(pdf_files)} PDF(s) in input/sr/:")
    for p in pdf_files:
        print(f"     {p.name}")

    # -- Step 2: PICO Management ----------------------------------------------
    print("\n" + "=" * 58)
    print("  PICO Configuration")
    print("=" * 58)
    
    # Find existing PICO JSON files
    pico_files = sorted(INPUT_SR.glob("pico_*.json"),
                        key=lambda f: f.stat().st_mtime, reverse=True)
    
    pico_data = None
    pico_path = None
    
    if pico_files:
        # Show existing PICO files
        print(f"\n  Found {len(pico_files)} PICO file(s) in input/sr/:")
        for idx, pf in enumerate(pico_files[:5], 1):
            print(f"    {idx}. {pf.name}")
        print(f"    0. Create new PICO")
        print()
        
        try:
            choice = input("  Select a PICO file to use (0 for new): ").strip()
        except (EOFError, KeyboardInterrupt):
            choice = "1"
        
        if choice.isdigit() and 1 <= int(choice) <= len(pico_files[:5]):
            pico_path = pico_files[int(choice) - 1]
            try:
                pico_data = _json.loads(pico_path.read_text(encoding="utf-8"))
                print(f"\n   - PICO loaded from: {pico_path.name}")
                print(f"     Population:   {pico_data.get('population', 'N/A')}")
                print(f"     Intervention: {pico_data.get('intervention', 'N/A')}")
                print(f"     Comparator:   {pico_data.get('comparator', 'N/A')}")
                print(f"     Outcome:      {pico_data.get('outcome', 'N/A')}")
                print(f"     Effect:       {pico_data.get('effect_measure', 'SMD')}")
                
                # Ask if user wants to modify
                try:
                    modify = input("\n  Modify this PICO? [y/N]: ").strip().upper()
                except (EOFError, KeyboardInterrupt):
                    modify = "N"
                
                if modify == "Y":
                    pico_data = _modify_pico(pico_data)
                    # Save modified version
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    new_path = INPUT_SR / f"pico_{timestamp}.json"
                    new_path.write_text(
                        _json.dumps(pico_data, indent=2, ensure_ascii=False),
                        encoding="utf-8"
                    )
                    pico_path = new_path
                    print(f"   - Modified PICO saved to: {new_path.name}")
                    
            except Exception as e:
                print(f"  [!] Could not load PICO: {e}")
                pico_data = None
        
        elif choice == "0":
            pico_data = _create_pico()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            pico_path = INPUT_SR / f"pico_{timestamp}.json"
            pico_path.write_text(
                _json.dumps(pico_data, indent=2, ensure_ascii=False),
                encoding="utf-8"
            )
            print(f"   - New PICO saved to: {pico_path.name}")
    
    if pico_data is None:
        # No PICO loaded - create new
        print("\n  No PICO configuration found.")
        try:
            create = input("  Create a new PICO configuration? [Y/n]: ").strip().upper()
        except (EOFError, KeyboardInterrupt):
            create = "Y"
        
        if create != "N":
            pico_data = _create_pico()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            pico_path = INPUT_SR / f"pico_{timestamp}.json"
            pico_path.write_text(
                _json.dumps(pico_data, indent=2, ensure_ascii=False),
                encoding="utf-8"
            )
            print(f"   - New PICO saved to: {pico_path.name}")
        else:
            print("  [i] No PICO configured. Using existing config.")
    
    # -- Step 3: Update prisma_criteria.yaml with PICO data -----------------
    cfg_yaml: dict = {}
    if pico_data:
        cfg_yaml = {
            "review_title":       pico_data.get("topic", pico_data.get("topic", "Systematic Review")),
            "effect_measure":     pico_data.get("effect_measure", "SMD"),
            "pico": {
                "population":     pico_data.get("population", ""),
                "intervention":   pico_data.get("intervention", ""),
                "comparator":     pico_data.get("comparator", ""),
                "outcome":        pico_data.get("outcome", ""),
            },
            "pubmed_query_cleaned": pico_data.get("pubmed_query_cleaned", ""),
            "inclusion_criteria": ["RCT", "Adult participants"],
            "exclusion_criteria": ["Non-RCT", "Animal studies"],
        }
        
        prisma_path = SOURCE_CODE_DIR / "pipelines" / "sr" / "config" / "prisma_criteria.yaml"
        prisma_path.parent.mkdir(parents=True, exist_ok=True)
        prisma_path.write_text(
            _yaml.dump(cfg_yaml, allow_unicode=True, sort_keys=False),
            encoding="utf-8",
        )
        print(f"\n   - prisma_criteria.yaml updated with PICO")
        print(f"     PubMed query: {cfg_yaml.get('pubmed_query_cleaned', 'n/a')}")
    else:
        print("\n  [i] Using existing prisma_criteria.yaml")

    # -- Step 4: resolve provider and model ------------------------------------

    _DEFAULT_MODELS = {
        "qwen":      QWEN_VISION_MODEL,
        "deepseek":  DEEPSEEK_MODEL,
        "openai":    OPENAI_MODEL,
        "anthropic": ANTHROPIC_MODEL,
        "groq":      GROQ_MODEL,
        "ollama":    OLLAMA_MODEL,
    }
    
    _provider = provider or "qwen"
    
    # --- Provider check for vision support ---
    if _provider == "deepseek":
        print("\n" + "=" * 58)
        print("  [!]  WARNING: DeepSeek does NOT support vision API")
        print("=" * 58)
        print("  The SR pipeline uses vision-based extraction (images of PDF pages).")
        print("  DeepSeek's API only accepts text, not images.")
        print()
        print("  Please use one of these providers with vision support:")
        print("     - qwen     (recommended) - Qwen vision model")
        print("     - openai   - GPT-4 vision")
        print("     - anthropic - Claude vision")
        print()
        print("  To switch provider, run:")
        print("    python SOURCE_CODE/main.py --mode sr --provider qwen")
        print("=" * 58 + "\n")
        
        try:
            choice = input("Continue anyway? (This will likely fail) [y/N]: ").strip().upper()
            if choice != "Y":
                print("Aborting. Please use --provider qwen")
                return
        except (EOFError, KeyboardInterrupt):
            print("\nAborting.")
            return
    

    _model = model or _DEFAULT_MODELS.get(_provider, QWEN_VISION_MODEL)


    # -- Step 5: run sr/main.py (as a module so relative imports resolve) -----
    cmd = [
        sys.executable, "-m", "SOURCE_CODE.pipelines.sr.main",
        "--pdf-dir",        str(INPUT_SR),
        "--provider",       _provider,
        "--model",          _model,
        "--effect-measure", cfg_yaml.get("effect_measure", "SMD")
                            if pico_data else "SMD",
    ]
    print(f"\n[SR] Running: {' '.join(cmd)}\n")
    _run_started = time.time() - 2  # small slack for filesystem timestamp granularity
    result = _sp.run(cmd, cwd=str(BASE_DIR))


    if result.returncode == 0:
        mirror_reports = BASE_DIR / "output" / "sr" / "reports"
        mirror_figures = BASE_DIR / "output" / "sr" / "figures"
        runs_root      = BASE_DIR / "reports" / "sr"
        print("\n[SR] Pipeline complete.")
        # Only advertise artifacts this run actually produced. The PDF is
        # legitimately absent when WeasyPrint is unavailable (README #2),
        # and the mirror dir persists across runs, so also require the
        # file to be newer than this run's start - a bare exists() would
        # present a stale copy from an earlier run as this run's output.
        _artifacts = [
            ("DOCX", mirror_reports / "systematic_review.docx"),
            ("PDF",  mirror_reports / "systematic_review.pdf"),
            ("HTML", mirror_reports / "systematic_review.html"),
            ("Plot", mirror_figures / "forest_plot.png"),
        ]
        for _label, _path in _artifacts:
            if _path.exists() and _path.stat().st_mtime >= _run_started:
                print(f"[SR] {_label:<7}-> {_path}")
            elif _path.exists():
                print(f"[SR] {_label:<7}-> not generated by this run "
                      f"(stale copy from an earlier run: {_path})")
            else:
                print(f"[SR] {_label:<7}-> not generated "
                      f"(see pipeline log above)")
        print(f"[SR] Full run (with audit CSVs) -> {runs_root}\\<run_id>\\")
    else:
        print(f"\n[SR] Pipeline exited with code {result.returncode}.")
        print("[SR] Check the output above for errors.")



def _create_pico() -> dict:
    """Interactive PICO creation."""
    print("\n  Enter PICO components for your systematic review:")
    print("  (Press Enter to skip any field)")
    
    population = input("    Population   (P): ").strip()
    intervention = input("    Intervention (I): ").strip()
    comparator = input("    Comparator   (C): ").strip()
    outcome = input("    Outcome      (O): ").strip()
    
    topic = f"Effect of {intervention} on {outcome}" if intervention and outcome else "Systematic Review"
    
    pico_data = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "topic": topic,
        "population": population,
        "intervention": intervention,
        "comparator": comparator,
        "outcome": outcome,
        "study_design": "RCT",
        "effect_measure": "SMD",
        "source_mode": "manual_sr",
        "pubmed_query_raw": f"{population} AND {intervention} AND {outcome}" if all([population, intervention, outcome]) else "",
        "pubmed_query_cleaned": "",
    }
    
    return pico_data


def _modify_pico(pico_data: dict) -> dict:
    """Allow user to modify existing PICO."""
    print("\n  Modify PICO components (press Enter to keep current value):")
    
    fields = [
        ("population", "Population   (P)"),
        ("intervention", "Intervention (I)"),
        ("comparator", "Comparator   (C)"),
        ("outcome", "Outcome      (O)"),
    ]
    
    for key, label in fields:
        current = pico_data.get(key, "")
        new_val = input(f"    {label} [{current}]: ").strip()
        if new_val:
            pico_data[key] = new_val
    
    # Update topic if intervention or outcome changed
    if pico_data.get("intervention") and pico_data.get("outcome"):
        pico_data["topic"] = f"Effect of {pico_data['intervention']} on {pico_data['outcome']}"
    
    return pico_data


def list_roles(mode: str = "coding") -> None:
    """Print each role's prompt file and injected documentation for a mode."""
    roles = ALL_MODES.get(mode, {})
    print(f"\n{'='*55}")
    print(f"  Roles for mode: {mode}")
    print(f"{'='*55}")
    for role_name, role_cfg in roles.items():
        colour    = role_color(role_name)
        prompt    = Path(role_cfg["prompt"]).relative_to(BASE_DIR)
        doc_files = DOC_FILES_BY_ROLE.get(role_name, [])
        doc_names = [d.name for d in doc_files]
        print(f"\n{colour}{role_name}{RESET}")
        print(f"  Prompt : {prompt}")
        print(f"  Docs   : {', '.join(doc_names) if doc_names else 'none'}")
    print(f"\n{'='*55}\n")

# ---------------------------------------------------------------------------
# Coding mode dispatcher  (Builder pipeline / Reviewer / Tester standalone)
# ---------------------------------------------------------------------------
def handle_coding_mode(
    provider: str = "deepseek",
    model: str | None = None,
    dry_run: bool = False,
) -> None:
    """
    Terminal entry point for Coding mode.
    Presents Builder / Reviewer / Tester sub-mode menu, collects > direct
    instructions from the user, then delegates to the coding.py engine.
    """

    def _call_llm_fn(system_prompt: str, user_prompt: str) -> str:
        if dry_run:
            return "[DRY RUN] LLM would respond here."
        combined = (
            f"## System Instructions\n{system_prompt}\n\n"
            f"## User Request\n{user_prompt}"
        )
        return call_ai(prompt=combined, provider=provider, model=model)

    print("\n" + "=" * 60)
    print("  CODING MODE")
    print("=" * 60)
    print(f"  {COLOURS['Builder']}1. Builder {RESET} "
          f" - pipeline: build  - review  - test  - output/coding/")
    print(f"  {COLOURS['Reviewer']}2. Reviewer{RESET} "
          f" - standalone: review code in input/coding/")
    print(f"  {COLOURS['Tester']}3. Tester  {RESET} "
          f" - standalone: test code in input/coding/")
    print("  0. Back to menu")
    print("=" * 60)

    try:
        choice = input("Select sub-mode [0-3]: ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        choice = "0"

    if choice == "0" or not choice:
        print("Returning to menu.")
        return

    sub_mode_map = {"1": "Builder", "2": "Reviewer", "3": "Tester"}
    if choice not in sub_mode_map:
        print("Invalid choice  - returning to menu.")
        return

    sub_mode = sub_mode_map[choice]
    colour   = COLOURS.get(sub_mode, RESET)

    print(f"\n{colour}[{sub_mode.upper()}]{RESET} Enter instructions below.")
    print("  Lines starting with > are DIRECT TASK INSTRUCTIONS (highest priority).")
    print("  Press ENTER on a blank line when done.\n")

    lines: list[str] = []
    while True:
        try:
            line = input("  instruction> ")
        except (EOFError, KeyboardInterrupt):
            break
        if line.strip() == "":
            print("\n   - Instructions received  - starting processing...\n")
            break
        if not line.strip().startswith(">"):
            line = "> " + line.strip()
        lines.append(line)

    raw_input_text      = "\n".join(lines)
    direct_instructions = parse_direct_instructions(raw_input_text)

    if direct_instructions:
        print(f"\n{colour}[{sub_mode.upper()}]{RESET} "
              f"Direct instructions captured ({len(direct_instructions)}):")
        for instr in direct_instructions:
            print(f"   - {instr}")
    else:
        print(f"\n{colour}[{sub_mode.upper()}]{RESET} "
              "No direct instructions found  - using docs/coding/ guidelines only.")

    auto_load_input_files("coding")
    print()

    if sub_mode == "Builder":
        run_builder(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            verbose=True,
        )
    elif sub_mode == "Reviewer":
        run_reviewer(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            verbose=True,
        )
    elif sub_mode == "Tester":
        run_tester(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            verbose=True,
        )

    # Clear checkpoint on successful completion
    try:
        from checkpoint import PipelineCheckpoint
        cp = PipelineCheckpoint(mode="coding", provider=provider)
        cp.clear()
    except Exception:
        pass

    print(f"\n{colour}[{sub_mode.upper()}]{RESET} Done. "
          "Check output/coding/ and reports/coding/ for results.\n")

# ---------------------------------------------------------------------------
# Appraisal mode dispatcher
# ---------------------------------------------------------------------------
def handle_appraisal_mode(
    provider: str = "deepseek",
    model: str | None = None,
    dry_run: bool = False,
) -> None:
    """Terminal entry point for Appraisal mode."""
    from pipelines.appraisal.appraisal import run_appraisal, parse_direct_instructions as _pdi

    def _call_llm_fn(system_prompt: str, user_prompt: str) -> str:
        if dry_run:
            return "[DRY RUN] LLM would respond here."
        return call_ai(
            prompt=f"## System Instructions\n{system_prompt}\n\n## User Request\n{user_prompt}",
            provider=provider,
            model=model,
        )

    print("\n" + "=" * 60)
    print("  APPRAISAL MODE")
    print("=" * 60)
    print("  Place article(s) in input/appraisal/ before starting.")
    print("  Supported formats: PDF, DOCX, MD, TXT")
    print("  Output   - output/appraisal/")
    print("  Reports  - reports/appraisal/")
    print("=" * 60)

    lines: list[str] = []
    print("\n  Enter optional direct instructions below.")
    print("  Lines starting with > are highest priority.")
    print("  Press ENTER on a blank line to start.\n")
    while True:
        line = input("  instruction> ").strip()
        if not line:
            print("\n  Instructions received -- starting appraisal...\n")
            break
        if not line.startswith(">"):
            line = "> " + line
        lines.append(line)

    direct = _pdi("\n".join(lines))
    run_appraisal(direct, _call_llm_fn, verbose=True)


# ---------------------------------------------------------------------------
# Writing mode dispatcher  (Writer pipeline / Editor / QA standalone)
# ---------------------------------------------------------------------------

def handle_search_mode(provider: str, model: str, sub_mode: str = None) -> None:
    """Dispatcher for Search mode (Topic Search and Article Search)."""
    import os
    from pipelines.search.search import run_topic_search, run_article_search

    print("\n=== SEARCH MODE ===")
    print("  1. Topic Search  (web synopsis with reference links)")
    print("  2. Article Search (PubMed by article type + comparison)")
    
    # Determine sub-mode
    is_cloud = any([
        os.environ.get('RENDER'),
        os.environ.get('STREAMLIT_SERVER_PORT'),
        os.environ.get('STREAMLIT_SHARING'),
    ])
    
    if sub_mode:
        sub = sub_mode
        print(f"   - Using sub-mode: {sub}")
    elif is_cloud:
        # Cloud environment - use default
        print("  [i]  Cloud environment detected. Defaulting to Topic Search (1).")
        sub = "1"
    else:
        # Local - interactive
        try:
            sub = input("  Select sub-mode [1-2]: ").strip()
            if not sub:
                sub = "1"
        except (EOFError, KeyboardInterrupt):
            print("\n  [i] No input received. Defaulting to Topic Search (1).")
            sub = "1"

    if sub not in ("1", "2"):
        print("  Invalid selection. Defaulting to Topic Search (1).")
        sub = "1"

    # Collect direct instructions
    print("\n  Enter your search query below.")
    print("  Lines starting with > are instructions (auto-prefixed if omitted).")
    print("  Press ENTER on a blank line to start.\n")
    raw_lines: list[str] = []
    while True:
        try:
            line = input("  > ")
        except (EOFError, KeyboardInterrupt):
            break
        if line.strip() == "":
            print("\n  Instructions received  - starting search...\n")
            break
        if not line.strip().startswith(">"):
            line = "> " + line.strip()
        raw_lines.append(line)
    direct_instructions: list[str] = [
        l[1:].strip() for l in raw_lines if l.strip().startswith(">") and l[1:].strip()
    ]

    def call_llm_fn(system_prompt: str = "", user_prompt: str = "") -> str:
        combined = f"{system_prompt}\n\n{user_prompt}" if system_prompt else user_prompt
        return call_ai(
            prompt=combined,
            provider=provider,
            model=model,
        )

    if sub == "1":
        run_topic_search(
            direct_instructions=direct_instructions,
            call_llm_fn=call_llm_fn,
            verbose=True,
            model=model,
        )
    else:
        run_article_search(
            direct_instructions=direct_instructions,
            call_llm_fn=call_llm_fn,
            verbose=True,
            model=model,
        )

def handle_writing_mode(
    provider: str = "deepseek",
    model: str | None = None,
    dry_run: bool = False,
) -> None:
    """Terminal entry point for Writing mode (two-track: topic / article)."""
    try:
        from pipelines.writing.writing import (
            run_writer, run_editor, run_qa,
            parse_direct_instructions as _pdi,
            TRACK_TOPIC, TRACK_ARTICLE, DEFAULT_WORDS,
        )
    except ModuleNotFoundError:
        from pipelines.writing.writing import (
            run_writer, run_editor, run_qa,
            parse_direct_instructions as _pdi,
            TRACK_TOPIC, TRACK_ARTICLE, DEFAULT_WORDS,
        )

    def _call_llm_fn(system_prompt: str, user_prompt: str) -> str:
        if dry_run:
            return "[DRY RUN] LLM would respond here."
        combined = (
            f"## System Instructions\n{system_prompt}\n\n"
            f"## User Request\n{user_prompt}"
        )
        return call_ai(prompt=combined, provider=provider, model=model)

    # Track selection
    print("\n" + "=" * 60)
    print("  WRITING MODE")
    print("=" * 60)
    print("  1. Topic Track     - editorial / opinion (newspaper style)")
    print("  2. Article Track   - medical journal article")
    print("  0. Back to menu")
    print("=" * 60)
    try:
        track_choice = input("Select track [0-2]: ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        return
    if track_choice == "0" or not track_choice:
        print("Returning to menu.")
        return
    if track_choice == "1":
        track = TRACK_TOPIC
    elif track_choice == "2":
        track = TRACK_ARTICLE
    else:
        print("Invalid choice  - returning to menu.")
        return

    # Sub-mode selection
    print(f"\n  Track: {track.upper()}")
    print("  " + "-" * 40)
    print("  1. Writer   - full pipeline: Writer -> Editor -> QA")
    print("  2. Editor   - standalone: edit documents in input/writing/")
    print("  3. QA       - standalone: review documents in input/writing/")
    print("  0. Back")
    try:
        sub_choice = input("Select sub-mode [0-3]: ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        return
    if sub_choice == "0" or not sub_choice:
        print("Returning to menu.")
        return
    sub_mode_map = {"1": "Writer", "2": "Editor", "3": "QA"}
    if sub_choice not in sub_mode_map:
        print("Invalid choice  - returning to menu.")
        return
    sub_mode = sub_mode_map[sub_choice]

    # Word limit
    if sub_mode in ("Writer", "Editor"):
        default_wl = DEFAULT_WORDS[track]
        print(f"\n  Default word limit for {track} track: {default_wl}")
        try:
            wl_input = input(
                f"  Press Enter to keep {default_wl}, or type a number: "
            ).strip()
        except (EOFError, KeyboardInterrupt):
            wl_input = ""
        word_limit = int(wl_input) if wl_input.isdigit() and int(wl_input) > 0 \
            else default_wl
        print(f"  Word limit set to: {word_limit}\n")
    else:
        word_limit = DEFAULT_WORDS[track]

    # Instructions
    print(f"  [{sub_mode.upper()} | {track.upper()}] Enter instructions below.")
    print("  Lines starting with > are DIRECT TASK INSTRUCTIONS (highest priority).")
    print("  Press ENTER on a blank line when done.\n")

    lines: list[str] = []
    while True:
        try:
            line = input("  instruction> ")
        except (EOFError, KeyboardInterrupt):
            break
        if line.strip() == "":
            print("\n  Instructions received -- starting processing...\n")
            break
        if not line.strip().startswith(">"):
            line = "> " + line.strip()
        lines.append(line)

    direct_instructions = _pdi("\n".join(lines))
    print()

    if sub_mode == "Writer":
        run_writer(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            track=track,
            word_limit=word_limit,
            verbose=True,
        )
    elif sub_mode == "Editor":
        run_editor(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            track=track,
            word_limit=word_limit,
            verbose=True,
        )
    elif sub_mode == "QA":
        run_qa(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            track=track,
            verbose=True,
        )

    print(f"\n  [{sub_mode.upper()} | {track.upper()}] Done. "
          "Check output/writing/ and reports/writing/ for results.\n")


# ---------------------------------------------------------------------------
# Main interactive session loop
# ---------------------------------------------------------------------------
def main(
    model_override: str | None = None,
    dry_run: bool = False,
    mode: str = "coding",
    provider: str = "deepseek",
) -> None:
    """Run an interactive AI session."""
    session_id  = uuid.uuid4().hex[:8]
    transcript  = start_session_transcript(REPORTS_DIR)
    step_count  = 0
    role_counts: dict[str, int] = {
        r: 0 for mode_roles in ALL_MODES.values() for r in mode_roles
    }

    print(f"\n{'='*55}")
    print(f"  AI Automation Tool  v{VERSION}")
    print(f"  Mode    : {mode}")
    print(f"  Provider: {provider}")
    print(f"  Session : {session_id}")
    if dry_run:
        print("  DRY RUN -- AI calls will be skipped")
    print(f"{'='*55}\n")

    input_files  = auto_load_input_files(mode)
    rag_enabled  = False
    input_folder = INPUT_DIR / mode
    input_folder.mkdir(parents=True, exist_ok=True)

    if not dry_run:
        try:
            from utils import rag as rag_module
            n_chunks = rag_module.index_uploads(
                mode=mode,
                session_id=session_id,
                upload_base=str(INPUT_DIR),
            )
            if n_chunks > 0:
                rag_enabled = True
                print(f"[RAG] Indexed {n_chunks} chunk(s) from input/{mode}/\n")
        except Exception as exc:  # noqa: BLE001
            print(f"[RAG] Indexing skipped: {exc}\n")

    role_name, role_cfg = choose_role(mode)
    colour      = role_color(role_name)
    prompt_path = Path(role_cfg["prompt"])

    if not prompt_path.exists():
        print(f"Prompt file not found: {prompt_path}")
        sys.exit(1)

    role_prompt = prompt_path.read_text(encoding="utf-8", errors="replace")
    print(f"\n{colour}Starting session as {role_name}{RESET}")
    print("Type your task below. Press Ctrl+C to end the session.\n")

    previous_response = ""

    article_context = ""
    if mode == "appraisal":
        articles = _read_article_files(INPUT_APPRAISAL)
        if articles:
            parts_art = [
                f"### Article: {art['name']}\n{art['content']}"
                for art in articles
            ]
            article_context = "\n\n".join(parts_art)
            print(f"[Appraisal] {len(articles)} article(s) loaded from input/appraisal/.")
        else:
            print("[Appraisal] No articles in input/appraisal/ -- using RAG only.")

    try:
        while True:
            try:
                task = input(f"{colour}{role_name} >{RESET} ").strip()
            except (EOFError, KeyboardInterrupt):
                print()
                break
            if not task:
                continue

            context = build_project_context(
                role_name=role_name,
                query=task if rag_enabled else "",
                mode=mode,
                session_id=session_id,
            )

            parts = [role_prompt]
            if mode == "appraisal" and article_context:
                parts.append(f"## Articles for Appraisal\n{article_context}")
            if context:
                parts.append(f"## Project Context\n{context}")
            if previous_response:
                parts.append(
                    f"## Previous Response\n{truncate_context(previous_response)}"
                )
            parts.append(f"## Task\n{task}")
            full_prompt = "\n\n".join(parts)

            start = time.time()
            if dry_run:
                response = f"[DRY RUN] {role_name} would respond here."
            else:
                try:
                    response = call_ai(
                        prompt=full_prompt,
                        provider=provider,
                        model=model_override,
                        stream=(not getattr(args, 'no_stream', False)) if 'args' in dir() else True,
                    )
                except RuntimeError as exc:
                    response = f"[ERROR] {exc}"

            elapsed     = time.time() - start
            step_count += 1
            role_counts[role_name] = role_counts.get(role_name, 0) + 1

            print(f"\n{colour}--- {role_name} response ({elapsed:.1f}s) ---{RESET}")
            print(response)
            print()

            append_to_transcript(
                path=transcript,
                role_name=role_name,
                step=step_count,
                task=task,
                response=response,
            )
            previous_response = response

            if mode == "rct_search" and not dry_run:
                save_rct_search_links(response=response, reports_dir=REPORTS_DIR / "rct_search")


    except KeyboardInterrupt:
        print(f"\n\n{colour}Session ended.{RESET}")

    finally:
        print_session_summary(transcript, step_count, role_counts)
        if rag_enabled and not dry_run:
            try:
                from utils import rag as rag_module
                rag_module.clear_session(mode=mode, session_id=session_id)
                print("[RAG] Session collection cleared.")
            except Exception as exc:  # noqa: BLE001
                print(f"[RAG] Clear warning: {exc}")


# ---------------------------------------------------------------------------
# API key validation
# ---------------------------------------------------------------------------
PROVIDER_ENV_VARS = {
    "anthropic": "ANTHROPIC_API_KEY",
    "openai":    "OPENAI_API_KEY",
    "deepseek":  "DEEPSEEK_API_KEY",
    "groq":      "GROQ_API_KEY",
    "qwen":      "DASHSCOPE_API_KEY",
    "ollama":    None,
}


def validate_api_keys(provider: str) -> None:
    """
    Validate the required API key is set for the selected provider.
    Raises EnvironmentError if missing. Ollama requires no key.
    """
    provider = provider.lower().strip()
    if provider not in PROVIDER_ENV_VARS:
        raise ValueError(
            f"Unknown provider '{provider}'. "
            f"Valid options: {', '.join(PROVIDER_ENV_VARS.keys())}"
        )
    env_var = PROVIDER_ENV_VARS[provider]
    if env_var is None:
        return
    value = os.environ.get(env_var, "").strip()
    if not value:
        raise EnvironmentError(
            f"\n"
            f"  [API KEY ERROR]\n"
            f"  Provider '{provider}' requires {env_var} to be set.\n"
            f"  Add it to your .env file in the project root:\n"
            f"  {env_var}=your-key-here\n"
        )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def open_help_guide() -> None:
    """Open the HTML flashcard help guide in the default browser."""
    import webbrowser
    guide_path = PROJECT_ROOT / "Readme" / "flashcard-help.html"
    if not guide_path.exists():
        print(f"Help guide not found: {guide_path}")
        print("Expected location: docs/flashcard-help.html")
        sys.exit(1)
    webbrowser.open(guide_path.as_uri())
    print(f"Opening help guide in browser: {guide_path.name}")


def parse_args(args: list[str] | None = None) -> argparse.Namespace:
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(
        description="AI Automation Tool",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python SOURCE_CODE/main.py                           # coding session\n"
            "  python SOURCE_CODE/main.py --mode writing            # writing mode\n"
            "  python SOURCE_CODE/main.py --mode rct_search         # RCT search mode\n"
            "  python SOURCE_CODE/main.py --mode appraisal          # appraisal mode\n"
            "  python SOURCE_CODE/main.py --mode search             # medical search\n"
            "  python SOURCE_CODE/main.py --mode sr                 # SR pipeline\n"
            "  python SOURCE_CODE/main.py --model llama3.2:3b       # different model\n"
            "  python SOURCE_CODE/main.py --provider qwen           # use Qwen\n"
            "  python SOURCE_CODE/main.py --list-sessions           # list transcripts\n"
            "  python SOURCE_CODE/main.py --list-roles              # show roles/docs\n"
            "  python SOURCE_CODE/main.py --list-roles --mode rct_search\n"
            "  python SOURCE_CODE/main.py --dry-run                 # simulate session\n"
            "  python SOURCE_CODE/main.py --version                 # show version\n"
        ),
    )
    parser.add_argument("--model",          type=str, default=None)
    parser.add_argument("--mode",           type=str, default="coding",
                        choices=["coding", "writing", "rct_search",
                                 "appraisal", "search", "sr"])
    parser.add_argument("--report",         action="store_true", default=False,
                        help="Generate a writing report from input/writing/ files")
    parser.add_argument("--revise",         action="store_true", default=False,
                        help="Run legacy code revision pipeline from docs/coding/")
    parser.add_argument("--role",           type=str, default="Builder",
                        choices=["Builder", "Reviewer", "Tester"],
                        help="Starting role for --revise pipeline (default: Builder)")
    parser.add_argument("--provider",       type=str, default=DEFAULT_PROVIDER,
                        choices=["ollama", "openai", "anthropic", "deepseek", "groq", "qwen"])
    parser.add_argument("--list-sessions",  action="store_true", default=False)
    parser.add_argument("--read-session",   type=str, default=None, metavar="FILENAME")
    parser.add_argument("--delete-session", type=str, default=None, metavar="FILENAME")
    parser.add_argument("--export-session", type=str, default=None, metavar="FILENAME")
    parser.add_argument("--rename-session", type=str, default=None, metavar="FILENAME")
    parser.add_argument("--stats",          action="store_true", default=False)
    parser.add_argument("--dry-run",        action="store_true", default=False)
    parser.add_argument("--no-stream",      action="store_true", default=False,
                        help="Disable streaming (show output all at once)")
    parser.add_argument("--resume",         action="store_true", default=False,
                        help="Resume from last checkpoint if available")
    parser.add_argument("--version",        action="version",
                        version=f"AI Automation Tool v{VERSION}")
    parser.add_argument("--list-roles",     action="store_true", default=False)
    parser.add_argument("--help-guide",     action="store_true", default=False,
                        help="Open the interactive HTML help guide in your browser")
    parser.add_argument("--ui",             action="store_true", default=False,
                        help="Launch the main Streamlit UI")
    parser.add_argument('--sub', '--submode', type=str, default=None,
                    help='Sub-mode selection: 1=Topic Search, 2=Article Search')
    return parser.parse_args(args)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
def run_cli(args: argparse.Namespace) -> None:
    """Dispatch parsed CLI arguments to the correct handler.

    Extracted from the ``if __name__ == "__main__"`` block so mode
    routing is testable (issue #43): SR mode must route to
    run_sr_launcher() and never reach main()/choose_role(), because
    ALL_MODES has no "sr" key and choose_role("sr") would raise
    KeyError.

    KeyboardInterrupt is intentionally NOT caught here; the
    entry-point guard below owns the clean-exit message.
    """

    # Anthropic-path tripwire gap (#61): _extract_anthropic and
    # assess_by_file_id bypass the source-quote, SD/SE, and
    # group/timepoint checks that the qwen path runs. Warn at point
    # of use until the checks are ported to the Anthropic path.
    if args.mode == "sr" and args.provider == "anthropic":
        print(
            "\n[!] Warning: --provider anthropic runs SR extraction without "
            "the source-quote, SD/SE, or group-timepoint tripwires.\n"
            "    See Known Issue #61. For full tripwire coverage use "
            "--provider qwen (default vision model).\n"
            "    Manual verification against source PDFs remains required "
            "either way.\n"
        )

    if args.ui:
        import subprocess as _sp
        proc = _sp.Popen(
            [sys.executable, "-m", "streamlit", "run",
             str(SOURCE_CODE_DIR / "ui" / "app.py"),
             "--server.runOnSave", "false"],
            cwd=str(Path(__file__).resolve().parent.parent),
        )
        print("UI launched -- open http://localhost:8501 in your browser.")
        print("Press Ctrl+C to stop the UI server.")
        try:
            proc.wait()
        except KeyboardInterrupt:
            proc.terminate()
        raise SystemExit(0)

    non_ai_flags = (
        args.list_sessions
        or args.read_session
        or args.delete_session
        or args.export_session
        or args.rename_session
        or args.stats
        or args.list_roles
        or args.help_guide
        or args.dry_run
    )
    if not non_ai_flags:
        try:
            validate_api_keys(args.provider)
        except (EnvironmentError, ValueError) as e:
            print(e)
            sys.exit(1)

    if args.list_sessions:
        list_sessions(reports_dir=str(REPORTS_DIR))
    elif args.read_session:
        read_session(filename=args.read_session, reports_dir=str(REPORTS_DIR))
    elif args.delete_session:
        delete_session(filename=args.delete_session, reports_dir=str(REPORTS_DIR))
    elif args.export_session:
        export_session(filename=args.export_session, reports_dir=str(REPORTS_DIR))
    elif args.rename_session:
        rename_session(filename=args.rename_session, reports_dir=str(REPORTS_DIR))
    elif args.stats:
        show_stats(reports_dir=str(REPORTS_DIR))
    elif args.list_roles:
        list_roles(mode=args.mode)
    elif args.help_guide:
        open_help_guide()
    elif args.revise:
        if args.mode != "coding":
            print("--revise is only available in coding mode.")
            print("Use: python SOURCE_CODE/main.py --mode coding --revise")
            sys.exit(1)
        generate_code_revision(
            start_role=args.role,
            provider=args.provider,
            model=args.model,
            dry_run=args.dry_run,
        )
    elif args.report:
        if args.mode != "writing":
            print("--report is only available in writing mode.")
            print("Use: python SOURCE_CODE/main.py --mode writing --report")
            sys.exit(1)
        generate_writing_report(provider=args.provider, model=args.model)
    elif args.mode == "search":
        sub_mode = getattr(args, 'sub', None)
        handle_search_mode(
            provider=args.provider,
            model=args.model,
            sub_mode=sub_mode,
        )


    elif args.mode == "sr":
        run_sr_launcher(
            provider=args.provider,
            model=args.model,
        )
    elif args.mode == "rct_search":
        run_rct_search_pipeline(
            provider=args.provider,
            model=args.model,
            dry_run=args.dry_run,
        )
    elif args.mode == "coding":
        handle_coding_mode(
            provider=args.provider,
            model=args.model,
            dry_run=args.dry_run,
        )
    elif args.mode == "writing":
        handle_writing_mode(
            provider=args.provider,
            model=args.model,
            dry_run=args.dry_run,
        )
    elif args.mode == "appraisal":
        handle_appraisal_mode(
            provider=args.provider,
            model=args.model,
            dry_run=args.dry_run,
        )
    else:
        # Unreachable via the CLI: parse_args restricts --mode to the six
        # choices dispatched above. A silent fallthrough to main() is how
        # the SR-mode KeyError path (#43) arose, so fail loudly instead
        # if a programmatic caller passes an undispatched mode.
        raise ValueError(
            f"run_cli: no dispatch branch for mode {args.mode!r}"
        )



# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    try:
        run_cli(parse_args())
    except KeyboardInterrupt:
        print("\n\nSession stopped. Returning to menu...\n")
        raise SystemExit(0)
