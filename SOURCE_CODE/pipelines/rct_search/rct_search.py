"""
RCT Search Mode - PICO-driven multi-database search and AI ranking
Supports: PubMed + Europe PMC
"""
import json
import re
import shutil
import sys
import urllib.parse
from datetime import datetime
from pathlib import Path
from typing import Optional
import random

# Add SOURCE_CODE to path
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent.parent
SOURCE_CODE_DIR = PROJECT_ROOT / "SOURCE_CODE"
sys.path.insert(0, str(SOURCE_CODE_DIR))

from utils.path_utils import PATH_MANAGER, get_input_dir, get_output_dir
from utils.document_reader import DocumentReader
from utils.rag import RAGUtils

# Project paths
BASE = PROJECT_ROOT
INPUT_DIR = BASE / "input"
OUTPUT_RCT_SEARCH = BASE / "output" / "rct_search"
REPORTS_DIR = BASE / "reports"
DOCS_RCT_SEARCH = PROJECT_ROOT / "docs" / "rct_search"
AI_DIR = PROJECT_ROOT / "prompts"
INPUT_SR = INPUT_DIR / "sr"


# ANSI colors for console output
RESET = "\033[0m"
ACCENT = "\033[38;5;121m"
DIM = "\033[0;31;40m"


def role_color(role: str) -> str:
    """Return color for a specific role"""
    colors = {
        "Formulator": "\033[38;5;51m",
        "Searcher": "\033[38;5;121m",
        "Validator": "\033[38;5;215m",
    }
    return colors.get(role, ACCENT)


def _read_topic_file(topic_file: Path) -> Optional[str]:
    """Read topic from file if it exists"""
    if topic_file.exists():
        try:
            return topic_file.read_text(encoding="utf-8").strip()
        except Exception:
            return None
    return None


def _clean_pico_term(t: str) -> str:
    """Clean PICO terms for PubMed query"""
    t = re.sub(r"\([^)]*\)", "", t)           # remove parenthetical qualifiers
    t = re.sub(r"\b(at\s+)?(post[\s\-]intervention|post[\s\-]treatment|"
               r"follow[\s\-]up|immediately after|baseline|"
               r"at\s+\d+\s+weeks?|at\s+\d+\s+months?)\b", "", t, flags=re.IGNORECASE)
    t = re.sub(r"[^\w\s\-]", " ", t)          # remove special chars except hyphen
    t = re.sub(r"\s+", " ", t).strip()         # collapse whitespace
    # keep only first 4 words to avoid over-specific phrases
    words = t.split()
    return " ".join(words[:4]) if len(words) > 4 else t


def fetch_pubmed_articles(
    query: str,
    max_results: int = 100,
) -> list[dict]:
    """
    Search PubMed via NCBI E-utilities (no API key required).
    Returns list of dicts: pmid, title, abstract, url.
    Returns empty list on any error.
    """
    import urllib.request
    import xml.etree.ElementTree as ET

    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    search_url = (
        f"{base}/esearch.fcgi?db=pubmed&term={urllib.parse.quote(query)}"
        f"&retmax={max_results}&retmode=json"
    )

    try:
        with urllib.request.urlopen(search_url) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except Exception:
        return []

    ids = data.get("esearchresult", {}).get("idlist", [])
    if not ids:
        return []

    ids_str = ",".join(ids)
    fetch_url = (
        f"{base}/efetch.fcgi?db=pubmed&id={ids_str}"
        f"&rettype=abstract&retmode=xml"
    )

    try:
        with urllib.request.urlopen(fetch_url) as resp:
            xml_data = resp.read().decode("utf-8")
    except Exception:
        return []

    try:
        root = ET.fromstring(xml_data)
    except ET.ParseError:
        return []

    articles = []
    for article in root.findall(".//PubmedArticle"):
        pmid_el = article.find(".//PMID")
        title_el = article.find(".//ArticleTitle")
        abstract_el = article.find(".//AbstractText")

        pmid = pmid_el.text if pmid_el is not None and pmid_el.text else ""
        title = title_el.text if title_el is not None and title_el.text else ""
        abstract = abstract_el.text if abstract_el is not None and abstract_el.text else ""

        if pmid and title:
            articles.append({
                "pmid": pmid,
                "title": title,
                "abstract": abstract,
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
                "source": "PubMed"
            })

    return articles


def fetch_europepmc_articles(
    query: str,
    max_results: int = 100,
) -> list[dict]:
    """
    Search Europe PMC via REST API (free, no API key required).
    Returns list of dicts with identifier, title, abstract, url.
    Returns empty list on any error.
    """
    import urllib.request
    import json
    
    # Europe PMC API endpoint
    base = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
    
    # Build query - Europe PMC uses different syntax
    clean_query = re.sub(r'[^\w\s]', ' ', query)
    
    search_url = (
        f"{base}?query={urllib.parse.quote(clean_query)}"
        f"&resultType=core"
        f"&pageSize={min(max_results, 100)}"
        f"&format=json"
    )
    
    try:
        req = urllib.request.Request(search_url, headers={"Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except Exception as e:
        print(f"[Europe PMC] Search failed: {e}")
        return []
    
    articles = []
    results = data.get("resultList", {}).get("result", [])
    
    for item in results:
        title = item.get("title", "")
        abstract = item.get("abstractText", "")
        pmid = item.get("pmid", "")
        doi = item.get("doi", "")
        source = item.get("source", "")
        
        # Skip if no title
        if not title:
            continue
        
        # Use pmid if available, otherwise use doi, otherwise generate hash
        if pmid:
            identifier = pmid
            url = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
        elif doi:
            identifier = doi.replace('/', '_')
            url = f"https://doi.org/{doi}"
        else:
            # Generate a unique identifier from title
            identifier = f"europe_{abs(hash(title)) % 10000000}"
            url = ""
        
        articles.append({
            "pmid": identifier,
            "title": title,
            "abstract": abstract or "No abstract available.",
            "url": url,
            "source": f"Europe PMC ({source})" if source else "Europe PMC",
            "original_pmid": pmid if pmid else "",
            "doi": doi if doi else ""
        })
    
    return articles


def merge_search_results(
    pubmed_articles: list[dict],
    europepmc_articles: list[dict],
) -> list[dict]:
    """
    Merge and deduplicate results from multiple sources.
    Uses PMID and title similarity for deduplication.
    """
    import re
    
    merged = []
    seen_pmids = set()
    seen_titles = set()
    
    def normalize_title(title: str) -> str:
        # Remove common patterns for better matching
        title = re.sub(r'[^a-zA-Z0-9\s]', '', title)
        title = ' '.join(title.lower().split())
        # Remove common stop words
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'for', 'nor', 'on', 'at', 'to', 'by', 'in', 'of', 'with'}
        title = ' '.join([w for w in title.split() if w not in stop_words])
        return title[:50]
    
    # Add PubMed articles (highest priority)
    for article in pubmed_articles:
        pmid = article.get("pmid", "")
        norm_title = normalize_title(article["title"])
        if pmid not in seen_pmids and norm_title not in seen_titles:
            if pmid:
                seen_pmids.add(pmid)
            seen_titles.add(norm_title)
            merged.append(article)
    
    # Add Europe PMC articles (deduplicated by PMID and title)
    for article in europepmc_articles:
        pmid = article.get("pmid", "")
        norm_title = normalize_title(article["title"])
        original_pmid = article.get("original_pmid", "")
        
        # Check if this article is already in the merged list
        is_duplicate = False
        
        # Check by original PMID if available
        if original_pmid and original_pmid in seen_pmids:
            is_duplicate = True
        
        # Check by generated identifier
        if pmid in seen_pmids:
            is_duplicate = True
        
        # Check by title similarity
        if norm_title in seen_titles:
            is_duplicate = True
        
        if not is_duplicate:
            if pmid:
                seen_pmids.add(pmid)
            if original_pmid:
                seen_pmids.add(original_pmid)
            seen_titles.add(norm_title)
            merged.append(article)
    
    return merged


def call_ai(prompt: str, provider: str = "deepseek", model: Optional[str] = None) -> str:
    """
    Call AI provider with prompt.
    Uses the main module's call_ai function.
    """
    import sys
    from pathlib import Path
    
    # Add SOURCE_CODE to path if needed
    source_code_dir = Path(__file__).resolve().parent.parent.parent
    if str(source_code_dir) not in sys.path:
        sys.path.insert(0, str(source_code_dir))
    
    try:
        # Import from main module
        from main import call_ai as _call_ai
        return _call_ai(prompt=prompt, provider=provider, model=model)
    except (ImportError, AttributeError) as e:
        # Try alternative import
        try:
            from ..main import call_ai as _call_ai
            return _call_ai(prompt=prompt, provider=provider, model=model)
        except ImportError:
            # Fallback for testing: mock response
            if provider == "ollama":
                return f"[MOCK] Ollama response for: {prompt[:50]}..."
            elif provider == "qwen":
                return f"[MOCK] Qwen response for: {prompt[:50]}..."
            elif provider == "openai":
                return f"[MOCK] OpenAI response for: {prompt[:50]}..."
            elif provider == "anthropic":
                return f"[MOCK] Anthropic response for: {prompt[:50]}..."
            else:
                return f"[MOCK] {provider} response for: {prompt[:50]}..."

def _add_hyperlink(paragraph, text: str, url: str):
    """Add a clickable hyperlink to a paragraph."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )
        
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        new_run.append(rPr)
        
        new_t = OxmlElement("w:t")
        new_t.text = text
        new_run.append(new_t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        # Fallback: plain text
        paragraph.add_run(text)


def _md_to_docx(md_content: str, title: str, out_path: Path) -> None:
    """Convert markdown to DOCX using python-docx"""
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = docx.Document()
        # Title
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True

        doc.add_paragraph()

        # Process markdown
        for line in md_content.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("|") and "|" in line[1:]:
                # Skip table headers for now
                pass
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Pt(36)
            elif line.strip():
                doc.add_paragraph(line)

        doc.save(str(out_path))
    except ImportError:
        # Fallback: save as text
        out_path.with_suffix(".txt").write_text(md_content, encoding="utf-8")


def _ranked_articles_to_docx(
    ranked: list[dict],
    title: str,
    out_path: Path,
    topic: str = ""
) -> None:
    """Write ranked article list as DOCX with table"""
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = docx.Document()
        # Title
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True

        if topic:
            doc.add_paragraph(f"Topic: {topic}")

        doc.add_paragraph()

        if not ranked:
            doc.add_paragraph("No articles retrieved from PubMed or Europe PMC.")
            doc.save(str(out_path))
            return

        # Determine which sources are present for dynamic caption
        sources_present = set()
        for r in ranked:
            source = r.get("source", "")
            if source:
                # Extract just the main source name (e.g., "PubMed", "Europe")
                main_source = source.split()[0] if source.startswith("Europe") else source
                sources_present.add(main_source)
        
        source_text = " and ".join(sorted(sources_present)) if sources_present else "PubMed and Europe PMC"

        # Caption
        cap = doc.add_paragraph()
        cap.add_run(
            f"All {len(ranked)} articles retrieved from {source_text}, "
            "ordered by PICO relevance score (10 = most relevant)."
        ).italic = True

        doc.add_paragraph()

        # Table: Rank | Score | Source | Title | PMID | Link
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        # Header
        hdr = table.rows[0].cells
        hdr[0].text = "Rank"
        hdr[1].text = "Score"
        hdr[2].text = "Source"
        hdr[3].text = "Title"
        hdr[4].text = "PMID"
        hdr[5].text = "Link"

        # Data rows
        for r in ranked:
            row = table.add_row().cells
            row[0].text = str(r["rank"])
            row[1].text = f"{r['score']}/10"
            row[2].text = r.get("source", "Unknown")
            row[3].text = r["title"]
            row[4].text = r["pmid"]
            # Add hyperlink
            p = row[5].paragraphs[0]
            if r.get("url"):
                _add_hyperlink(p, "View", r["url"])
            else:
                p.add_run("No link")

        # Add comments after the table
        doc.add_paragraph()
        note1 = doc.add_paragraph()
        note1.add_run(
            "Select your top 5 articles, download PDFs and place them in input/sr/ to run the SR pipeline."
        ).italic = True

        doc.add_paragraph()
        note2 = doc.add_paragraph()
        note2.add_run(
            "For explanation on ranking, please refer to the full report in the reports folder."
        ).italic = True

        doc.save(str(out_path))
    except ImportError:
        # Fallback: CSV
        import csv
        with open(out_path.with_suffix(".csv"), "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Rank", "Score", "Source", "Title", "PMID", "URL"])
            for r in ranked:
                writer.writerow([r["rank"], r["score"], r.get("source", "Unknown"), r["title"], r["pmid"], r.get("url", "")])


def run_rct_search_pipeline(
    provider: str = "deepseek",
    model: str | None = None,
    reports_dir: Path = REPORTS_DIR,
    dry_run: bool = False,
) -> Path:
    """
    Single-pass RCT search pipeline:
      1. Formulator - structures user topic into PICO question
      2. Searcher - builds Boolean search strategy for all 7 databases
      3. Validator - validates alignment and approves or requests refinement
      4. PubMed + Europe PMC Fetch, merge, and AI ranking
    Saves output as output/rct_search/rct_search_{ts}.md and .docx.
    Returns path of the .md report.
    """
    print("\n" + "=" * 55)
    print("  RCT SEARCH PIPELINE")
    print("=" * 55)
    print("  This pipeline will:")
    print("  1. Structure your topic into a PICO question")
    print("  2. Build a search strategy for all 7 SR databases")
    print("  3. Validate the strategy before download")
    print("  4. Fetch and rank articles from PubMed and Europe PMC")
    print("  No article appraisal - use --mode appraisal for that.")
    print("=" * 55 + "\n")

    topic_file = DOCS_RCT_SEARCH / "topic.md"
    file_topic = _read_topic_file(topic_file)
    if file_topic:
        topic = file_topic
        print(f"[RCT Search] Topic loaded from docs/rct_search/topic.md: {topic}")
    else:
        try:
            topic = input("Enter your research topic: ").strip()
        except (EOFError, KeyboardInterrupt):
            print()
            topic = ""
    if not topic:
        print("No topic entered. Exiting.")
        sys.exit(0)

    # -- PICO input: offer to import saved JSON or prompt manually --------------

    _pico_input_dir = INPUT_DIR / "rct_search"
    pico_json_files = sorted(
        _pico_input_dir.glob("pico_*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    ) if _pico_input_dir.exists() else []
    if not pico_json_files and OUTPUT_RCT_SEARCH.exists():
        pico_json_files = sorted(
            OUTPUT_RCT_SEARCH.glob("pico_*.json"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )
    _pico_source = "input/rct_search" if (_pico_input_dir.exists() and list(_pico_input_dir.glob("pico_*.json"))) else "output/rct_search"

    pico_population = ""
    pico_intervention = ""
    pico_comparator = ""
    pico_outcome = ""

    if pico_json_files:
        print()
        print(f"[RCT Search] Saved PICO files found in {_pico_source}:")
        for idx, pf in enumerate(pico_json_files[:5], 1):
            print(f"  {idx}. {pf.name}")
        print("  0. Enter new PICO manually")
        print()
        try:
            choice = input("Import a saved PICO? Enter number or 0 for new: ").strip()
        except (EOFError, KeyboardInterrupt):
            choice = "0"
        if choice.isdigit() and 1 <= int(choice) <= len(pico_json_files[:5]):
            chosen = pico_json_files[int(choice) - 1]
            try:
                pico_data = json.loads(chosen.read_text(encoding="utf-8"))
                pico_population = pico_data.get("population", "")
                pico_intervention = pico_data.get("intervention", "")
                pico_comparator = pico_data.get("comparator", "")
                pico_outcome = pico_data.get("outcome", "")
                if not topic:
                    topic = pico_data.get("topic", topic)
                print(f"[RCT Search] PICO imported from {chosen.name}")
            except Exception as _e:
                print(f"[RCT Search] Could not load PICO JSON: {_e}")

    if not any([pico_population, pico_intervention, pico_comparator, pico_outcome]):
        print()
        print("[RCT Search] Enter PICO components (press Enter to skip any field):")
        try:
            pico_population = input("  Population   (P): ").strip()
            pico_intervention = input("  Intervention (I): ").strip()
            pico_comparator = input("  Comparator   (C): ").strip()
            pico_outcome = input("  Outcome      (O): ").strip()
        except (EOFError, KeyboardInterrupt):
            print()

    pico_manual_context = ""
    if any([pico_population, pico_intervention, pico_comparator, pico_outcome]):
        pico_manual_context = (
            "\n\nUser-provided PICO components:\n"
            f"  Population:   {pico_population or '(not specified)'}\n"
            f"  Intervention: {pico_intervention or '(not specified)'}\n"
            f"  Comparator:   {pico_comparator or '(not specified)'}\n"
            f"  Outcome:      {pico_outcome or '(not specified)'}"
        )
    # -- end PICO input ---------------------------------------------------------

    pico_path = DOCS_RCT_SEARCH / "pico-framework.md"
    db_guide = DOCS_RCT_SEARCH / "database-guide.md"
    val_criteria = DOCS_RCT_SEARCH / "validation-criteria.md"
    pico_context = pico_path.read_text(encoding="utf-8", errors="replace") if pico_path.exists() else ""
    db_guide_text = db_guide.read_text(encoding="utf-8", errors="replace") if db_guide.exists() else ""
    val_criteria_text = val_criteria.read_text(encoding="utf-8", errors="replace") if val_criteria.exists() else ""

    stages = [
        {
            "role": "Formulator",
            "prompt_file": AI_DIR / "formulator-prompt.md",
            "extra_context": pico_context,
            "task": f"The user's research topic is: {topic}{pico_manual_context}\n\nStructure this into a formal PICO question.",
        },
        {
            "role": "Searcher",
            "prompt_file": AI_DIR / "searcher-prompt.md",
            "extra_context": db_guide_text,
            "task": "Build a comprehensive Boolean search strategy for all 7 SR databases based on the PICO question above.",
        },
        {
            "role": "Validator",
            "prompt_file": AI_DIR / "validator-prompt.md",
            "extra_context": val_criteria_text,
            "task": "Validate the search strategy above. Check PICO alignment, database coverage, syntax, and RCT filters. Return APPROVED FOR DOWNLOAD or REQUIRES REFINEMENT with specific justification.",
        },
    ]

    report_parts = [
        "# RCT Search Strategy Report",
        f"Topic: {topic}",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "---",
        "",
        "> **Note:** This report contains the validated search strategy only.",
        "> For article appraisal, copy URLs into --mode appraisal.",
        "",
    ]

    previous_response = ""

    for stage in stages:
        role = stage["role"]
        colour = role_color(role)
        print(f"Running {role}...")

        prompt_file = Path(stage["prompt_file"])
        if not prompt_file.exists():
            print(f"  Prompt not found: {prompt_file} - skipping {role}.")
            continue

        role_prompt = prompt_file.read_text(encoding="utf-8", errors="replace")
        parts = [role_prompt]
        if pico_context:
            parts.append(f"## PICO Framework Reference\n{pico_context}")
        if stage["extra_context"]:
            parts.append(f"## Reference Document\n{stage['extra_context']}")
        if previous_response:
            parts.append(f"## Previous Stage Output\n{previous_response}")
        parts.append(f"## Task\n{stage['task']}")
        full_prompt = "\n\n".join(parts)

        if dry_run:
            response = f"[DRY RUN] {role} would respond here."
        else:
            try:
                response = call_ai(prompt=full_prompt, provider=provider, model=model)
            except RuntimeError as exc:
                response = f"[ERROR in {role}: {exc}]"

        previous_response = response
        report_parts.append(f"## {role} Output\n\n{response}\n")
        print(f"{colour}{role} complete.{RESET}\n")

        # -- save PICO JSON after Formulator completes -------------------------
        if role == "Formulator" and not dry_run:
            pico_json_dir = OUTPUT_RCT_SEARCH
            pico_json_dir.mkdir(parents=True, exist_ok=True)
            pico_json_path = pico_json_dir / f"pico_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            pico_data = {
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "topic": topic,
                "formulator_output": response,
                "population": "",
                "intervention": "",
                "comparator": "",
                "outcome": "",
                "study_design": "RCT",
                "effect_measure": "SMD",
                "source_mode": "rct_search",
                "pubmed_query_raw": "",
                "pubmed_query_cleaned": "",
            }

            def _extract_pico_field(text, labels):
                for lbl in labels:
                    m = re.search(r"(?i)" + lbl + r"[^:]*:\s*(.+)", text)
                    if m:
                        return m.group(1).strip()
                return ""

            pico_data["population"] = pico_population or _extract_pico_field(response, [r"P\s*\(Population\)", r"Population"])
            pico_data["intervention"] = pico_intervention or _extract_pico_field(response, [r"I\s*\(Intervention\)", r"Intervention"])
            pico_data["comparator"] = pico_comparator or _extract_pico_field(response, [r"C\s*\(Comp\w+\)", r"Compar"])
            pico_data["outcome"] = pico_outcome or _extract_pico_field(response, [r"O\s*\(Outcome\)", r"Outcome"])
            pico_json_path.write_text(
                json.dumps(pico_data, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
            print(f"[RCT Search] PICO saved to: {pico_json_path.name}")

        # ----------------------------------------------------------------------

    validator_output = previous_response.upper()
    if "APPROVED FOR DOWNLOAD" in validator_output:
        status = "APPROVED FOR DOWNLOAD"
        print("\n Search strategy APPROVED FOR DOWNLOAD.")
    else:
        status = "REQUIRES REFINEMENT"
        print("\n Search strategy REQUIRES REFINEMENT - see Validator output.")

    report_parts.append(f"## Final Status\n\n**{status}**\n")
    report_parts.append(
        "## Next Steps\n\n"
        "- If APPROVED: copy database search strings into each SR database platform\n"
        "- Download article lists from each database\n"
        "- Run python src/main.py --mode appraisal to appraise individual articles\n"
        "- Run python src/main.py --mode sr for full systematic review pipeline\n"
    )

    # -- Stage 4: Multi-Database Search (PubMed + Europe PMC) ----------

    # Clean PICO terms for PubMed
    _pico_terms = [_clean_pico_term(t) for t in [pico_population, pico_intervention, pico_outcome] if t]
    _pico_terms = [t for t in _pico_terms if t]
    _pubmed_query = " AND ".join(_pico_terms) if _pico_terms else topic

    # Broadened RCT filter for PubMed
    _pubmed_query = _pubmed_query + " AND (randomized controlled trial[pt] OR clinical trial[pt] OR random*[tiab] OR RCT[tiab])"
    print(f"[RCT Search] PubMed query: {_pubmed_query}")
    print("[RCT Search] Searching PubMed for relevant articles...")

    # Search PubMed
    pubmed_articles = fetch_pubmed_articles(_pubmed_query, max_results=100) if not dry_run else []

    # Search Europe PMC
    print("[RCT Search] Searching Europe PMC for relevant articles...")
    _europe_query = f"{pico_population} {pico_intervention} {pico_outcome}" if pico_population else topic
    europepmc_articles = fetch_europepmc_articles(_europe_query, max_results=100) if not dry_run else []

    # Merge results
    articles = merge_search_results(pubmed_articles, europepmc_articles) if not dry_run else [
        {"pmid": "00000001", "title": "[DRY RUN] Test Article",
         "abstract": "Dry run abstract.",
         "url": "https://pubmed.ncbi.nlm.nih.gov/00000001/",
         "source": "Dry Run"}
    ]

    # Shuffle to prevent position bias in AI ranking
    if not dry_run and articles:
        random.shuffle(articles)
        print(f"[RCT Search] Shuffled {len(articles)} articles for unbiased ranking")

    print(f"[RCT Search] Found {len(pubmed_articles)} PubMed + {len(europepmc_articles)} Europe PMC = {len(articles)} unique articles")

    if articles:
        print(f"[RCT Search] Ranking {len(articles)} articles by PICO relevance...")
        
        # Build abstracts block with clear source information
        abstracts_block = ""
        for idx, art in enumerate(articles, 1):
            source = art.get("source", "PubMed")
            pmid = art.get("pmid", "N/A")
            title = art.get("title", "No title")
            url = art.get("url", "")
            abstract = art.get("abstract", "No abstract available.")
            
            abstracts_block += (
                f"\n### Article {idx}\n"
                f"Source: {source}\n"
                f"PMID: {pmid}\n"
                f"Title: {title}\n"
                f"URL: {url}\n"
                f"Abstract: {abstract}\n"
            )
        
        rank_prompt = (
            "You are a systematic review expert. Your ONLY task is to rate articles.\n\n"
            "PICO Question:\n" + previous_response + "\n\n"
            "INSTRUCTIONS:\n"
            "- Rate each article for relevance to the PICO question on a scale of 1-10.\n"
            "- 10 = highly relevant RCT directly matching the PICO.\n"
            "- 1 = not relevant.\n"
            "- IMPORTANT: Score MUST be a single digit between 1 and 10. Do NOT use 0, 100, or any other scale.\n"
            "- Articles come from PubMed and Europe PMC. Both are valid sources.\n"
            "- Output ONLY the rating lines. No summaries, no headers, no extra text.\n"
            "- Each line MUST follow this EXACT format:\n"
            "ARTICLE_RANK: 1 | SCORE: 9 | PMID: 12345678 | SOURCE: PubMed | TITLE: Example title here | URL: https://pubmed.ncbi.nlm.nih.gov/12345678/\n\n"
            "EXAMPLE OUTPUT (do not copy - use real values):\n"
            "ARTICLE_RANK: 1 | SCORE: 9 | PMID: 11111111 | SOURCE: PubMed | TITLE: CBT for fibromyalgia RCT | URL: https://pubmed.ncbi.nlm.nih.gov/11111111/\n"
            "ARTICLE_RANK: 2 | SCORE: 7 | PMID: 22222222 | SOURCE: Europe PMC | TITLE: Mindfulness for chronic pain | URL: https://pubmed.ncbi.nlm.nih.gov/22222222/\n"
            "ARTICLE_RANK: 3 | SCORE: 3 | PMID: 33333333 | SOURCE: Europe PMC | TITLE: Ketamine infusion study | URL: https://pubmed.ncbi.nlm.nih.gov/33333333/\n\n"
            "NOW RATE THESE ARTICLES - output rating lines ONLY:\n\n"
            + abstracts_block
        )
        
        if dry_run:
            rank_response = "\n".join(
                "ARTICLE_RANK: " + str(i) + " | SCORE: " + str(10 - i) + " | "
                + "PMID: " + art["pmid"] + " | SOURCE: " + art.get("source", "Unknown") + " | TITLE: " + art["title"]
                + " | URL: " + art["url"]
                for i, art in enumerate(articles, 1)
            )
        else:
            try:
                rank_response = call_ai(
                    prompt=rank_prompt, provider=provider, model=model
                )
            except RuntimeError as exc:
                rank_response = "[ERROR ranking articles: " + str(exc) + "]"
        
        ranked = []
        if rank_response.startswith("[ERROR"):
            print(f"[RCT Search] Ranking error: {rank_response}")
            # Fallback: use articles in fetch order
            ranked = [
                {"rank": i, "score": min(10, len(articles) - i + 1),
                 "pmid": a["pmid"], "title": a["title"], "url": a["url"], "source": a.get("source", "Unknown")}
                for i, a in enumerate(articles, 1)
            ]
        else:
            # Parse the AI response
            for ln in rank_response.splitlines():
                ln = ln.strip()
                if not ln:
                    continue
                
                # Try to match the format with SOURCE field
                m = re.match(
                    r"ARTICLE_RANK:\s*(\d+)\s*\|\s*SCORE:\s*(\d+)\s*\|"
                    r"\s*PMID:\s*(\S+)\s*\|\s*SOURCE:\s*([^|]+)\s*\|\s*TITLE:\s*(.+?)\s*\|\s*URL:\s*(\S+)",
                    ln
                )
                if m:
                    score = int(m.group(2))
                    original_score = score
                    
                    # Normalize: AI seems to use 0-100 or 0-160 scale
                    if score > 10:
                        # Try dividing by 10 first
                        normalized_score = round(score / 10)
                        # If still > 10, divide by 16
                        if normalized_score > 10:
                            normalized_score = round(score / 16)
                        # Cap at 1-10
                        if normalized_score > 10:
                            normalized_score = 10
                        elif normalized_score < 1:
                            normalized_score = 1
                        score = normalized_score
                        print(f"[RCT Search] Normalized score: {original_score}  - {score} for {m.group(3)}")
                    
                    ranked.append({
                        "rank": int(m.group(1)),
                        "score": score,
                        "pmid": m.group(3).strip(),
                        "source": m.group(4).strip(),
                        "title": m.group(5).strip(),
                        "url": m.group(6).strip(),
                    })
                else:
                    # Try without SOURCE field (fallback for older format)
                    m2 = re.match(
                        r"ARTICLE_RANK:\s*(\d+)\s*\|\s*SCORE:\s*(\d+)\s*\|"
                        r"\s*PMID:\s*(\S+)\s*\|\s*TITLE:\s*(.+?)\s*\|\s*URL:\s*(\S+)",
                        ln
                    )
                    if m2:
                        score = int(m2.group(2))
                        original_score = score
                        
                        # Normalize
                        if score > 10:
                            normalized_score = round(score / 10)
                            if normalized_score > 10:
                                normalized_score = round(score / 16)
                            if normalized_score > 10:
                                normalized_score = 10
                            elif normalized_score < 1:
                                normalized_score = 1
                            score = normalized_score
                            print(f"[RCT Search] Normalized score: {original_score}  - {score} for {m2.group(3)}")
                        
                        # Find the article in our list to get its source
                        pmid = m2.group(3).strip()
                        source = "PubMed"  # Default
                        for art in articles:
                            if art.get("pmid") == pmid:
                                source = art.get("source", "PubMed")
                                break
                        ranked.append({
                            "rank": int(m2.group(1)),
                            "score": score,
                            "pmid": pmid,
                            "source": source,
                            "title": m2.group(4).strip(),
                            "url": m2.group(5).strip(),
                        })
        
        # If no articles were parsed or less than expected, use fallback
        if not ranked or len(ranked) < len(articles) * 0.5:
            print(f"[RCT Search] Warning: Only {len(ranked)} articles parsed from AI response. Using fallback ranking.")
            ranked = [
                {"rank": i, "score": min(10, len(articles) - i + 1),
                 "pmid": a["pmid"], "title": a["title"], "url": a["url"], "source": a.get("source", "Unknown")}
                for i, a in enumerate(articles, 1)
            ]
        
        # Sort by score (highest first)
        ranked.sort(key=lambda x: x["score"], reverse=True)
        for new_rank, r in enumerate(ranked, 1):
            r["rank"] = new_rank
        
        # Build the table
        table_lines = [
            "## Ranked Article List\n",
            f"_All {len(ranked)} articles retrieved from PubMed and Europe PMC, ordered by PICO relevance score (10 = most relevant)._\n",
            "| Rank | Score | Source | Title | PMID | Link |",
            "|------|-------|--------|-------|------|------|",
        ]
        for r in ranked:
            source = r.get("source", "Unknown")
            url_display = r["url"] if r["url"] else "#"
            table_lines.append(
                "| " + str(r["rank"]) + " | " + str(r["score"]) + "/10 | "
                + source + " | "
                + r["title"][:60] + ("..." if len(r["title"]) > 60 else "") + " | " 
                + r["pmid"] + " | [Link](" + url_display + ") |"
            )
        table_lines.append(
            "\n> Select your top 5 articles, download PDFs "
            "and place them in input/sr/ to run the SR pipeline."
        )
        table_lines.append(
            "\n> For explanation on ranking, please refer to the full report "
            "in the reports folder."
        )

        report_parts.append("\n".join(table_lines))
        top = ranked[0]["title"] if ranked else "N/A"
        print("[RCT Search] Ranking complete. Top article: " + top)
        
        # Update PICO JSON
        pico_files = sorted(OUTPUT_RCT_SEARCH.glob("pico_*.json"), reverse=True)
        if pico_files:
            pico_data = json.loads(pico_files[0].read_text(encoding="utf-8"))
            pico_data["ranked_articles"] = ranked
            pico_data["pubmed_query_raw"] = " AND ".join(
                [t for t in [pico_population, pico_intervention, pico_outcome] if t]
            ) if any([pico_population, pico_intervention, pico_outcome]) else topic
            pico_data["pubmed_query_cleaned"] = _pubmed_query
            pico_files[0].write_text(
                json.dumps(pico_data, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
            print("[RCT Search] PICO JSON updated with "
                  + str(len(ranked)) + " ranked article(s).")
    else:
        print("[RCT Search] No articles found on PubMed or Europe PMC.")
        report_parts.append(
            "## Ranked Article List\n\n_No articles retrieved from PubMed or Europe PMC._\n"
        )

    # -- End Stage 4 ----------------------------------------------------------

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if reports_dir != REPORTS_DIR:
        out_dir = reports_dir
        final_dir = reports_dir
    else:
        out_dir = REPORTS_DIR / "rct_search"
        final_dir = OUTPUT_RCT_SEARCH
    out_dir.mkdir(parents=True, exist_ok=True)
    final_dir.mkdir(parents=True, exist_ok=True)
    md_path = out_dir / f"rct_search_{timestamp}.md"
    docx_path = out_dir / f"rct_search_{timestamp}.docx"

    md_content = "\n".join(report_parts)
    md_path.write_text(md_content, encoding="utf-8")
    print(f"Intermediate report   : reports/rct_search/{md_path.name}")

    final_parts = [p for p in report_parts if "Ranked Article List" in p or "Final Status" in p]
    if final_parts and final_dir != out_dir:
        final_md_path = final_dir / f"rct_search_{timestamp}.md"
        final_md_content = "\n".join(final_parts)
        final_md_path.write_text(final_md_content, encoding="utf-8")
        print(f"Final report saved    : output/rct_search/{final_md_path.name}")
        final_docx_path = final_dir / f"rct_search_{timestamp}.docx"
        try:
            _ranked_articles_to_docx(ranked, f"RCT Search Results - {timestamp}", final_docx_path, topic=topic)
            print(f"Results DOCX saved    : output/rct_search/{final_docx_path.name}")
        except Exception as exc:
            print(f"  Warning: could not generate results .docx - {exc}")

    try:
        _md_to_docx(md_content, f"RCT Search Strategy - {timestamp}", docx_path)
        print(f"Word document saved   : {out_dir.name}\\{docx_path.name}")
    except Exception as exc:
        print(f"  Warning: could not generate .docx - {exc}")

    # -- Auto-copy PICO JSON to input/sr/ -------------------------------------
    if not dry_run:
        pico_files = sorted(OUTPUT_RCT_SEARCH.glob("pico_*.json"), reverse=True)
        if pico_files:
            latest_pico = pico_files[0]
            print()
            print(f"[RCT Search] Latest PICO JSON: {latest_pico.name}")
            try:
                copy_choice = input(
                    "[RCT Search] Copy pico_*.json to input/sr/ for SR pipeline? [Y/N]: "
                ).strip().upper()
            except (EOFError, KeyboardInterrupt):
                copy_choice = "N"
            if copy_choice == "Y":
                INPUT_SR.mkdir(parents=True, exist_ok=True)
                dest = INPUT_SR / latest_pico.name
                shutil.copy2(str(latest_pico), str(dest))
                print(f"[RCT Search] PICO JSON copied to input/sr/{latest_pico.name}")
            else:
                print("[RCT Search] PICO JSON not copied - you can copy it manually later.")
    # -------------------------------------------------------------------------

    return md_path


