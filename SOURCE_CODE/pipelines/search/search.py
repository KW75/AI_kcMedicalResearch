"""
src/modes/search.py
Search mode engine for AI kcMedicalResearch.

Sub-modes:
  - Topic Search  : DuckDuckGo web search  - LLM synopsis with reference links
  - Article Search: PubMed search by article type  - LLM summary + comparison
"""
from __future__ import annotations

import datetime
import json
import re
import time
import threading
import urllib.parse
import sys
import requests
from pathlib import Path
from typing import Optional

# Add SOURCE_CODE to path
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent.parent
SOURCE_CODE_DIR = PROJECT_ROOT / "SOURCE_CODE"
sys.path.insert(0, str(SOURCE_CODE_DIR))

from utils.path_utils import PATH_MANAGER, get_input_dir, get_output_dir
from utils.document_reader import DocumentReader
from utils.rag import RAGUtils

try:
    from docx import Document
    from docx.shared import Pt
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
PUBMED_ESEARCH = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
PUBMED_EFETCH  = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi"
PUBMED_ESUM    = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi"
DDG_API        = "https://api.duckduckgo.com/"
DDG_HTML       = "https://html.duckduckgo.com/html/"

ARTICLE_TYPES = {
    "1": ("Systematic Review",  "systematic[sb]"),
    "2": ("Meta-analysis",      "meta-analysis[pt]"),
    "3": ("RCT",                "randomized controlled trial[pt]"),
    "4": ("Guideline",          "guideline[pt]"),
    "5": ("Protocol",           "research support[pt]"),
    "6": ("Review Article",     "review[pt]"),
    "7": ("Clinical Trial",     "clinical trial[pt]"),
    "8": ("Custom",             ""),
}

DISCLOSURE = (
    "This report was generated with AI assistance. "
    "Always verify search results against primary sources."
)

MAX_RESULTS_TOPIC   = 10
MAX_RESULTS_ARTICLE = 15

# Models known to be small/slow  - use reduced result counts
_SMALL_MODELS = {"llama3.2", "llama3.2:latest", "qwen2.5-coder:3b", "qwen2.5-coder:3b:latest", "llama3.1:8b"}

def _result_limits(model: str | None) -> tuple[int, int]:
    """Return (topic_limit, article_limit) based on model capability."""
    m = (model or "").lower().strip()
    if any(m == s or m.startswith(s.split(":")[0]) for s in _SMALL_MODELS):
        return 5, 8
    return MAX_RESULTS_TOPIC, MAX_RESULTS_ARTICLE

# ---------------------------------------------------------------------------
# Path helpers
# ---------------------------------------------------------------------------
def _project_root() -> Path:
    return Path(__file__).resolve().parent.parent.parent.parent.parent

def _ts() -> str:
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

def _paths(root: Path) -> dict:
    return {
        "output":  PROJECT_ROOT / "output"  / "search",
        "reports": PROJECT_ROOT / "reports" / "search",
        "docs":    PROJECT_ROOT / "docs"    / "search",
    }

# ---------------------------------------------------------------------------
# Spinner
# ---------------------------------------------------------------------------
class _Spinner:
    def __init__(self, message: str = "Processing"):
        self.message = message
        self._stop  = threading.Event()
        self._thread = threading.Thread(target=self._spin, daemon=True)

    def _spin(self):
        frames = ["|", "/", "-", "\\"]
        i = 0
        while not self._stop.is_set():
            print(f"{frames[i % 4]}  {self.message}...", end="", flush=True)
            i += 1
            time.sleep(0.12)
        print(f" -  {self.message}  - done.      ")

    def __enter__(self):
        self._thread.start()
        return self

    def __exit__(self, *_):
        self._stop.set()
        self._thread.join()

# ---------------------------------------------------------------------------
# Guideline loader
# ---------------------------------------------------------------------------
def _load_guidelines(docs_path: Path) -> str:
    parts = []
    if not docs_path.exists():
        return ""
    for f in sorted(docs_path.iterdir()):
        if f.is_file() and f.suffix.lower() == ".md":
            parts.append(f"### {f.name}\n{f.read_text(encoding='utf-8')}")
    return "\n\n".join(parts)

# ---------------------------------------------------------------------------
# DuckDuckGo web search
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Europe PMC search (replaces DuckDuckGo for Topic Search)
# ---------------------------------------------------------------------------
EUROPEPMC_SEARCH = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"

def _europepmc_search(query: str, max_results: int = MAX_RESULTS_TOPIC) -> list[dict]:
    """
    Search Europe PMC and return list of {title, url, snippet} dicts.
    Covers PubMed, PubMed Central, bioRxiv, WHO, and clinical guidelines.
    No API key required.
    """
    results = []
    try:
        params = {
            "query":    query,
            "format":   "json",
            "pageSize": str(max_results),
            "resultType": "core",
        }
        r = requests.get(EUROPEPMC_SEARCH, params=params, timeout=15)
        if r.status_code == 200:
            data = r.json()
            for item in data.get("resultList", {}).get("result", []):
                pmid  = item.get("pmid", "")
                pmcid = item.get("pmcid", "")
                title = item.get("title", "").strip()
                abstract = item.get("abstractText", "").strip()
                snippet  = abstract[:300] if abstract else title
                if pmid:
                    url = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
                elif pmcid:
                    url = f"https://europepmc.org/article/PMC/{pmcid}"
                else:
                    url = f"https://europepmc.org/search?query={urllib.parse.quote(title)}"
                if title and url:
                    results.append({
                        "title":   title,
                        "url":     url,
                        "snippet": snippet,
                    })
                if len(results) >= max_results:
                    break
    except Exception as exc:
        print(f"  [SEARCH] Europe PMC error: {exc}")
    return results[:max_results]

def _pubmed_search(
    query: str,
    article_filter: str,
    max_results: int = MAX_RESULTS_ARTICLE,
    api_key: str = "",
) -> list[dict]:
    """Search PubMed and return list of article metadata dicts."""
    search_term = f"({query})"
    if article_filter:
        search_term += f" AND {article_filter}"

    params: dict = {
        "db":      "pubmed",
        "term":    search_term,
        "retmax":  max_results,
        "retmode": "json",
    }
    if api_key:
        params["api_key"] = api_key

    try:
        r = requests.get(PUBMED_ESEARCH, params=params, timeout=20)
        r.raise_for_status()
        ids = r.json()["esearchresult"]["idlist"]
    except Exception as exc:
        print(f"  [SEARCH] PubMed search error: {exc}")
        return []

    if not ids:
        return []

    # Fetch summaries
    sum_params: dict = {
        "db":      "pubmed",
        "id":      ",".join(ids),
        "retmode": "json",
    }
    if api_key:
        sum_params["api_key"] = api_key

    try:
        r = requests.get(PUBMED_ESUM, params=sum_params, timeout=20)
        r.raise_for_status()
        uids = r.json()["result"]["uids"]
        raw  = r.json()["result"]
    except Exception as exc:
        print(f"  [SEARCH] PubMed summary error: {exc}")
        return []

    articles = []
    for uid in uids:
        item = raw.get(uid, {})
        authors = ", ".join(
            a.get("name", "") for a in item.get("authors", [])[:3]
        )
        if len(item.get("authors", [])) > 3:
            authors += " et al."
        pub_date = item.get("pubdate", "")[:4]
        journal  = item.get("fulljournalname", item.get("source", ""))
        title    = item.get("title", "Untitled")
        doi      = ""
        for art_id in item.get("articleids", []):
            if art_id.get("idtype") == "doi":
                doi = art_id.get("value", "")
        url = f"https://pubmed.ncbi.nlm.nih.gov/{uid}/"
        articles.append({
            "uid":      uid,
            "title":    title,
            "authors":  authors,
            "year":     pub_date,
            "journal":  journal,
            "doi":      doi,
            "url":      url,
            "abstract": "",   # fetched separately if needed
        })

    return articles

def _pubmed_fetch_abstracts(
    articles: list[dict],
    api_key: str = "",
) -> list[dict]:
    """Fetch abstracts for a list of articles (adds 'abstract' key)."""
    if not articles:
        return articles
    ids = [a["uid"] for a in articles]
    params: dict = {
        "db":       "pubmed",
        "id":       ",".join(ids),
        "rettype":  "abstract",
        "retmode":  "text",
    }
    if api_key:
        params["api_key"] = api_key
    try:
        r = requests.get(PUBMED_EFETCH, params=params, timeout=30)
        r.raise_for_status()
        # Split by double newline blocks  - each article separated by blank lines
        blocks = re.split(r"\n{3,}", r.text.strip())
        for i, article in enumerate(articles):
            article["abstract"] = blocks[i].strip() if i < len(blocks) else ""
    except Exception:
        pass
    return articles

# ---------------------------------------------------------------------------
# LLM call wrapper
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# LLM call wrapper  (direct call  - no threading, Ctrl+C works on Windows)
# ---------------------------------------------------------------------------
def _call_llm(
    system_prompt: str,
    user_prompt: str,
    call_llm_fn,
    spinner_message: str = "LLM processing",
) -> str:
    """Call the LLM directly. Spinner is shown but does not block interrupts."""
    combined = f"{system_prompt}\n\n{user_prompt}" if system_prompt else user_prompt
    print(f"  / {spinner_message}...", end="", flush=True)
    try:
        result = call_llm_fn(system_prompt=system_prompt, user_prompt=user_prompt)
        print(f"\r   -  {spinner_message}  - done.      ")
        return result
    except KeyboardInterrupt:
        print(f"\r   -  {spinner_message}  - interrupted.")
        raise
    except Exception as exc:
        print(f"\r   -  {spinner_message}  - error.")
        return f"[ERROR] {exc}"


# ---------------------------------------------------------------------------
# Output writers
# ---------------------------------------------------------------------------
def _write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    print(f"  [WRITTEN] {path.name}")

def _add_inline_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))")
    for match in pattern.finditer(text):
        if match.group(2):
            paragraph.add_run(match.group(2)).bold = True
        elif match.group(3):
            paragraph.add_run(match.group(3)).italic = True
        elif match.group(4):
            paragraph.add_run(match.group(4))

def _write_docx(path: Path, content: str, title: str = "") -> None:
    if not _DOCX_AVAILABLE:
        print("  [WARN] python-docx not available -- skipping .docx export.")
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    for line in content.splitlines():
        s = line.rstrip()
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+[.)]\s", s):
            doc.add_paragraph(re.sub(r"^\d+[.)]\s*", "", s).strip(), style="List Number")
        elif s.startswith("---"):
            doc.add_paragraph("")
        else:
            _add_inline_runs(doc.add_paragraph(), s)
    disc = doc.add_paragraph()
    run  = disc.add_run(DISCLOSURE)
    run.italic = True
    run.font.size = Pt(9)
    doc.save(str(path))
    print(f"  [DOCX] {path.name}")

def _write_report(
    reports_path: Path,
    stem: str,
    role: str,
    mode: str,
    content: str,
    ts: str,
    extra_meta: str = "",
) -> Path:
    header = (
        f"# {role} Report\n"
        f"**Timestamp:** {ts}  \n"
        f"**Mode:** {mode}  \n"
        f"**Query:** {stem}  \n"
    )
    if extra_meta:
        header += extra_meta + "  \n"
    header += f"\n**Preview:** {content[:300]}...\n\n---\n\n"
    report_path = reports_path / f"{role}_{mode}_{stem}_{ts}.md"
    _write_text(report_path, header + content)
    return report_path

# ---------------------------------------------------------------------------
# Prompt builders
# ---------------------------------------------------------------------------
def _topic_system_prompt(guidelines: str) -> str:
    base = (
        "You are an expert medical research librarian and science communicator.\n"
        "You synthesise search results into clear, accurate research synopses.\n"
        "You write in plain English suitable for a medical professional audience.\n\n"
        "CRITICAL RULES:\n"
        "1. ONLY cite articles that appear in the ## Search Results section provided.\n"
        "2. Do NOT invent, fabricate, or hallucinate any article, author, or URL.\n"
        "3. Copy every URL exactly as given in the ## Search Results list.\n"
        "4. If a result has no URL, omit it from References entirely.\n"
        "5. Your References section must contain ONLY the articles listed below.\n"
    )
    if guidelines:
        return base + "\n\n## Guidelines\n\n" + guidelines
    return base

def _topic_user_prompt(query: str, results: list[dict]) -> str:
    results_text = ""
    ref_list = ""
    for i, r in enumerate(results, 1):
        results_text += (
            f"### Result {i}\n"
            f"**Title:** {r['title']}\n"
            f"**URL:** {r['url']}\n"
            f"**Snippet:** {r['snippet']}\n\n"
        )
        ref_list += f"{i}. [{r['title']}]({r['url']})\n"
    return (
        f"## Search Query\n{query}\n\n"
        f"## Source Articles (AUTHORITATIVE  - do NOT modify or add to this list)\n"
        f"{ref_list}\n"
        f"## Full Search Results\n{results_text}\n"
        "## Task\n"
        "Write a comprehensive synopsis of these search results.\n"
        "Structure your response with EXACTLY these four sections:\n"
        "1. ## Overview (2-3 paragraph summary of the topic)\n"
        "2. ## Key Findings (thematic bullet points from the articles above)\n"
        "3. ## Clinical Relevance (practical implications for clinicians)\n"
        "4. ## References\n"
        "IMPORTANT for section 4:\n"
        "- Copy the numbered list from ## Source Articles above VERBATIM.\n"
        "- Do NOT add any article not listed in ## Source Articles.\n"
        "- Do NOT invent authors, journals, PMIDs, or URLs.\n"
        "- The References section must be identical to ## Source Articles."
    )

def _article_system_prompt(guidelines: str) -> str:
    base = (
        "You are an expert medical research librarian specialising in evidence synthesis.\n"
        "You summarise PubMed articles accurately and objectively.\n"
        "You identify similarities and differences between related documents.\n"
        "You always include PubMed URLs and DOIs where available.\n"
    )
    if guidelines:
        return base + "\n\n## Guidelines\n\n" + guidelines
    return base

def _article_user_prompt(
    query: str,
    article_type: str,
    articles: list[dict],
) -> str:
    articles_text = ""
    for i, a in enumerate(articles, 1):
        doi_str = f"DOI: {a['doi']}  " if a["doi"] else ""
        articles_text += (
            f"### Article {i}\n"
            f"**Title:** {a['title']}\n"
            f"**Authors:** {a['authors']}\n"
            f"**Year:** {a['year']}  "
            f"**Journal:** {a['journal']}\n"
            f"{doi_str}**URL:** [{a['url']}]({a['url']})\n"
            f"**Abstract:** {a['abstract'] or '[Abstract not available]'}\n\n"
        )
    comparison_instruction = ""
    if len(articles) >= 2:
        comparison_instruction = (
            "\n5. ## Similarities and Differences\n"
            "   Compare the 2 most similar articles. Summarise:\n"
            "   - Shared findings, methodology, and recommendations\n"
            "   - Key differences in scope, population, outcomes, or conclusions\n"
        )
    return (
        f"## Search Query\n{query}\n\n"
        f"## Article Type\n{article_type}\n\n"
        f"## Articles Found\n{articles_text}\n"
        "## Task\n"
        "Write a structured summary report containing:\n"
        "1. ## Overview (brief description of search and results)\n"
        "2. ## Article Summaries (150-200 words each, include URL)\n"
        "3. ## Key Themes (common themes across articles)\n"
        "4. ## References (numbered list with markdown links)\n"
        f"{comparison_instruction}"
    )

# ---------------------------------------------------------------------------
# Main runners
# ---------------------------------------------------------------------------
def run_topic_search(
    direct_instructions: list[str],
    call_llm_fn,
    verbose: bool = True,
    model: str | None = None,
) -> None:
    """Run Topic Search sub-mode."""
    root  = _project_root()
    paths = _paths(root)
    ts    = _ts()

    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)

    guidelines = _load_guidelines(paths["docs"])

    # Get query from direct instructions or prompt user
    if direct_instructions:
        query = " ".join(direct_instructions).lstrip("0123456789. ").strip()
    else:
        query = input("  Enter search topic: ").strip()
        if not query:
            print("  [SEARCH] No query entered. Aborting.")
            return

    print(f"\n  [TOPIC SEARCH] Query: {query}")

    # Web search
    topic_limit, _ = _result_limits(model)
    with _Spinner("Searching Europe PMC"):
        results = _europepmc_search(query, max_results=topic_limit)

    print(f"  [SEARCH] Found {len(results)} web results.")

    if not results:
        print("  [SEARCH] No results found. Try a different query.")
        return


    # Build ground-truth references from fetched results (used to overwrite LLM refs)
    ref_lines = []
    for i, r in enumerate(results, 1):
        title = r['title']
        url   = r['url']
        ref_lines.append(f"{i}. [{title}]({url})")
    ground_truth_refs = "## References\n\n" + "\n".join(ref_lines) + "\n"

    # LLM synopsis
    system_prompt = _topic_system_prompt(guidelines)
    user_prompt   = _topic_user_prompt(query, results)
    synopsis      = _call_llm(
        system_prompt, user_prompt, call_llm_fn,
        spinner_message="Generating synopsis",
    )

    # Strip LLM-generated References section and replace with ground-truth URLs
    synopsis = re.sub(
        r'##\s*References[\s\S]*',
        '',
        synopsis,
        flags=re.IGNORECASE,
    ).rstrip()
    synopsis = synopsis + "\n\n" + ground_truth_refs


    # Sanitise stem for filename
    stem = re.sub(r"[^\w\-]", "_", query[:50]).strip("_")

    # Write outputs
    md_path   = paths["output"] / f"TOPIC_{stem}_{ts}.md"
    docx_path = paths["output"] / f"TOPIC_{stem}_{ts}.docx"
    _write_text(md_path, synopsis)
    _write_docx(docx_path, synopsis, title=f"Topic Search: {query}")

    # Write report log
    _write_report(
        paths["reports"], stem, "TOPIC_SEARCH", "topic",
        synopsis, ts,
        extra_meta=f"**Results found:** {len(results)}",
    )

    print(f"\n  [TOPIC SEARCH] Complete.")
    print(f"  Output : output/search/TOPIC_{stem}_{ts}.md")


def run_article_search(
    direct_instructions: list[str],
    call_llm_fn,
    verbose: bool = True,
    model: str | None = None,
) -> None:
    """Run Article Search sub-mode."""
    root  = _project_root()
    paths = _paths(root)
    ts    = _ts()

    for p in paths.values():
        p.mkdir(parents=True, exist_ok=True)

    guidelines = _load_guidelines(paths["docs"])

    # Get query
    if direct_instructions:
        query = " ".join(direct_instructions)
    else:
        query = input("  Enter search query: ").strip()
        if not query:
            print("  [SEARCH] No query entered. Aborting.")
            return

    # Select article type
    print("\n  Select article type:")
    for k, (label, _) in ARTICLE_TYPES.items():
        print(f"    {k}. {label}")
    choice = input("  Choice: ").strip()

    if choice in ARTICLE_TYPES:
        article_type, pubmed_filter = ARTICLE_TYPES[choice]
        if choice == "8":  # Custom
            article_type  = input("  Enter custom article type: ").strip()
            pubmed_filter = f"{article_type}[tiab]"
    else:
        print("  [SEARCH] Invalid choice. Using Review Article as default.")
        article_type  = "Review Article"
        pubmed_filter = "review[pt]"

    print(f"\n  [ARTICLE SEARCH] Query: {query} | Type: {article_type}")

    # PubMed search
    api_key = ""
    try:
        import os
        api_key = os.environ.get("PUBMED_API_KEY", "")
    except Exception:
        pass

    with _Spinner("Searching PubMed"):
        articles = _pubmed_search(
            query, pubmed_filter,
            max_results=MAX_RESULTS_ARTICLE,
            api_key=api_key,
        )

    print(f"  [SEARCH] Found {len(articles)} articles.")

    if not articles:
        print("  [SEARCH] No articles found. Try a different query or article type.")
        return

    # Fetch abstracts
    with _Spinner("Fetching abstracts"):
        articles = _pubmed_fetch_abstracts(articles, api_key=api_key)

    # LLM summary
    system_prompt = _article_system_prompt(guidelines)
    user_prompt   = _article_user_prompt(query, article_type, articles)
    summary       = _call_llm(
        system_prompt, user_prompt, call_llm_fn,
        spinner_message="Generating article summary",
    )

    # Sanitise stem
    stem = re.sub(r"[^\w\-]", "_", f"{article_type}_{query}"[:60]).strip("_")

    # Write outputs
    md_path   = paths["output"] / f"ARTICLE_{stem}_{ts}.md"
    docx_path = paths["output"] / f"ARTICLE_{stem}_{ts}.docx"
    _write_text(md_path, summary)
    _write_docx(docx_path, summary, title=f"Article Search: {query} ({article_type})")

    # Write report log
    _write_report(
        paths["reports"], stem, "ARTICLE_SEARCH", "article",
        summary, ts,
        extra_meta=f"**Article type:** {article_type}  \n**Articles found:** {len(articles)}",
    )

    print(f"\n  [ARTICLE SEARCH] Complete.")
    print(f"  Output : output/search/ARTICLE_{stem}_{ts}.md")


