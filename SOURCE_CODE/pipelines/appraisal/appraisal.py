"""
src/modes/appraisal.py
======================
Appraisal mode engine ??critical appraisal of medical literature.

Pipeline (per article):
    APPRAISER  ?? output/appraisal/APPRAISAL_[stem]_[ts].md + .docx
                  reports/appraisal/APPRAISAL_[stem]_[ts].md  (process log)
                  reports/appraisal/APPRAISAL_SESSION_[ts]_SUMMARY.md
"""

from __future__ import annotations

import datetime
import sys
import threading
from pathlib import Path
from typing import Callable

# Add SOURCE_CODE to path
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent.parent.parent.parent
SOURCE_CODE_DIR = PROJECT_ROOT / "SOURCE_CODE"
sys.path.insert(0, str(SOURCE_CODE_DIR))

from utils.path_utils import PATH_MANAGER, get_input_dir, get_output_dir
from utils.document_reader import DocumentReader
from utils.rag import RAGUtils

# ---------------------------------------------------------------------------
# Optional dependencies
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
DISCLOSURE = (
    "This appraisal was produced with AI assistance. "
    "All ratings and judgements must be verified by a qualified clinician "
    "before being used to inform clinical decisions."
)

DOCS_DIR_NAME  = "appraisal"
INPUT_DIR_NAME = "appraisal"

# ---------------------------------------------------------------------------
# Path helpers
# ---------------------------------------------------------------------------
def _project_root() -> Path:
    return Path(__file__).resolve().parent.parent.parent.parent.parent


def _ts() -> str:
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")


def _paths(root: Path) -> dict:
    return {
        "doc_root":    PROJECT_ROOT / "docs"    / DOCS_DIR_NAME,
        "input":   PROJECT_ROOT / "input"   / INPUT_DIR_NAME,
        "output":  PROJECT_ROOT / "output"  / INPUT_DIR_NAME,
        "reports": PROJECT_ROOT / "reports" / INPUT_DIR_NAME,
    }


# ---------------------------------------------------------------------------
# Spinner
# ---------------------------------------------------------------------------
class _Spinner:
    def __init__(self, message: str = "Processing"):
        self._message = message
        self._stop_event = threading.Event()
        self._thread = threading.Thread(target=self._spin, daemon=True)

    def _spin(self) -> None:
        frames = ["|", "/", "-", "\\"]
        idx = 0
        while not self._stop_event.is_set():
            sys.stdout.write(f"\r  {frames[idx % len(frames)]}  {self._message}... ")
            sys.stdout.flush()
            self._stop_event.wait(0.15)
            idx += 1
        sys.stdout.write(f"\r  OK  {self._message} -- done.          \n")
        sys.stdout.flush()

    def __enter__(self):
        self._thread.start()
        return self

    def __exit__(self, *_):
        self._stop_event.set()
        self._thread.join()


# ---------------------------------------------------------------------------
# LLM wrapper
# ---------------------------------------------------------------------------
def _call_llm(
    system_prompt: str,
    user_prompt: str,
    call_llm_fn: Callable,
    spinner_message: str = "Appraising",
) -> str:
    with _Spinner(spinner_message):
        try:
            result = call_llm_fn(system_prompt, user_prompt)
        except Exception as exc:
            result = f"[ERROR] LLM call failed: {exc}"
    return result


# ---------------------------------------------------------------------------
# I/O helpers
# ---------------------------------------------------------------------------
def _load_guidelines(docs_dir: Path) -> str:
    """Load all .md files from docs/appraisal/ as guidelines."""
    parts: list[str] = []
    if not docs_dir.exists():
        return ""
    for f in sorted(docs_dir.rglob("*.md")):
        parts.append(f"### {f.name}\n{f.read_text(encoding='utf-8')}")
    return "\n\n".join(parts)


def _load_input_files(input_dir: Path) -> list:
    """Load all supported files from input_dir.
    Returns list of (stem, content, suffix) tuples.
    """
    extensions = {".md", ".txt", ".docx", ".pdf"}
    results: list = []
    if not input_dir.exists():
        return results
    for f in sorted(input_dir.iterdir()):
        if not f.is_file() or f.suffix.lower() not in extensions:
            continue
        suffix = f.suffix.lower()
        stem   = f.stem
        if suffix == ".docx":
            try:
                doc = Document(str(f))
                content = "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                content = f"[Could not read {f.name}]"
        elif suffix == ".pdf":
            try:
                import fitz
                pdf = fitz.open(str(f))
                content = "\n\n".join(page.get_text() for page in pdf)
                pdf.close()
                if not content.strip():
                    content = f"[PDF {f.name} contained no extractable text]"
            except ImportError:
                content = "[PDF reading requires PyMuPDF: pip install pymupdf]"
            except Exception as exc:
                content = f"[Could not read PDF {f.name}: {exc}]"
        else:
            content = f.read_text(encoding="utf-8", errors="replace")
        results.append((stem, content, suffix))
    return results


def _write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def _add_inline_runs(paragraph, text: str) -> None:
    import re
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))")
    for match in pattern.finditer(text):
        if match.group(2):
            paragraph.add_run(match.group(2)).bold = True
        elif match.group(3):
            paragraph.add_run(match.group(3)).italic = True
        elif match.group(4):
            paragraph.add_run(match.group(4))


def _write_docx(path: Path, content: str, title: str = "") -> None:
    if not _DOCX_AVAILABLE:
        print("  [WARN] python-docx not available -- skipping .docx export.")
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    for line in content.splitlines():
        s = line.rstrip()
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif len(s) > 2 and s[0].isdigit() and s[1] in ".)":
            doc.add_paragraph(s[2:].strip(), style="List Number")
        elif s.startswith("---"):
            doc.add_paragraph("")
        else:
            _add_inline_runs(doc.add_paragraph(), s)
    disc = doc.add_paragraph()
    run  = disc.add_run(DISCLOSURE)
    run.italic = True
    run.font.size = Pt(9)
    doc.save(str(path))


def _strip_fences(text: str) -> str:
    import re
    text = re.sub(r"^```[a-zA-Z]*\n", "", text.strip())
    text = re.sub(r"\n```$", "", text.strip())
    return text.strip()


# ---------------------------------------------------------------------------
# Prompt builders
# ---------------------------------------------------------------------------
def _system_prompt(guidelines: str) -> str:
    return (
        "You are an medical research appraisal expert performing a "
        "critical appraisal of a medical research article.\n\n"
        "You must follow the appraisal guide and scoring criteria exactly. "
        "Produce a structured 7-section appraisal report followed by a "
        "plain-language summary. Apply GRADE evidence rating. "
        "If a custom checklist is present in the guidelines, apply it as well.\n\n"
        "## Appraisal Guidelines\n\n"
        f"{guidelines}"
    )


def _appraiser_user_prompt(
    direct_instructions: list,
    article_content: str,
    stem: str,
) -> str:
    instr_block = ""
    if direct_instructions:
        instr_block = (
            "## DIRECT TASK INSTRUCTIONS\n"
            + "\n".join(
                i.lstrip("> ").strip() if isinstance(i, str) else str(i)
                for i in direct_instructions
            )
            + "\n\n"
        )
    return (
        f"{instr_block}"
        f"## Article to Appraise: {stem}\n\n"
        f"{article_content}\n\n"
        "## Task\n"
        "Produce a complete critical appraisal report following the mandatory "
        "7-section structure defined in your guidelines. "
        "End with a plain-language summary of ??200 words. "
        "Append the disclosure statement at the very end."
    )


# ---------------------------------------------------------------------------
# Report helper
# ---------------------------------------------------------------------------
def _write_process_log(
    reports_path: Path,
    stem: str,
    timestamp: str,
    content: str,
) -> Path:
    """Write a brief process log to reports/appraisal/."""
    filename   = f"APPRAISAL_{stem}_{timestamp}.md"
    word_count = len(content.split())
    preview    = content[:300].replace("\n", " ").strip()
    if len(content) > 300:
        preview += "..."
    body = (
        f"# Appraisal Process Log ??{stem}\n"
        f"**Timestamp:** {timestamp}  \n"
        f"**Word count:** {word_count}  \n\n"
        "---\n\n"
        f"## Preview\n\n{preview}\n\n"
        "---\n\n"
        f"*Full report saved to output/appraisal/APPRAISAL_{stem}_{timestamp}.docx*\n"
    )
    path = reports_path / filename
    _write_text(path, body)
    print(f"  [LOG]    {filename}")
    return path


def _write_session_summary(
    reports_path: Path,
    timestamp: str,
    articles: list,
) -> None:
    filename = f"APPRAISAL_SESSION_{timestamp}_SUMMARY.md"
    lines = [
        "# Appraisal Session Summary\n",
        f"**Timestamp:** {timestamp}  ",
        f"**Articles appraised:** {len(articles)}\n",
        "---\n",
    ]
    for i, a in enumerate(articles, 1):
        lines.append(f"### {i}. {a['stem']}")
        lines.append(f"- Output: `output/appraisal/APPRAISAL_{a['stem']}_{a['ts']}.md`")
        lines.append(f"- DOCX:   `output/appraisal/APPRAISAL_{a['stem']}_{a['ts']}.docx`")
        lines.append(f"- Log:    `reports/appraisal/APPRAISAL_{a['stem']}_{a['ts']}.md`\n")
    _write_text(reports_path / filename, "\n".join(lines))
    print(f"  [SUMMARY] {filename}")


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------
def run_appraisal(
    direct_instructions: list,
    call_llm_fn: Callable,
    verbose: bool = True,
) -> None:
    root      = _project_root()
    paths     = _paths(root)
    timestamp = _ts()

    for p in [paths["input"], paths["output"], paths["reports"]]:
        p.mkdir(parents=True, exist_ok=True)

    input_files = _load_input_files(paths["input"])
    if not input_files:
        print("  [APPRAISAL] No files in input/appraisal/ -- place articles there first.")
        return

    guidelines = _load_guidelines(paths["doc_root"])
    session_log: list[dict] = []

    print(f"\n  [APPRAISAL] Found {len(input_files)} article(s) to appraise.\n")

    for stem, content, suffix in input_files:
        if verbose:
            print(f"  {'='*56}")
            print(f"  Appraising: {stem}{suffix}")
            print(f"  {'='*56}")

        appraisal_out = _call_llm(
            _system_prompt(guidelines),
            _appraiser_user_prompt(direct_instructions, content, stem),
            call_llm_fn,
            f"Appraising {stem[:40]}",
        )
        appraisal_out = _strip_fences(appraisal_out)

        # Write full report to output/
        out_md   = paths["output"] / f"APPRAISAL_{stem}_{timestamp}.md"
        out_docx = paths["output"] / f"APPRAISAL_{stem}_{timestamp}.docx"
        _write_text(out_md, appraisal_out)
        _write_docx(out_docx, appraisal_out,
                    title=f"Critical Appraisal ??{stem}")

        # Write process log to reports/
        _write_process_log(paths["reports"], stem, timestamp, appraisal_out)

        if verbose:
            print(f"  [MD]     output/appraisal/APPRAISAL_{stem}_{timestamp}.md")
            print(f"  [DOCX]   output/appraisal/APPRAISAL_{stem}_{timestamp}.docx")

        session_log.append({"stem": stem, "ts": timestamp})

    _write_session_summary(paths["reports"], timestamp, session_log)
    print(f"\n  [APPRAISAL] Done. {len(session_log)} article(s) processed.")
    print("  Check output/appraisal/ for reports and reports/appraisal/ for logs.\n")


# ---------------------------------------------------------------------------
# Direct-instruction parser (shared utility)
# ---------------------------------------------------------------------------
def parse_direct_instructions(raw: str) -> list:
    lines = []
    for line in raw.splitlines():
        stripped = line.strip()
        if stripped.startswith(">"):
            lines.append(stripped)
    return lines



