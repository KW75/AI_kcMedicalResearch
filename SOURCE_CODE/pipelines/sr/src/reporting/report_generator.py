import logging
from pathlib import Path
logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _OK = True
except ImportError:
    _OK = False

class ReportGenerator:
    def generate(self, title, authors, pico, ma_result, extraction_results,
                 screening_results, forest_plot_path, effect_measure="OR",
                 output_path="outputs/reports/systematic_review.docx"):
        if not _OK:
            logger.error("python-docx required."); return
        doc = Document()
        h = doc.add_heading(title, 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if authors:
            p=doc.add_paragraph(authors); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading("Abstract",1)
        doc.add_paragraph(
            f"k={ma_result.get('k','?')} RCTs included. "
            f"Pooled {effect_measure}={ma_result.get('pooled_effect',0):.3f} "
            f"(95% CI {ma_result.get('ci_lower',0):.3f}-{ma_result.get('ci_upper',0):.3f}). "
            f"I\u00b2={ma_result.get('I2',0):.1f}%, \u03c4\u00b2={ma_result.get('tau2',0):.4f}.")
        doc.add_heading("Methods",1)
        doc.add_heading("Eligibility Criteria",2)
        doc.add_paragraph(" | ".join(f"{k.upper()}: {v}" for k,v in pico.items()))
        doc.add_heading("Results",1)
        doc.add_heading("Meta-Analysis",2)
        doc.add_paragraph(
            f"k={ma_result.get('k','?')} studies. "
            f"Pooled {effect_measure}={ma_result.get('pooled_effect',0):.3f} "
            f"(95% CI {ma_result.get('ci_lower',0):.3f}-{ma_result.get('ci_upper',0):.3f}; "
            f"p={ma_result.get('p_value',0):.4f}). "
            f"I\u00b2={ma_result.get('I2',0):.1f}%.")
        if forest_plot_path and Path(forest_plot_path).exists():
            doc.add_heading("Forest Plot",2)
            doc.add_picture(forest_plot_path, width=Inches(6.0))
        doc.add_heading("Discussion",1); doc.add_paragraph("[Author interpretation required]")
        doc.add_heading("Conclusion",1); doc.add_paragraph("[Author conclusion required]")
        out=Path(output_path); out.parent.mkdir(parents=True,exist_ok=True)
        doc.save(str(out)); logger.info(f"DOCX saved: {out}")
