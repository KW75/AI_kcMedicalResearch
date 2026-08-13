# SOURCE_CODE/utils/document_reader.py
import os
import re
from pathlib import Path
from typing import List, Dict, Optional, Union, Any
import json
from PIL import Image
try:
    import pytesseract
except ImportError:
    pytesseract = None  # OCR unavailable (CI/server)
import io

# Document-specific imports
try:
    from docx import Document
    import docx2txt
except ImportError:
    Document = None
    docx2txt = None

try:
    import easyocr
except ImportError:
    easyocr = None

try:
    import cv2
except ImportError:
    cv2 = None

# PDF imports
try:
    import pypdf
    from pdf2image import convert_from_path
except ImportError:
    pypdf = None
    convert_from_path = None


class DocumentReader:
    """Universal document reader for all file formats"""
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'word': ['.doc', '.docx'],
        'image': ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif'],
        'text': ['.txt', '.md', '.rtf'],
        'excel': ['.xls', '.xlsx'],
    }
    
    @classmethod
    def detect_format(cls, file_path: Union[str, Path]) -> str:
        """Detect the format of a file"""
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        
        for format_type, extensions in cls.SUPPORTED_FORMATS.items():
            if ext in extensions:
                return format_type
        
        return 'unknown'
    
    @classmethod
    def extract_text(cls, file_path: Union[str, Path], **kwargs) -> Dict[str, Any]:
        """Extract text from any supported document format"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                'success': False,
                'error': f'File not found: {file_path}'
            }
        
        format_type = cls.detect_format(file_path)
        result = {
            'file': str(file_path),
            'format': format_type,
            'success': True,
            'text': '',
            'metadata': {},
            'pages': 0
        }
        
        try:
            if format_type == 'pdf':
                result.update(cls._extract_pdf(file_path, **kwargs))
            elif format_type == 'word':
                result.update(cls._extract_word(file_path, **kwargs))
            elif format_type == 'image':
                result.update(cls._extract_image(file_path, **kwargs))
            elif format_type == 'text':
                result.update(cls._extract_text_file(file_path, **kwargs))
            else:
                result.update({
                    'success': False,
                    'error': f'Unsupported format: {format_type}'
                })
        
        except Exception as e:
            result.update({
                'success': False,
                'error': str(e)
            })
        
        return result
    
    @classmethod
    def _extract_pdf(cls, pdf_path: Path, **kwargs) -> Dict:
        """Extract text from PDF"""
        result = {'pages': 0, 'text': '', 'metadata': {}}
        use_ocr = kwargs.get('use_ocr', False)
        
        try:
            # Try text extraction first
            if not use_ocr and pypdf:
                with open(pdf_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    result['pages'] = len(reader.pages)
                    
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            result['text'] += page_text + '\n\n'
                    
                    # Extract metadata
                    if reader.metadata:
                        result['metadata'] = {
                            'title': reader.metadata.get('/Title', ''),
                            'author': reader.metadata.get('/Author', ''),
                            'subject': reader.metadata.get('/Subject', ''),
                            'keywords': reader.metadata.get('/Keywords', ''),
                            'creator': reader.metadata.get('/Creator', ''),
                            'producer': reader.metadata.get('/Producer', ''),
                            'creation_date': str(reader.metadata.get('/CreationDate', '')),
                            'modification_date': str(reader.metadata.get('/ModDate', ''))
                        }
            
            # If no text or OCR requested, try OCR
            if not result['text'].strip() or use_ocr:
                if convert_from_path:
                    images = convert_from_path(pdf_path)
                    result['pages'] = len(images)
                    ocr_text = ''
                    
                    for i, image in enumerate(images):
                        # Preprocess image for better OCR
                        processed = cls._preprocess_image_for_ocr(image)
                        page_text = pytesseract.image_to_string(processed)
                        ocr_text += page_text + '\n\n'
                    
                    result['text'] = ocr_text
                    result['metadata']['ocr_used'] = True
        
        except Exception as e:
            result['error'] = str(e)
            result['success'] = False
        
        return result
    
    @classmethod
    def _extract_word(cls, doc_path: Path, **kwargs) -> Dict:
        """Extract text from Word document"""
        result = {'pages': 0, 'text': '', 'metadata': {}}
        
        try:
            if doc_path.suffix.lower() == '.docx':
                # DOCX extraction
                if Document:
                    doc = Document(doc_path)
                    result['text'] = '\n\n'.join([p.text for p in doc.paragraphs])
                    
                    # Extract tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join([cell.text for cell in row.cells])
                            result['text'] += row_text + '\n'
                    
                    # Extract metadata
                    core_props = doc.core_properties
                    result['metadata'] = {
                        'title': core_props.title or '',
                        'author': core_props.author or '',
                        'subject': core_props.subject or '',
                        'keywords': core_props.keywords or '',
                        'category': core_props.category or '',
                        'comments': core_props.comments or '',
                        'created': str(core_props.created) if core_props.created else '',
                        'modified': str(core_props.modified) if core_props.modified else '',
                        'last_modified_by': core_props.last_modified_by or ''
                    }
                    
                    # Estimate pages (roughly 250 words per page)
                    result['pages'] = max(1, len(result['text'].split()) // 250)
                
                elif docx2txt:
                    result['text'] = docx2txt.process(str(doc_path))
                    result['pages'] = max(1, len(result['text'].split()) // 250)
            
            else:
                # .doc (old format) - fallback to textract
                try:
                    import textract
                    result['text'] = textract.process(str(doc_path)).decode('utf-8')
                    result['pages'] = max(1, len(result['text'].split()) // 250)
                except:
                    result['text'] = cls._extract_with_alternative(doc_path)
        
        except Exception as e:
            result['error'] = str(e)
            result['success'] = False
        
        return result
    
    @classmethod
    def _extract_image(cls, image_path: Path, **kwargs) -> Dict:
        """Extract text from image using OCR"""
        result = {'pages': 0, 'text': '', 'metadata': {}}
        
        try:
            # Load image
            image = Image.open(image_path)
            result['pages'] = 1
            
            # Extract metadata
            result['metadata'] = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.width,
                'height': image.height,
                'filename': image_path.name
            }
            
            # Perform OCR
            processed = cls._preprocess_image_for_ocr(image)
            result['text'] = pytesseract.image_to_string(processed)
            
            # Try EasyOCR if available and text is poor
            if len(result['text'].strip()) < 50 and easyocr:
                reader = easyocr.Reader(['en'])
                result_easy = reader.readtext(str(image_path))
                if result_easy:
                    easy_text = '\n'.join([item[1] for item in result_easy])
                    if len(easy_text.strip()) > len(result['text'].strip()):
                        result['text'] = easy_text
        
        except Exception as e:
            result['error'] = str(e)
            result['success'] = False
        
        return result
    
    @classmethod
    def _extract_text_file(cls, txt_path: Path, **kwargs) -> Dict:
        """Extract text from plain text file"""
        result = {'pages': 0, 'text': '', 'metadata': {}}
        
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as f:
                        result['text'] = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            result['pages'] = max(1, len(result['text'].split()) // 250)
            result['metadata'] = {
                'filename': txt_path.name,
                'size': txt_path.stat().st_size,
                'encoding': encoding if 'encoding' in locals() else 'unknown'
            }
        
        except Exception as e:
            result['error'] = str(e)
            result['success'] = False
        
        return result
    
    @classmethod
    def _preprocess_image_for_ocr(cls, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR results"""
        try:
            if cv2:
                # Convert PIL to OpenCV
                import numpy as np
                img_array = np.array(image)
                img = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                
                # Grayscale
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                
                # Binarization (thresholding)
                _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                
                # Denoising
                denoised = cv2.fastNlMeansDenoising(thresh, h=10)
                
                # Convert back to PIL
                return Image.fromarray(denoised)
            else:
                # Simple fallback: grayscale
                if image.mode != 'L':
                    return image.convert('L')
                return image
        
        except:
            return image
    
    @classmethod
    def _extract_with_alternative(cls, doc_path: Path) -> str:
        """Fallback extraction using alternative methods"""
        try:
            # Try reading as binary and extracting text patterns
            with open(doc_path, 'rb') as f:
                content = f.read()
                # Try to extract readable text
                text = content.decode('latin-1', errors='ignore')
                # Remove non-printable characters
                text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)
                # Clean up multiple spaces
                text = re.sub(r'\s+', ' ', text)
                return text.strip()
        except:
            return ""
    
    @classmethod
    def batch_extract(cls, directory: Union[str, Path], **kwargs) -> List[Dict]:
        """Extract text from all supported files in a directory"""
        directory = Path(directory)
        results = []
        
        for ext in ['.pdf', '.docx', '.doc', '.txt', '.md', '.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            for file_path in directory.glob(f'*{ext}'):
                result = cls.extract_text(file_path, **kwargs)
                results.append(result)
        
        return results
    
    @classmethod
    def get_file_info(cls, file_path: Union[str, Path]) -> Dict:
        """Get detailed information about a file"""
        file_path = Path(file_path)
        
        info = {
            'filename': file_path.name,
            'extension': file_path.suffix,
            'size': file_path.stat().st_size,
            'format': cls.detect_format(file_path),
            'path': str(file_path),
            'exists': file_path.exists()
        }
        
        return info
    
    @classmethod
    def chunk_text(cls, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
        
        return chunks