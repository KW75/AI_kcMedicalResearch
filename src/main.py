"""
main.py — AI Automation Tool  v2.2.0
Supports six workflow modes: coding, writing, rct_search, appraisal, search, sr.
Supports six AI providers: ollama (default), openai, anthropic, deepseek, groq, qwen.
RAG layer: per-session, mode-specific input/ folder indexing via rag.py.
Coding mode: Builder (pipeline), Reviewer (standalone), Tester (standalone).
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
from urllib.request import urlopen  # bare name so tests can patch src.main.urlopen
import uuid
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
try:
    # When imported as a package by pytest: from src.main import ...
    from src.modes.coding import (
        run_builder,
        run_reviewer,
        run_tester,
        parse_direct_instructions,
    )
except ModuleNotFoundError:
    # When run directly: python src/main.py
    from modes.coding import (
        run_builder,
        run_reviewer,
        run_tester,
        parse_direct_instructions,
    )


load_dotenv()

# Ensure project root is on sys.path so `from src import rag` works
_ROOT = Path(__file__).resolve().parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))


# ---------------------------------------------------------------------------
# Version
# ---------------------------------------------------------------------------
VERSION = "2.2.0"

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE_DIR            = Path(__file__).resolve().parent.parent   # D:\AI_kcMedicalResearch

# Docs folders (guidance .md files per mode)
DOCS_DIR            = BASE_DIR / "docs"
DOCS_CODING         = DOCS_DIR / "coding"
DOCS_WRITING        = DOCS_DIR / "writing"
DOCS_APPRAISAL      = DOCS_DIR / "appraisal"
DOCS_RCT_SEARCH     = DOCS_DIR / "rct_search"
DOCS_SEARCH         = DOCS_DIR / "search"
DOCS_SR             = DOCS_DIR / "sr"

# AI prompt files
AI_DIR              = BASE_DIR / "ai"

# Input folders (auto-loaded at startup per mode)
INPUT_DIR           = BASE_DIR / "input"
INPUT_CODING        = INPUT_DIR / "coding"
INPUT_WRITING       = INPUT_DIR / "writing"
INPUT_APPRAISAL     = INPUT_DIR / "appraisal"
INPUT_RCT_SEARCH    = INPUT_DIR / "rct_search"
INPUT_SEARCH        = INPUT_DIR / "search"
INPUT_SR            = INPUT_DIR / "sr"

# Output folders (deliverables)
OUTPUT_DIR          = BASE_DIR / "output"
OUTPUT_CODING       = OUTPUT_DIR / "coding"
OUTPUT_WRITING      = OUTPUT_DIR / "writing"
OUTPUT_APPRAISAL    = OUTPUT_DIR / "appraisal"
OUTPUT_RCT_SEARCH   = OUTPUT_DIR / "rct_search"
OUTPUT_SEARCH       = OUTPUT_DIR / "search"
OUTPUT_SR           = OUTPUT_DIR / "sr"

# Reports folder (session transcripts / operation logs)
REPORTS_DIR         = BASE_DIR / "reports"

# Ensure all directories exist at startup
for _d in [
    DOCS_CODING, DOCS_WRITING, DOCS_APPRAISAL,
    DOCS_RCT_SEARCH, DOCS_SEARCH, DOCS_SR,
    INPUT_CODING, INPUT_WRITING, INPUT_APPRAISAL,
    INPUT_RCT_SEARCH, INPUT_SEARCH, INPUT_SR,
    OUTPUT_CODING, OUTPUT_WRITING, OUTPUT_APPRAISAL,
    OUTPUT_RCT_SEARCH, OUTPUT_SEARCH, OUTPUT_SR,
    REPORTS_DIR,
    # reports/ subfolders — each mode writes here
    REPORTS_DIR / "coding",
    REPORTS_DIR / "writing",
    REPORTS_DIR / "appraisal",
    REPORTS_DIR / "rct_search",
    REPORTS_DIR / "search",
    REPORTS_DIR / "systematic_review",
    REPORTS_DIR / "transcripts",
]:
    _d.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Auto-load input files
# ---------------------------------------------------------------------------
import mimetypes  # noqa: E402  (after path setup)

_MODE_EXTENSIONS = {
    "coding":     {".py", ".js", ".ts", ".html", ".css", ".java", ".c",
                   ".cpp", ".cs", ".rb", ".go", ".rs", ".txt", ".md",
                   ".php", ".swift", ".kt", ".r", ".sh", ".sql", ".svg"},
    "writing":    {".txt", ".md", ".docx", ".pdf"},
    "appraisal":  {".pdf", ".txt", ".md", ".docx"},
    "rct_search": {".txt", ".md", ".pdf", ".docx"},
    "search":     {".txt", ".md"},
    "sr":         {".pdf"},
}


def auto_load_input_files(mode: str) -> list[Path]:
    """
    Scan input/<mode>/ and return a sorted list of accepted files.
    Prints what was found. Returns empty list if folder is empty.
    """
    folder_map = {
        "coding":     INPUT_CODING,
        "writing":    INPUT_WRITING,
        "appraisal":  INPUT_APPRAISAL,
        "rct_search": INPUT_RCT_SEARCH,
        "search":     INPUT_SEARCH,
        "sr":         INPUT_SR,
    }
    folder  = folder_map.get(mode, INPUT_DIR / mode)
    allowed = _MODE_EXTENSIONS.get(mode, set())
    files   = sorted(
        p for p in folder.iterdir()
        if p.is_file() and p.suffix.lower() in allowed
    )
    if files:
        print(f"[auto-load] Found {len(files)} file(s) in input/{mode}/:")
        for f in files:
            print(f"  • {f.name}")
    else:
        print(f"[auto-load] No input files in input/{mode}/ — proceeding without pre-loaded context.")
    return files


# ---------------------------------------------------------------------------
# Environment / provider config
# ---------------------------------------------------------------------------
OLLAMA_HOST         = os.getenv("OLLAMA_HOST",         "http://localhost:11434")
OLLAMA_MODEL        = os.getenv("OLLAMA_MODEL",        "llama3.2")
OPENAI_API_KEY      = os.getenv("OPENAI_API_KEY",      "")
OPENAI_MODEL        = os.getenv("OPENAI_MODEL",        "gpt-4o-mini")
ANTHROPIC_API_KEY   = os.getenv("ANTHROPIC_API_KEY",   "")
ANTHROPIC_MODEL     = os.getenv("ANTHROPIC_MODEL",     "claude-sonnet-5")
DEEPSEEK_API_KEY    = os.getenv("DEEPSEEK_API_KEY",    "")
DEEPSEEK_MODEL      = os.getenv("DEEPSEEK_MODEL",      "deepseek-v4-flash")
GROQ_API_KEY        = os.getenv("GROQ_API_KEY",        "")
GROQ_MODEL          = os.getenv("GROQ_MODEL",          "llama-3.3-70b-versatile")
DASHSCOPE_API_KEY   = os.environ.get("DASHSCOPE_API_KEY", "")
DASHSCOPE_BASE_URL  = os.environ.get(
    "DASHSCOPE_BASE_URL",
    "https://dashscope-intl.aliyuncs.com/compatible-mode/v1",
)
QWEN_MODEL          = "qwen3.7-plus"
EMBEDDING_PROVIDER  = os.getenv("EMBEDDING_PROVIDER",  "ollama")
EMBEDDING_MODEL     = os.getenv("EMBEDDING_MODEL",     "nomic-embed-text")


# ---------------------------------------------------------------------------
# Role / mode definitions
# ---------------------------------------------------------------------------
ROLE_FILES_CODING = {
    "Builder":  {"prompt": AI_DIR / "builder-prompt.md",  "report": "builder-report.md"},
    "Reviewer": {"prompt": AI_DIR / "reviewer-prompt.md", "report": "reviewer-report.md"},
    "Tester":   {"prompt": AI_DIR / "tester-prompt.md",   "report": "tester-report.md"},
}

ROLE_FILES_WRITING = {
    "Writer": {"prompt": AI_DIR / "writer-prompt.md", "report": "writer-report.md"},
    "Editor": {"prompt": AI_DIR / "editor-prompt.md", "report": "editor-report.md"},
    "QA":     {"prompt": AI_DIR / "qa-prompt.md",     "report": "qa-report.md"},
}

ROLE_FILES_RCT_SEARCH = {
    "Formulator": {"prompt": AI_DIR / "formulator-prompt.md", "report": "formulator-report.md"},
    "Searcher":   {"prompt": AI_DIR / "searcher-prompt.md",   "report": "searcher-report.md"},
    "Validator":  {"prompt": AI_DIR / "validator-prompt.md",  "report": "validator-report.md"},
}

ROLE_FILES_APPRAISAL = {
    "Appraiser": {
        "prompt": AI_DIR / "appraisal-prompt.md",
        "report": "appraisal-report.md",
    },
    "Methodologist": {
        "prompt": AI_DIR / "methodologist-prompt.md",
        "report": "methodologist-report.md",
    },
    "Summariser": {
        "prompt": AI_DIR / "summariser-prompt.md",
        "report": "summariser-report.md",
    },
}

ROLE_FILES_SEARCH = {
    "Researcher": {
        "prompt": AI_DIR / "researcher-prompt.md",
        "report": "researcher-report.md",
    },
}

ALL_MODES: dict[str, dict] = {
    "coding":     ROLE_FILES_CODING,
    "writing":    ROLE_FILES_WRITING,
    "rct_search": ROLE_FILES_RCT_SEARCH,
    "appraisal":  ROLE_FILES_APPRAISAL,
    "search":     ROLE_FILES_SEARCH,
}


# ---------------------------------------------------------------------------
# Role-specific documentation injection (least-privilege)
# ---------------------------------------------------------------------------
DOC_FILES_BY_ROLE: dict[str, list[Path]] = {
    "Builder":  [DOCS_CODING / "PRD.md",
                 DOCS_CODING / "architecture.md",
                 DOCS_CODING / "coding-standards.md"],
    "Reviewer": [DOCS_CODING / "PRD.md",
                 DOCS_CODING / "architecture.md",
                 DOCS_CODING / "decision-log.md"],
    "Tester":   [DOCS_CODING / "PRD.md",
                 DOCS_CODING / "architecture.md",
                 DOCS_CODING / "test-strategy.md"],
    "Writer":   [DOCS_WRITING / "project-brief.md",
                 DOCS_WRITING / "style-guide.md"],
    "Editor":   [DOCS_WRITING / "project-brief.md",
                 DOCS_WRITING / "editorial-standards.md"],
    "QA":       [DOCS_WRITING / "project-brief.md",
                 DOCS_WRITING / "qa-checklist.md"],
    "Formulator": [DOCS_RCT_SEARCH / "pico-framework.md"],
    "Searcher":   [DOCS_RCT_SEARCH / "pico-framework.md",
                   DOCS_RCT_SEARCH / "database-guide.md"],
    "Validator":  [DOCS_RCT_SEARCH / "pico-framework.md",
                   DOCS_RCT_SEARCH / "validation-criteria.md"],
    "Appraiser":     [],
    "Methodologist": [],
    "Summariser":    [],
    "Researcher":    [],
}


# ---------------------------------------------------------------------------
# Provider registry
# ---------------------------------------------------------------------------
def call_openai_provider(
    prompt: str,
    model: str | None = None,
    max_tokens: int = 2048,
) -> str:
    """Send a prompt to the OpenAI chat completions endpoint."""
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY is not set. Add it to your .env file.")
    model   = model or OPENAI_MODEL
    url     = "https://api.openai.com/v1/chat/completions"
    payload = json.dumps({
        "model":      model,
        "messages":   [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
    }).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        headers={
            "Content-Type":  "application/json",
            "Authorization": f"Bearer {OPENAI_API_KEY}",
        },
        method="POST",
    )
    try:
        with urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
        choices = data.get("choices", [])
        if not choices or not choices[0].get("message", {}).get("content"):
            raise RuntimeError("OpenAI returned an empty response.")
        return choices[0]["message"]["content"]
    except urllib.error.HTTPError as exc:
        raise RuntimeError(f"OpenAI HTTP error {exc.code}: {exc.reason}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"OpenAI connection error: {exc.reason}") from exc
    except (KeyError, IndexError) as exc:
        raise RuntimeError(f"OpenAI unexpected response format: {exc}") from exc


def call_anthropic_provider(
    prompt: str,
    model: str | None = None,
    max_tokens: int = 2048,
) -> str:
    """Send a prompt to the Anthropic messages endpoint."""
    if not ANTHROPIC_API_KEY:
        raise RuntimeError("ANTHROPIC_API_KEY is not set. Add it to your .env file.")
    model   = model or ANTHROPIC_MODEL
    url     = "https://api.anthropic.com/v1/messages"
    payload = json.dumps({
        "model":      model,
        "messages":   [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
    }).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        headers={
            "Content-Type":      "application/json",
            "x-api-key":         ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
        },
        method="POST",
    )
    try:
        with urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
        content = data.get("content", [])
        if not content or not content[0].get("text"):
            raise RuntimeError("Anthropic returned an empty response.")
        return content[0]["text"]
    except urllib.error.HTTPError as exc:
        raise RuntimeError(f"Anthropic HTTP error {exc.code}: {exc.reason}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Anthropic connection error: {exc.reason}") from exc
    except (KeyError, IndexError) as exc:
        raise RuntimeError(f"Anthropic unexpected response format: {exc}") from exc


def call_ollama_provider(
    prompt: str,
    model: str | None = None,
    max_tokens: int = 2048,
) -> str:
    """Send a prompt to the local Ollama generate endpoint."""
    model   = model or OLLAMA_MODEL
    url     = f"{OLLAMA_HOST}/api/generate"
    payload = json.dumps({
        "model":   model,
        "prompt":  prompt,
        "stream":  False,
        "options": {"num_predict": max_tokens},
    }).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urlopen(req, timeout=300) as resp:
            data = json.loads(resp.read())
        response_text = data.get("response", "")
        if not response_text:
            raise RuntimeError("Ollama returned an empty response.")
        return response_text
    except urllib.error.HTTPError as exc:
        raise RuntimeError(f"Ollama HTTP error {exc.code}: {exc.reason}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Ollama connection error: {exc.reason}") from exc


def call_deepseek_provider(
    prompt: str,
    model: str | None = None,
    max_tokens: int = 2048,
) -> str:
    """Send a prompt to the DeepSeek chat completions endpoint."""
    if not DEEPSEEK_API_KEY:
        raise RuntimeError("DEEPSEEK_API_KEY is not set. Add it to your .env file.")
    model   = model or DEEPSEEK_MODEL
    url     = "https://api.deepseek.com/chat/completions"
    payload = json.dumps({
        "model":      model,
        "messages":   [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "stream":     False,
        "thinking":   {"type": "disabled"},
    }).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        headers={
            "Content-Type":  "application/json",
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        },
        method="POST",
    )
    try:
        with urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
        choices = data.get("choices", [])
        if not choices:
            raise RuntimeError("DeepSeek returned an empty response.")
        msg     = choices[0].get("message", {})
        content = msg.get("content") or msg.get("reasoning_content", "")
        if not content:
            raise RuntimeError("DeepSeek returned an empty response.")
        return content
    except urllib.error.HTTPError as exc:
        raise RuntimeError(f"DeepSeek HTTP error {exc.code}: {exc.reason}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"DeepSeek connection error: {exc.reason}") from exc
    except (KeyError, IndexError) as exc:
        raise RuntimeError(f"DeepSeek unexpected response format: {exc}") from exc


def call_groq_provider(
    prompt: str,
    model: str | None = None,
    max_tokens: int = 2048,
) -> str:
    """Send a prompt to the Groq inference endpoint."""
    if not GROQ_API_KEY:
        raise RuntimeError("GROQ_API_KEY is not set. Add it to your .env file.")
    model   = model or GROQ_MODEL
    url     = "https://api.groq.com/openai/v1/chat/completions"
    payload = json.dumps({
        "model":      model,
        "messages":   [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "stream":     False,
    }).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        headers={
            "Content-Type":  "application/json",
            "Authorization": f"Bearer {GROQ_API_KEY}",
        },
        method="POST",
    )
    try:
        with urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
        choices = data.get("choices", [])
        if not choices or not choices[0].get("message", {}).get("content"):
            raise RuntimeError("Groq returned an empty response.")
        return choices[0]["message"]["content"]
    except urllib.error.HTTPError as exc:
        raise RuntimeError(f"Groq HTTP error {exc.code}: {exc.reason}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Groq connection error: {exc.reason}") from exc
    except (KeyError, IndexError) as exc:
        raise RuntimeError(f"Groq unexpected response format: {exc}") from exc


def call_qwen_provider(
    prompt: str,
    model: str | None = None,
    max_tokens: int = 2048,
) -> str:
    """Send a prompt to Alibaba Cloud Model Studio (Qwen) via OpenAI-compatible API."""
    if not DASHSCOPE_API_KEY:
        raise RuntimeError("DASHSCOPE_API_KEY is not set. Add it to your .env file.")
    model   = model or QWEN_MODEL
    url     = f"{DASHSCOPE_BASE_URL.rstrip('/')}/chat/completions"
    payload = json.dumps({
        "model":           model,
        "messages":        [{"role": "user", "content": prompt}],
        "max_tokens":      max_tokens,
        "stream":          False,
        "enable_thinking": False,
    }).encode()
    req = urllib.request.Request(
        url,
        data=payload,
        headers={
            "Content-Type":  "application/json",
            "Authorization": f"Bearer {DASHSCOPE_API_KEY}",
        },
        method="POST",
    )
    try:
        with urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
        choices = data.get("choices", [])
        if not choices or not choices[0].get("message", {}).get("content"):
            raise RuntimeError("Qwen returned an empty response.")
        return choices[0]["message"]["content"]
    except urllib.error.HTTPError as exc:
        raise RuntimeError(f"Qwen HTTP error {exc.code}: {exc.reason}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Qwen connection error: {exc.reason}") from exc
    except (KeyError, IndexError) as exc:
        raise RuntimeError(f"Qwen unexpected response format: {exc}") from exc


PROVIDERS: dict[str, callable] = {
    "ollama":    call_ollama_provider,
    "openai":    call_openai_provider,
    "anthropic": call_anthropic_provider,
    "deepseek":  call_deepseek_provider,
    "groq":      call_groq_provider,
    "qwen":      call_qwen_provider,
}


def call_ai(
    prompt: str,
    provider: str = "ollama",
    model: str | None = None,
) -> str:
    """Dispatch an AI call to the correct provider function."""
    fn = PROVIDERS.get(provider, call_ollama_provider)
    return fn(prompt, model=model)


# ---------------------------------------------------------------------------
# Context builder (with optional RAG)
# ---------------------------------------------------------------------------
def build_project_context(
    role_name: str,
    query: str = "",
    mode: str = "",
    session_id: str = "",
) -> str:
    """
    Build the project context string injected into every AI prompt.
    Reads small docs/ files for the role, then optionally appends RAG chunks.
    """
    doc_files = DOC_FILES_BY_ROLE.get(role_name, [])
    sections: list[str] = []

    for doc_path in doc_files:
        if doc_path.exists():
            content = doc_path.read_text(encoding="utf-8", errors="replace").strip()
            if content:
                sections.append(f"### {doc_path.name}\n{content}")

    if query and mode and session_id:
        try:
            from src import rag
            rag_context = rag.retrieve(query, mode, session_id)
            if rag_context:
                sections.append(rag_context)
        except Exception as exc:  # noqa: BLE001
            print(f"[RAG] Retrieval warning: {exc}")

    return "\n\n".join(sections)


def truncate_context(text: str, max_chars: int = 2000) -> str:
    """Return text unchanged if within max_chars, otherwise truncate with ellipsis."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\u2026"


# ---------------------------------------------------------------------------
# File utilities
# ---------------------------------------------------------------------------
def read_text_file(path: Path) -> str:
    """Read a text file, strip whitespace. Returns '' if missing."""
    try:
        return path.read_text(encoding="utf-8", errors="replace").strip()
    except (FileNotFoundError, OSError):
        return ""


def save_report(
    path: Path,
    role_name: str,
    model: str,
    task: str,
    response: str,
) -> None:
    """Append a role interaction to a markdown report file."""
    path.parent.mkdir(parents=True, exist_ok=True)
    entry = (
        f"\n## {role_name}\n"
        f"**Model:** {model}\n\n"
        f"**Task:** {task}\n\n"
        f"**Response:**\n{response}\n"
    )
    with path.open("a", encoding="utf-8") as fh:
        fh.write(entry)


def start_session_transcript(reports_dir: Path) -> Path:
    """Create a new timestamped session transcript file. Returns its Path."""
    reports_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path      = reports_dir / f"session_{timestamp}.md"
    path.write_text(
        f"# Session Transcript\nStarted: {timestamp}\n",
        encoding="utf-8",
    )
    return path


def append_to_transcript(
    path: Path,
    role_name: str,
    step: int,
    task: str,
    response: str,
) -> None:
    """Append one interaction step to an existing transcript file."""
    entry = (
        f"\n## Step {step} — {role_name}\n"
        f"**Task:** {task}\n\n"
        f"**Response:**\n{response}\n"
    )
    with path.open("a", encoding="utf-8") as fh:
        fh.write(entry)


def print_session_summary(
    transcript_path: Path,
    step_count: int,
    role_counts: dict[str, int],
) -> None:
    """Print a summary of the completed session to stdout."""
    print(f"\n{'='*55}")
    print(f"  Session complete")
    print(f"  Transcript : {transcript_path.name}")
    print(f"  Total steps: {step_count}")
    if step_count == 0:
        print("  No steps recorded.")
    else:
        print("\n  Role usage:")
        for role, count in role_counts.items():
            if count > 0:
                print(f"    {role}: {count} step(s)")
    print(f"{'='*55}\n")


# ---------------------------------------------------------------------------
# Colour helpers
# ---------------------------------------------------------------------------
COLOURS = {
    "Builder":       "\033[94m",
    "Reviewer":      "\033[93m",
    "Tester":        "\033[92m",
    "Writer":        "\033[95m",
    "Editor":        "\033[96m",
    "QA":            "\033[91m",
    "Formulator":    "\033[94m",
    "Searcher":      "\033[92m",
    "Validator":     "\033[93m",
    "Appraiser":     "\033[95m",
    "Methodologist": "\033[96m",
    "Summariser":    "\033[92m",
    "Researcher":    "\033[94m",
}
RESET = "\033[0m"


def role_color(role_name: str) -> str:
    return COLOURS.get(role_name, RESET)


def choose_role(mode: str = "coding") -> tuple[str, dict]:
    """Prompt the user to pick a role and return (role_name, role_config)."""
    roles      = ALL_MODES[mode]
    role_names = list(roles.keys())

    print(f"\nSelect a role for {mode} mode:")
    for i, name in enumerate(role_names, start=1):
        colour = role_color(name)
        print(f"  {colour}{i}. {name}{RESET}")

    while True:
        try:
            choice = input("Enter number: ").strip()
        except (EOFError, KeyboardInterrupt):
            choice = "1"
            print("[no input detected - defaulting to role 1]")
        if choice.isdigit() and 1 <= int(choice) <= len(role_names):
            name = role_names[int(choice) - 1]
            return name, roles[name]
        print(f"Please enter a number between 1 and {len(role_names)}.")


# ---------------------------------------------------------------------------
# Session management
# ---------------------------------------------------------------------------
def list_sessions(reports_dir: str = str(REPORTS_DIR)) -> None:
    """Print all saved session transcript files, newest first."""
    path = Path(reports_dir)
    if not path.exists():
        print("No reports folder found.")
        return
    files = sorted(path.glob("session_*.md"), reverse=True)
    if not files:
        print("No session transcripts found.")
        return
    print(f"\n{'#':<4} {'Filename':<45} {'Size':>8}")
    print("-" * 60)
    for i, f in enumerate(files, start=1):
        print(f"{i:<4} {f.name:<45} {f.stat().st_size:>6} B")


def read_session(filename: str, reports_dir: str = str(REPORTS_DIR)) -> None:
    """Print the contents of a saved session transcript."""
    path = Path(reports_dir) / filename
    if not path.exists():
        print(f"File not found: {filename}. Use --list-sessions to see available files.")
        return
    print(f"--- {filename} ---")
    print(path.read_text(encoding="utf-8"))


def delete_session(filename: str, reports_dir: str = str(REPORTS_DIR)) -> None:
    """Delete a saved session transcript after confirmation."""
    path = Path(reports_dir) / filename
    if not path.exists():
        print(f"File not found: {filename}. Use --list-sessions to see available files.")
        return
    try:
        confirm = input(f"Delete '{filename}'? [y/N]: ").strip().lower()
    except (EOFError, KeyboardInterrupt):
        print()
        confirm = "n"
    if confirm == "y":
        path.unlink()
        print(f"Deleted: {filename}")
    else:
        print("Cancelled.")


def export_session(filename: str, reports_dir: str = str(REPORTS_DIR)) -> None:
    """Export a session transcript as a plain-text .txt file."""
    src = Path(reports_dir) / filename
    if not src.exists():
        print(f"File not found: {filename}. Use --list-sessions to see available files.")
        return
    raw     = src.read_text(encoding="utf-8")
    cleaned = re.sub(r"#{1,6}\s*", "", raw)
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*(.+?)\*",    r"\1", cleaned)
    txt_name = Path(filename).stem + ".txt"
    dest     = Path(reports_dir) / txt_name
    dest.write_text(cleaned, encoding="utf-8")
    print(f"Exported to: {dest}")


def rename_session(filename: str, reports_dir: str = str(REPORTS_DIR)) -> None:
    """Rename a session transcript file, appending .md automatically."""
    src = Path(reports_dir) / filename
    if not src.exists():
        print(f"File not found: {filename}. Use --list-sessions to see available files.")
        return
    try:
        raw_name = input("New filename (without extension): ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        raw_name = ""
    if not raw_name:
        print("Name cannot be empty. Cancelled.")
        return
    new_name = raw_name if raw_name.endswith(".md") else raw_name + ".md"
    dest = Path(reports_dir) / new_name
    if dest.exists():
        print(f"A file named '{new_name}' already exists. Cancelled.")
        return
    src.rename(dest)
    print(f"Renamed to: {new_name}")


def show_stats(reports_dir: str = str(REPORTS_DIR)) -> None:
    """Print session statistics for all saved transcripts."""
    path = Path(reports_dir)
    if not path.exists():
        print("No reports folder found.")
        return
    files = sorted(path.glob("session_*.md"))
    if not files:
        print("No session transcripts found.")
        return

    total_size  = 0
    role_counts: dict[str, int] = {}
    all_role_names = [r for mode_roles in ALL_MODES.values() for r in mode_roles]
    for role in all_role_names:
        role_counts[role] = 0

    for f in files:
        total_size += f.stat().st_size
        content     = f.read_text(encoding="utf-8", errors="replace")
        for role in all_role_names:
            role_counts[role] += content.count(f"## {role}")

    print(f"\nTotal sessions    : {len(files)}")
    print(f"Total size        : {total_size} bytes")
    print("\nRole usage across all sessions:")
    for role, count in role_counts.items():
        colour = role_color(role)
        print(f"  {colour}{role:<14}{RESET}: {count} interaction(s)")


# ---------------------------------------------------------------------------
# File readers
# ---------------------------------------------------------------------------
def _read_docx(path: Path) -> str:
    """Extract plain text from a .docx file using python-docx."""
    try:
        import docx as _docx
        doc = _docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:  # noqa: BLE001
        return f"[Could not read DOCX: {exc}]"


def _read_pdf_pymupdf(path: Path, max_chars: int = 30_000) -> str:
    """Extract plain text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz
        doc  = fitz.open(str(path))
        text = "\n\n".join(page.get_text() for page in doc)
        doc.close()
        return text[:max_chars]
    except Exception as exc:  # noqa: BLE001
        return f"[Could not read PDF: {exc}]"


def _md_to_docx(md_text: str, title: str, out_path: Path) -> None:
    """Convert a markdown report string to a .docx file using python-docx."""
    import docx as _docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document   = _docx.Document()
    title_para = document.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in md_text.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
        elif stripped.startswith("#### "):
            document.add_heading(stripped[5:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            document.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        elif stripped.startswith("**") and stripped.endswith("**"):
            p   = document.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        else:
            document.add_paragraph(stripped)

    document.save(str(out_path))

def _add_hyperlink(paragraph, text: str, url: str):
    """Add a clickable hyperlink run to an existing paragraph."""
    import docx as _docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    new_t = OxmlElement("w:t")
    new_t.text = text
    new_run.append(new_t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _ranked_articles_to_docx(
    ranked: list[dict],
    title: str,
    out_path: Path,
    topic: str = "",
) -> None:
    """
    Write a ranked article list as a proper Word table with clickable
    PubMed hyperlinks. Each row: Rank | Score | Title | PMID | Link.
    """
    import docx as _docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn

    doc = _docx.Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if topic:
        p = doc.add_paragraph()
        run = p.add_run(f"Topic: {topic}")
        run.italic = True

    doc.add_paragraph("")

    if not ranked:
        doc.add_paragraph("No articles retrieved from PubMed.")
        doc.save(str(out_path))
        return

    # Caption
    cap = doc.add_paragraph()
    cap.add_run(
        f"All {len(ranked)} RCT articles retrieved from PubMed, "
        "ordered by PICO relevance score (10 = most relevant)."
    ).italic = True

    doc.add_paragraph("")

    # Table: Rank | Score | Title | PMID | Link
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for i, label in enumerate(["Rank", "Score", "Title", "PMID", "PubMed Link"]):
        hdr[i].text = label
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True

    # Set column widths
    widths = [Cm(1.2), Cm(1.5), Cm(9.0), Cm(2.2), Cm(2.8)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    # Data rows
    for r in ranked:
        row = table.add_row()
        row.cells[0].text = str(r["rank"])
        row.cells[1].text = f"{r['score']}/10"
        row.cells[2].text = r["title"]
        row.cells[3].text = r["pmid"]
        # Clickable hyperlink in Link cell
        cell_para = row.cells[4].paragraphs[0]
        _add_hyperlink(cell_para, "PubMed", r["url"])

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.add_run(
        "Select your top articles, download PDFs and place them in "
        "input/sr/ to run the SR pipeline."
    ).italic = True

    doc.add_paragraph("")
    ref = doc.add_paragraph()
    ref.add_run(
        "For explanation on ranking, please refer to the full report "
        "in the reports folder."
    ).italic = True

    doc.save(str(out_path))

# ---------------------------------------------------------------------------
# Writing report
# ---------------------------------------------------------------------------

def generate_writing_report(
    docs_dir: Path = DOCS_WRITING,
    reports_dir: Path = REPORTS_DIR,
    provider: str = "ollama",
    model: str | None = None,
    input_dir: Path | None = None,
) -> Path:

    """
    Read input files, send to AI with writing-report prompt, save as
    output/writing/writing_report_{ts}.md and .docx.
    Returns the path of the saved .md report.
    """
    prompt_path = AI_DIR / "writing-report-prompt.md"
    if not prompt_path.exists():
        raise FileNotFoundError(f"Writing report prompt not found: {prompt_path}")

    if input_dir is not None:
        _SUPPORTED = {".txt", ".md", ".pdf", ".docx"}
        files = sorted(
            p for p in input_dir.iterdir()
            if p.is_file() and p.suffix.lower() in _SUPPORTED
        ) if input_dir.exists() else []
    else:
        files = auto_load_input_files("writing")

    if not files and docs_dir.exists():
        SUPPORTED = {".txt", ".md", ".pdf", ".docx"}
        files = sorted(
            f for f in docs_dir.rglob("*")
            if f.is_file() and f.suffix.lower() in SUPPORTED
        )

    if not files:
        print(f"No files found in input/writing/ or {docs_dir}.")
        print("Add .txt, .md, .pdf, or .docx files to: input/writing/")
        return OUTPUT_WRITING / "writing_report_empty.md"

    system_prompt = prompt_path.read_text(encoding="utf-8", errors="replace")

    sections: list[str] = []
    for f in files:
        suffix = f.suffix.lower()
        if suffix == ".pdf":
            content = _read_pdf_pymupdf(f)
        elif suffix == ".docx":
            content = _read_docx(f)
        else:
            content = read_text_file(f)
        if content.strip():
            sections.append(f"### {f.name}\n{content.strip()}")
        else:
            print(f"  Warning: no readable content in {f.name} — skipped.")

    if not sections:
        print("No readable content found in any file. Exiting.")
        return OUTPUT_WRITING / "writing_report_empty.md"

    combined    = "\n\n".join(sections)
    full_prompt = f"{system_prompt}\n\n## Documents to Summarise\n\n{combined}"

    print(f"Generating writing report from {len(sections)} file(s)...")
    try:
        response = call_ai(prompt=full_prompt, provider=provider, model=model)
    except RuntimeError as exc:
        response = f"[ERROR generating report: {exc}]"

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    OUTPUT_WRITING.mkdir(parents=True, exist_ok=True)

    md_path    = OUTPUT_WRITING / f"writing_report_{timestamp}.md"
    md_content = f"# Writing Report\nGenerated: {timestamp}\n\n{response}\n"
    md_path.write_text(md_content, encoding="utf-8")
    print(f"Markdown report saved : output\\writing\\{md_path.name}")

    docx_path = OUTPUT_WRITING / f"writing_report_{timestamp}.docx"
    try:
        _md_to_docx(response, f"Writing Report — {timestamp}", docx_path)
        print(f"Word document saved   : output\\writing\\{docx_path.name}")
    except Exception as exc:  # noqa: BLE001
        print(f"  Warning: could not generate .docx — {exc}")

    return md_path


# ---------------------------------------------------------------------------
# Code revision pipeline (--revise flag)
# ---------------------------------------------------------------------------
CODE_EXTENSIONS = {".py", ".js", ".ts", ".cs", ".java", ".sql", ".html", ".css"}
GUIDANCE_DOCS   = {
    "PRD.md", "architecture.md", "coding-standards.md",
    "decision-log.md", "test-strategy.md",
}


def _read_code_files(docs_coding: Path) -> list[dict]:
    """Return list of {name, content} for code files in docs/coding/."""
    results = []
    if not docs_coding.exists():
        return results
    for f in sorted(docs_coding.iterdir()):
        if (f.is_file()
                and f.suffix.lower() in CODE_EXTENSIONS
                and f.name not in GUIDANCE_DOCS):
            content = f.read_text(encoding="utf-8", errors="replace")
            if content.strip():
                results.append({"name": f.name, "content": content})
    return results


def generate_code_revision(
    start_role: str = "Builder",
    docs_dir: Path = DOCS_CODING,
    reports_dir: Path = REPORTS_DIR,
    provider: str = "ollama",
    model: str | None = None,
    dry_run: bool = False,
) -> Path:
    """
    Single-pass code revision pipeline (legacy --revise flag).
    Builder -> Build -> Review -> Test
    Reviewer -> Review -> Test
    Tester -> Test only
    Saves results as reports/code_revision_{ts}.md and .docx.
    Returns path of the .md report.
    """
    all_stages = ["Builder", "Reviewer", "Tester"]
    start_idx  = all_stages.index(start_role) if start_role in all_stages else 0
    stages     = all_stages[start_idx:]

    code_files = _read_code_files(docs_dir)
    if not code_files:
        print(f"No code files found in {docs_dir}.")
        print(f"Add .py/.js/.ts/.cs/.java/.sql/.html/.css files to: {docs_dir}")
        return reports_dir / "code_revision_empty.md"

    print(f"\n{'='*55}")
    print(f"  CODE REVISION PIPELINE (--revise)")
    print(f"  Starting role : {start_role}")
    print(f"  Pipeline      : {' -> '.join(stages)}")
    print(f"  Code files    : {len(code_files)}")
    print(f"{'='*55}\n")

    code_context = "\n\n".join(
        f"### {cf['name']}\n`\n{cf['content']}\n`"
        for cf in code_files
    )

    guidance_parts = []
    for doc_name in GUIDANCE_DOCS:
        doc_path = docs_dir / doc_name
        if doc_path.exists():
            guidance_parts.append(
                f"### {doc_name}\n"
                f"{doc_path.read_text(encoding='utf-8', errors='replace').strip()}"
            )
    guidance_context = "\n\n".join(guidance_parts)

    try:
        task = input("Describe the revision task (or press Enter for general review): ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        task = ""
    if not task:
        task = "Review and improve the code quality, readability, and correctness."

    reports_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    md_path   = reports_dir / f"code_revision_{timestamp}.md"
    docx_path = reports_dir / f"code_revision_{timestamp}.docx"

    full_report_parts = [
        "# Code Revision Report",
        f"Generated: {timestamp}",
        f"Pipeline: {' -> '.join(stages)}",
        f"Task: {task}",
        "",
    ]

    previous_response = ""

    for stage in stages:
        print(f"Running {stage}...")
        role_cfg    = ALL_MODES["coding"][stage]
        prompt_path = Path(role_cfg["prompt"])
        role_prompt = prompt_path.read_text(encoding="utf-8", errors="replace")

        parts = [role_prompt]
        if guidance_context:
            parts.append(f"## Project Guidance\n{guidance_context}")
        parts.append(f"## Code Files\n{code_context}")
        if previous_response:
            parts.append(f"## Previous Stage Output\n{previous_response}")
        parts.append(f"## Revision Task\n{task}")
        full_prompt = "\n\n".join(parts)

        if dry_run:
            response = f"[DRY RUN] {stage} would respond here."
        else:
            try:
                response = call_ai(prompt=full_prompt, provider=provider, model=model)
            except RuntimeError as exc:
                response = f"[ERROR in {stage}: {exc}]"

        previous_response = response
        full_report_parts.append(f"## {stage} Output\n\n{response}\n")
        print(f"{stage} complete.\n")

    md_content = "\n".join(full_report_parts)
    md_path.write_text(md_content, encoding="utf-8")
    print(f"Markdown report saved : reports\\{md_path.name}")

    try:
        _md_to_docx(md_content, f"Code Revision Report — {timestamp}", docx_path)
        print(f"Word document saved   : reports\\{docx_path.name}")
    except Exception as exc:  # noqa: BLE001
        print(f"  Warning: could not generate .docx — {exc}")

    return md_path


# ---------------------------------------------------------------------------
# Article helpers (appraisal mode)
# ---------------------------------------------------------------------------
ARTICLE_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
ARTICLE_SIZE_LIMIT = 8000


def _read_article_files(input_appraisal: Path) -> list[dict]:
    """Read appraisal article files from input/appraisal/."""
    results = []
    if not input_appraisal.exists():
        return results
    for f in sorted(input_appraisal.iterdir()):
        if f.suffix.lower() not in ARTICLE_EXTENSIONS:
            continue
        try:
            if f.suffix.lower() == ".pdf":
                text = _read_pdf_pymupdf(f)
            elif f.suffix.lower() == ".docx":
                text = _read_docx(f)
            else:
                text = f.read_text(encoding="utf-8", errors="replace")
            if not text.strip():
                continue
            if len(text) <= ARTICLE_SIZE_LIMIT:
                results.append({"name": f.name, "content": text})
            else:
                print(f"[Appraisal] {f.name} exceeds {ARTICLE_SIZE_LIMIT} chars — passed to RAG.")
        except Exception as exc:  # noqa: BLE001
            print(f"[Appraisal] Could not read {f.name}: {exc}")
    return results


def _read_topic_file(topic_path: Path) -> str:
    if not topic_path.exists():
        return ""
    return topic_path.read_text(encoding="utf-8", errors="replace").strip()


# ---------------------------------------------------------------------------
# RCT Search pipeline
# ---------------------------------------------------------------------------
def run_rct_search_pipeline(
    provider: str = "ollama",
    model: str | None = None,
    reports_dir: Path = REPORTS_DIR,
    dry_run: bool = False,
) -> Path:
    """
    Single-pass RCT search pipeline:
      1. Formulator  — structures user topic into PICO question
      2. Searcher    — builds Boolean search strategy for all 7 databases
      3. Validator   — validates alignment and approves or requests refinement
    Saves output as output/rct_search/rct_search_{ts}.md and .docx.
    Returns path of the .md report.
    """
    print("\n" + "=" * 55)
    print("  RCT SEARCH PIPELINE")
    print("=" * 55)
    print("  This pipeline will:")
    print("  1. Structure your topic into a PICO question")
    print("  2. Build a search strategy for all 7 SR databases")
    print("  3. Validate the strategy before download")
    print("  No article appraisal — use --mode appraisal for that.")
    print("=" * 55 + "\n")

    topic_file = DOCS_RCT_SEARCH / "topic.md"
    file_topic = _read_topic_file(topic_file)
    if file_topic:
        topic = file_topic
        print(f"[RCT Search] Topic loaded from docs/rct_search/topic.md: {topic}")
    else:
        try:
            topic = input("Enter your research topic: ").strip()
        except (EOFError, KeyboardInterrupt):
            print()
            topic = ""
    if not topic:
        print("No topic entered. Exiting.")
        sys.exit(0)

    # ── PICO input: offer to import saved JSON or prompt manually ──────────────

    _pico_input_dir = INPUT_DIR / "rct_search"
    pico_json_files = sorted(
        _pico_input_dir.glob("pico_*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    ) if _pico_input_dir.exists() else []
    if not pico_json_files and OUTPUT_RCT_SEARCH.exists():
        pico_json_files = sorted(
            OUTPUT_RCT_SEARCH.glob("pico_*.json"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )
    _pico_source = "input/rct_search" if (_pico_input_dir.exists() and list(_pico_input_dir.glob("pico_*.json"))) else "output/rct_search"

    pico_population   = ""
    pico_intervention = ""
    pico_comparator   = ""
    pico_outcome      = ""

    if pico_json_files and not dry_run:
        print()
        print(f"[RCT Search] Saved PICO files found in {_pico_source}:")

        for idx, pf in enumerate(pico_json_files[:5], 1):
            print(f"  {idx}. {pf.name}")
        print("  0. Enter new PICO manually")
        print()
        try:
            choice = input("Import a saved PICO? Enter number or 0 for new: ").strip()
        except (EOFError, KeyboardInterrupt):
            choice = "0"
        if choice.isdigit() and 1 <= int(choice) <= len(pico_json_files[:5]):
            chosen = pico_json_files[int(choice) - 1]
            try:
                import json as _json
                pico_data = _json.loads(chosen.read_text(encoding="utf-8"))
                pico_population   = pico_data.get("population", "")
                pico_intervention = pico_data.get("intervention", "")
                pico_comparator   = pico_data.get("comparator", "")
                pico_outcome      = pico_data.get("outcome", "")
                if not topic:
                    topic = pico_data.get("topic", topic)
                print(f"[RCT Search] PICO imported from {chosen.name}")
            except Exception as _e:
                print(f"[RCT Search] Could not load PICO JSON: {_e}")

    if not any([pico_population, pico_intervention, pico_comparator, pico_outcome]) and not dry_run:
        print()
        print("[RCT Search] Enter PICO components (press Enter to skip any field):")
        try:
            pico_population   = input("  Population   (P): ").strip()
            pico_intervention = input("  Intervention (I): ").strip()
            pico_comparator   = input("  Comparator   (C): ").strip()
            pico_outcome      = input("  Outcome      (O): ").strip()
        except (EOFError, KeyboardInterrupt):
            print()

    pico_manual_context = ""
    if any([pico_population, pico_intervention, pico_comparator, pico_outcome]):
        pico_manual_context = (
            "\n\nUser-provided PICO components:\n"
            f"  Population:   {pico_population or '(not specified)'}\n"
            f"  Intervention: {pico_intervention or '(not specified)'}\n"
            f"  Comparator:   {pico_comparator or '(not specified)'}\n"
            f"  Outcome:      {pico_outcome or '(not specified)'}"
        )
    # ── end PICO input ─────────────────────────────────────────────────────────

    pico_path         = DOCS_RCT_SEARCH / "pico-framework.md"
    db_guide          = DOCS_RCT_SEARCH / "database-guide.md"
    val_criteria      = DOCS_RCT_SEARCH / "validation-criteria.md"
    pico_context      = pico_path.read_text(encoding="utf-8", errors="replace")     if pico_path.exists()    else ""
    db_guide_text     = db_guide.read_text(encoding="utf-8", errors="replace")      if db_guide.exists()     else ""
    val_criteria_text = val_criteria.read_text(encoding="utf-8", errors="replace")  if val_criteria.exists() else ""

    stages = [
        {
            "role":          "Formulator",
            "prompt_file":   AI_DIR / "formulator-prompt.md",
            "extra_context": pico_context,
            "task":          f"The user's research topic is: {topic}{pico_manual_context}\n\nStructure this into a formal PICO question.",
        },
        {
            "role":          "Searcher",
            "prompt_file":   AI_DIR / "searcher-prompt.md",
            "extra_context": db_guide_text,
            "task":          "Build a comprehensive Boolean search strategy for all 7 SR databases based on the PICO question above.",
        },
        {
            "role":          "Validator",
            "prompt_file":   AI_DIR / "validator-prompt.md",
            "extra_context": val_criteria_text,
            "task":          "Validate the search strategy above. Check PICO alignment, database coverage, syntax, and RCT filters. Return APPROVED FOR DOWNLOAD or REQUIRES REFINEMENT with specific justification.",
        },
    ]

    report_parts = [
        "# RCT Search Strategy Report",
        f"Topic: {topic}",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "---",
        "",
        "> **Note:** This report contains the validated search strategy only.",
        "> For article appraisal, copy URLs into --mode appraisal.",
        "",
    ]

    previous_response = ""

    for stage in stages:
        role   = stage["role"]
        colour = role_color(role)
        print(f"Running {role}...")

        prompt_file = Path(stage["prompt_file"])
        if not prompt_file.exists():
            print(f"  Prompt not found: {prompt_file} — skipping {role}.")
            continue

        role_prompt = prompt_file.read_text(encoding="utf-8", errors="replace")
        parts = [role_prompt]
        if pico_context:
            parts.append(f"## PICO Framework Reference\n{pico_context}")
        if stage["extra_context"]:
            parts.append(f"## Reference Document\n{stage['extra_context']}")
        if previous_response:
            parts.append(f"## Previous Stage Output\n{previous_response}")
        parts.append(f"## Task\n{stage['task']}")
        full_prompt = "\n\n".join(parts)

        if dry_run:
            response = f"[DRY RUN] {role} would respond here."
        else:
            try:
                response = call_ai(prompt=full_prompt, provider=provider, model=model)
            except RuntimeError as exc:
                response = f"[ERROR in {role}: {exc}]"

        previous_response = response
        report_parts.append(f"## {role} Output\n\n{response}\n")
        print(f"{colour}{role} complete.{RESET}\n")

        # -- save PICO JSON after Formulator completes -------------------------
        if role == "Formulator" and not dry_run:
            pico_json_dir = OUTPUT_RCT_SEARCH
            pico_json_dir.mkdir(parents=True, exist_ok=True)
            pico_json_path = pico_json_dir / f"pico_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            pico_data = {
                "timestamp":         datetime.now().strftime("%Y%m%d_%H%M%S"),
                "topic":             topic,
                "formulator_output": response,
                "population":        "",
                "intervention":      "",
                "comparator":        "",
                "outcome":           "",
                "study_design":      "RCT",
                "effect_measure":    "SMD",
                "source_mode":       "rct_search",
            }
            import re as _re
            def _extract_pico_field(text, labels):
                for lbl in labels:
                    m = _re.search(r"(?i)" + lbl + r"[^:]*:\s*(.+)", text)
                    if m:
                        return m.group(1).strip()
                return ""
            pico_data["population"]   = pico_population or _extract_pico_field(response, [r"P\s*\(Population\)", r"Population"])
            pico_data["intervention"] = pico_intervention or _extract_pico_field(response, [r"I\s*\(Intervention\)", r"Intervention"])
            pico_data["comparator"]   = pico_comparator or _extract_pico_field(response, [r"C\s*\(Comp\w+\)", r"Compar"])
            pico_data["outcome"]      = pico_outcome or _extract_pico_field(response, [r"O\s*\(Outcome\)", r"Outcome"])
            pico_json_path.write_text(
                json.dumps(pico_data, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
            print(f"[RCT Search] PICO saved to: {pico_json_path.name}")

        # ----------------------------------------------------------------------


    validator_output = previous_response.upper()
    if "APPROVED FOR DOWNLOAD" in validator_output:
        status = "APPROVED FOR DOWNLOAD"
        print("\n✅ Search strategy APPROVED FOR DOWNLOAD.")
    else:
        status = "REQUIRES REFINEMENT"
        print("\n⚠️  Search strategy REQUIRES REFINEMENT — see Validator output.")

    report_parts.append(f"## Final Status\n\n**{status}**\n")
    report_parts.append(
        "## Next Steps\n\n"
        "- If APPROVED: copy database search strings into each SR database platform\n"
        "- Download article lists from each database\n"
        "- Run python src/main.py --mode appraisal to appraise individual articles\n"
        "- Run python src/main.py --mode sr for full systematic review pipeline\n"
    )

    # -- Stage 4: PubMed Fetch + AI Ranking ----------------------------------
    import re as _re2
    def _clean_pico_term(t: str) -> str:
        t = _re2.sub(r"\([^)]*\)", "", t)           # remove parenthetical qualifiers e.g. (>=18 years)
        t = _re2.sub(r"\b(at\s+)?(post[\s\-]intervention|post[\s\-]treatment|"
                     r"follow[\s\-]up|immediately after|baseline|"
                     r"at\s+\d+\s+weeks?|at\s+\d+\s+months?)\b", "", t, flags=_re2.IGNORECASE)
        t = _re2.sub(r"[^\w\s\-]", " ", t)          # remove special chars except hyphen
        t = _re2.sub(r"\s+", " ", t).strip()         # collapse whitespace
        # keep only first 4 words to avoid over-specific phrases
        words = t.split()
        return " ".join(words[:4]) if len(words) > 4 else t
    _pico_terms = [_clean_pico_term(t) for t in [pico_population, pico_intervention, pico_outcome] if t]
    _pico_terms = [t for t in _pico_terms if t]     # drop any that became empty after cleaning
    _pubmed_query = " AND ".join(_pico_terms) if _pico_terms else topic
    _pubmed_query = _pubmed_query + " AND randomized controlled trial[pt]"
    print(f"[RCT Search] PubMed query (cleaned): {_pubmed_query}")
    print("[RCT Search] Searching PubMed for relevant articles...")
    articles = fetch_pubmed_articles(_pubmed_query, max_results=20) if not dry_run else [

        {"pmid": "00000001", "title": "[DRY RUN] Test Article",
         "abstract": "Dry run abstract.",
         "url": "https://pubmed.ncbi.nlm.nih.gov/00000001/"}
    ]

    if articles:
        print(f"[RCT Search] Found {len(articles)} article(s). Ranking by PICO relevance...")
        abstracts_block = ""
        for idx, art in enumerate(articles, 1):
            abstracts_block += (
                "\n### Article " + str(idx) + "\n"
                + "Title: " + art["title"] + "\n"
                + "PMID: " + art["pmid"] + "\n"
                + "URL: " + art["url"] + "\n"
                + "Abstract: " + art["abstract"] + "\n"
            )
        rank_prompt = (
            "You are a systematic review expert. Your ONLY task is to rate articles.\n\n"
            "PICO Question:\n" + previous_response + "\n\n"
            "INSTRUCTIONS:\n"
            "- Rate each article for relevance to the PICO question on a scale of 1-10.\n"
            "- 10 = highly relevant RCT directly matching the PICO.\n"
            "- 1 = not relevant.\n"
            "- Output ONLY the rating lines. No summaries, no headers, no extra text.\n"
            "- Each line MUST follow this EXACT format:\n"
            "ARTICLE_RANK: 1 | SCORE: 9 | PMID: 12345678 | TITLE: Example title here | URL: https://pubmed.ncbi.nlm.nih.gov/12345678/\n\n"
            "EXAMPLE OUTPUT (do not copy — use real values):\n"
            "ARTICLE_RANK: 1 | SCORE: 9 | PMID: 11111111 | TITLE: CBT for fibromyalgia RCT | URL: https://pubmed.ncbi.nlm.nih.gov/11111111/\n"
            "ARTICLE_RANK: 2 | SCORE: 7 | PMID: 22222222 | TITLE: Mindfulness for chronic pain | URL: https://pubmed.ncbi.nlm.nih.gov/22222222/\n"
            "ARTICLE_RANK: 3 | SCORE: 3 | PMID: 33333333 | TITLE: Ketamine infusion study | URL: https://pubmed.ncbi.nlm.nih.gov/33333333/\n\n"
            "NOW RATE THESE ARTICLES — output rating lines ONLY:\n\n"
            + abstracts_block
        )
        if dry_run:
            rank_response = "\n".join(
                "ARTICLE_RANK: " + str(i) + " | SCORE: " + str(10 - i) + " | "
                + "PMID: " + art["pmid"] + " | TITLE: " + art["title"]
                + " | URL: " + art["url"]
                for i, art in enumerate(articles, 1)
            )
        else:
            try:
                rank_response = call_ai(
                    prompt=rank_prompt, provider=provider, model=model
                )
            except RuntimeError as exc:
                rank_response = "[ERROR ranking articles: " + str(exc) + "]"
        ranked = []
        if rank_response.startswith("[ERROR"):
            print(f"[RCT Search] Ranking error: {rank_response}")
        for ln in rank_response.splitlines():
            m = re.match(
                r"ARTICLE_RANK:\s*(\d+)\s*\|\s*SCORE:\s*(\d+)\s*\|"
                r"\s*PMID:\s*(\S+)\s*\|\s*TITLE:\s*(.+?)\s*\|\s*URL:\s*(\S+)",
                ln.strip()
            )
            if m:
                ranked.append({
                    "rank":  int(m.group(1)),
                    "score": int(m.group(2)),
                    "pmid":  m.group(3),
                    "title": m.group(4),
                    "url":   m.group(5),
                })
        ranked.sort(key=lambda x: x["score"], reverse=True)
        for new_rank, r in enumerate(ranked, 1):
            r["rank"] = new_rank
        if not ranked:
            print("[RCT Search] Warning: AI ranking returned no parseable results.")
            print("[RCT Search] Raw ranking response (first 500 chars):")
            print(rank_response[:500])
            # fallback: use fetch order, score descending by position
            ranked = [
                {"rank": i, "score": len(articles) - i + 1,
                 "pmid": a["pmid"], "title": a["title"], "url": a["url"]}
                for i, a in enumerate(articles, 1)
    ]
        table_lines = [
            "## Ranked Article List\n",
            f"_All {len(ranked)} articles retrieved from PubMed, ordered by relevance score (1 = most relevant)._\n",
            "| Rank | Score | Title | PMID | Link |",
            "|------|-------|-------|------|------|",
        ]
        for r in ranked:
            table_lines.append(
                "| " + str(r["rank"]) + " | " + str(r["score"]) + "/10 | "
                + r["title"] + " | " + r["pmid"]
                + " | [PubMed](" + r["url"] + ") |"
            )
        table_lines.append(
            "\n> Select your top articles, download PDFs "
            "and place them in input/sr/ to run the SR pipeline."
        )
        
        report_parts.append("\n".join(table_lines))
        top = ranked[0]["title"] if ranked else "N/A"
        print("[RCT Search] Ranking complete. Top article: " + top)
        pico_files = sorted(OUTPUT_RCT_SEARCH.glob("pico_*.json"), reverse=True)
        if pico_files:
            pico_data = json.loads(pico_files[0].read_text(encoding="utf-8"))
            pico_data["ranked_articles"] = ranked
            pico_files[0].write_text(
                json.dumps(pico_data, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
            print("[RCT Search] PICO JSON updated with "
                  + str(len(ranked)) + " ranked article(s).")
    else:
        print("[RCT Search] No articles found on PubMed for this topic.")
        report_parts.append(
            "## Ranked Article List\n\n_No articles retrieved from PubMed._\n"
        )

    # -- End Stage 4 ----------------------------------------------------------

    timestamp  = datetime.now().strftime("%Y%m%d_%H%M%S")
    if reports_dir != REPORTS_DIR:
        out_dir   = reports_dir
        final_dir = reports_dir
    else:
        out_dir   = REPORTS_DIR / "rct_search"
        final_dir = OUTPUT_RCT_SEARCH
    out_dir.mkdir(parents=True, exist_ok=True)
    final_dir.mkdir(parents=True, exist_ok=True)
    md_path    = out_dir / f"rct_search_{timestamp}.md"
    docx_path  = out_dir / f"rct_search_{timestamp}.docx"

    md_content = "\n".join(report_parts)
    md_path.write_text(md_content, encoding="utf-8")
    print(f"Intermediate report   : reports/rct_search/{md_path.name}")

    final_parts = [p for p in report_parts if "Ranked Article List" in p or "Final Status" in p]
    if final_parts and final_dir != out_dir:
        final_md_path = final_dir / f"rct_search_{timestamp}.md"
        final_md_content = (
            "\n".join(final_parts)
            + "\n\n---\n\n> For explanation on ranking, please refer to the full report in the reports folder."
        )
        final_md_path.write_text(final_md_content, encoding="utf-8")
        print(f"Final report saved    : output/rct_search/{final_md_path.name}")
        final_docx_path = final_dir / f"rct_search_{timestamp}.docx"
        try:
            _ranked_articles_to_docx(ranked, f"RCT Search Results — {timestamp}", final_docx_path, topic=topic)
            print(f"Results DOCX saved    : output/rct_search/{final_docx_path.name}")
        except Exception as exc:  # noqa: BLE001
            print(f"  Warning: could not generate results .docx — {exc}")        

    try:
        _md_to_docx(md_content, f"RCT Search Strategy — {timestamp}", docx_path)
        print(f"Word document saved   : {out_dir.name}\\{docx_path.name}")
    except Exception as exc:  # noqa: BLE001
        print(f"  Warning: could not generate .docx — {exc}")

    # -- Auto-copy PICO JSON to input/sr/ -------------------------------------
    if not dry_run:
        pico_files = sorted(OUTPUT_RCT_SEARCH.glob("pico_*.json"), reverse=True)
        if pico_files:
            latest_pico = pico_files[0]
            print()
            print(f"[RCT Search] Latest PICO JSON: {latest_pico.name}")
            try:
                copy_choice = input(
                    "[RCT Search] Copy pico_*.json to input/sr/ for SR pipeline? [Y/N]: "
                ).strip().upper()
            except (EOFError, KeyboardInterrupt):
                copy_choice = "N"
            if copy_choice == "Y":
                import shutil as _shutil
                INPUT_SR.mkdir(parents=True, exist_ok=True)
                dest = INPUT_SR / latest_pico.name
                _shutil.copy2(str(latest_pico), str(dest))
                print(f"[RCT Search] PICO JSON copied to input/sr/{latest_pico.name}")
            else:
                print("[RCT Search] PICO JSON not copied — you can copy it manually later.")
    # -------------------------------------------------------------------------

    return md_path
    

def rct_search_reminder() -> None:
    """Print a reminder to edit the PICO framework before starting a search."""
    print("\n" + "=" * 55)
    print("  RCT SEARCH MODE")
    print("=" * 55)
    print("  IMPORTANT: Before proceeding, ensure you have edited")
    print("  your PICO framework file:")
    print(f"  {DOCS_RCT_SEARCH / 'pico-framework.md'}")
    print()
    print("  Population  : who are the patients?")
    print("  Intervention: what is being tested?")
    print("  Comparison  : what is the control/comparator?")
    print("  Outcome     : what are you measuring?")
    print()
    print("  Search links will be saved to the output/rct_search/ folder.")
    print("=" * 55 + "\n")
    try:
        confirm = input("Have you edited your PICO file? [y/N]: ").strip().lower()
    except (EOFError, KeyboardInterrupt):
        print()
        confirm = "n"
    if confirm != "y":
        print("Please edit the PICO file first, then re-run rct_search mode.")
        sys.exit(0)


def save_rct_search_links(
    response: str,
    reports_dir: Path = REPORTS_DIR,
) -> Path:
    """
    Extract URLs from an AI response and save them as a markdown link list.
    Returns the path of the saved file.
    """
    urls = re.findall(r'https?://\S+', response)
    urls = [u.rstrip(".,);\"'") for u in urls]

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path  = reports_dir / f"rct_search_{timestamp}.md"
    reports_dir.mkdir(parents=True, exist_ok=True)

    lines = [
        "# RCT Search Links",
        f"Generated: {timestamp}",
        "",
        "Download this file from the `reports/` folder.",
        "",
        "## Search Result Links",
        "",
    ]
    if urls:
        for url in urls:
            lines.append(f"- [{url}]({url})")
    else:
        lines.append("_No URLs found in the AI response._")

    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"\n[RCT Search] Links saved to: {out_path.name}")
    print(f"[RCT Search] Open reports\\{out_path.name} to download your links.\n")
    return out_path


# ---------------------------------------------------------------------------
# PubMed fetch
# ---------------------------------------------------------------------------
def fetch_pubmed_articles(
    query: str,
    max_results: int = 10,
) -> list[dict]:
    """
    Search PubMed via NCBI E-utilities (no API key required).
    Returns list of dicts: pmid, title, abstract, url.
    Returns empty list on any error.
    """
    base       = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    search_url = (
        f"{base}/esearch.fcgi?db=pubmed&term={urllib.parse.quote(query)}"
        f"&retmax={max_results}&retmode=json"
    )
    try:
        with urlopen(search_url, timeout=15) as resp:
            search_data = json.loads(resp.read())
        pmids = search_data["esearchresult"]["idlist"]
    except Exception:  # noqa: BLE001
        return []

    if not pmids:
        return []

    ids_str   = ",".join(pmids)
    fetch_url = (
        f"{base}/efetch.fcgi?db=pubmed&id={ids_str}"
        f"&rettype=abstract&retmode=xml"
    )
    try:
        with urlopen(fetch_url, timeout=15) as resp:
            xml_data = resp.read().decode("utf-8", errors="replace")
    except Exception:  # noqa: BLE001
        return []

    import xml.etree.ElementTree as ET
    try:
        root = ET.fromstring(xml_data)
    except ET.ParseError:
        return []

    articles = []
    for article in root.findall(".//PubmedArticle"):
        pmid_el     = article.find(".//PMID")
        title_el    = article.find(".//ArticleTitle")
        abstract_el = article.find(".//AbstractText")

        pmid     = pmid_el.text.strip()     if pmid_el     is not None else "unknown"
        title    = title_el.text.strip()    if title_el    is not None else "No title"
        abstract = abstract_el.text.strip() if abstract_el is not None else "No abstract available."

        title    = re.sub(r"\s+", " ", title)
        abstract = re.sub(r"\s+", " ", abstract)

        articles.append({
            "pmid":     pmid,
            "title":    title,
            "abstract": abstract,
            "url":      f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
        })

    return articles


# ---------------------------------------------------------------------------
# SR launcher
# ---------------------------------------------------------------------------
def run_sr_launcher() -> None:
    """Launch the unified Pipeline UI (SR mode uses src/ui/app.py)."""
    import subprocess as _sp
    unified_ui = BASE_DIR / "src" / "ui" / "app.py"
    repo_root  = BASE_DIR

    print("\n" + "=" * 58)
    print("  SR Automation Pipeline")
    print("  PRISMA 2020  |  Cochrane Handbook v6.5")
    print("=" * 58)
    print("\n  Launching Pipeline UI (SR mode) in a new window...")
    print("  URL: http://localhost:8501")
    print("  Navigate to Systematic Review in the sidebar.")
    print("  Close the browser window to stop the server.\n")

    _launch_kw: dict = {"cwd": str(repo_root)}
    if os.name == "nt":
        import subprocess as _sub
        _launch_kw["creationflags"] = _sub.CREATE_NEW_CONSOLE
    _sp.Popen(
        [sys.executable, "-m", "streamlit", "run", str(unified_ui),
         "--server.runOnSave", "false"],
        **_launch_kw,
    )

    print("  Pipeline UI launched. Returning to menu...\n")

def list_roles(mode: str = "coding") -> None:
    """Print each role's prompt file and injected documentation for a mode."""
    roles = ALL_MODES.get(mode, {})
    print(f"\n{'='*55}")
    print(f"  Roles for mode: {mode}")
    print(f"{'='*55}")
    for role_name, role_cfg in roles.items():
        colour    = role_color(role_name)
        prompt    = Path(role_cfg["prompt"]).relative_to(BASE_DIR)
        doc_files = DOC_FILES_BY_ROLE.get(role_name, [])
        doc_names = [d.name for d in doc_files]
        print(f"\n{colour}{role_name}{RESET}")
        print(f"  Prompt : {prompt}")
        print(f"  Docs   : {', '.join(doc_names) if doc_names else 'none'}")
    print(f"\n{'='*55}\n")

# ---------------------------------------------------------------------------
# Coding mode dispatcher  (Builder pipeline / Reviewer / Tester standalone)
# ---------------------------------------------------------------------------
def handle_coding_mode(
    provider: str = "ollama",
    model: str | None = None,
    dry_run: bool = False,
) -> None:
    """
    Terminal entry point for Coding mode.
    Presents Builder / Reviewer / Tester sub-mode menu, collects > direct
    instructions from the user, then delegates to the coding.py engine.
    """

    def _call_llm_fn(system_prompt: str, user_prompt: str) -> str:
        if dry_run:
            return "[DRY RUN] LLM would respond here."
        combined = (
            f"## System Instructions\n{system_prompt}\n\n"
            f"## User Request\n{user_prompt}"
        )
        return call_ai(prompt=combined, provider=provider, model=model)

    print("\n" + "=" * 60)
    print("  CODING MODE")
    print("=" * 60)
    print(f"  {COLOURS['Builder']}1. Builder {RESET} "
          f"— pipeline: build → review → test → output/coding/")
    print(f"  {COLOURS['Reviewer']}2. Reviewer{RESET} "
          f"— standalone: review code in input/coding/")
    print(f"  {COLOURS['Tester']}3. Tester  {RESET} "
          f"— standalone: test code in input/coding/")
    print("  0. Back to menu")
    print("=" * 60)

    try:
        choice = input("Select sub-mode [0-3]: ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        choice = "0"

    if choice == "0" or not choice:
        print("Returning to menu.")
        return

    sub_mode_map = {"1": "Builder", "2": "Reviewer", "3": "Tester"}
    if choice not in sub_mode_map:
        print("Invalid choice — returning to menu.")
        return

    sub_mode = sub_mode_map[choice]
    colour   = COLOURS.get(sub_mode, RESET)

    print(f"\n{colour}[{sub_mode.upper()}]{RESET} Enter instructions below.")
    print("  Lines starting with > are DIRECT TASK INSTRUCTIONS (highest priority).")
    print("  Press ENTER on a blank line when done.\n")

    lines: list[str] = []
    while True:
        try:
            line = input("  instruction> ")
        except (EOFError, KeyboardInterrupt):
            break
        if line.strip() == "":
            print("\n  ✓ Instructions received — starting processing...\n")
            break
        if not line.strip().startswith(">"):
            line = "> " + line.strip()
        lines.append(line)

    raw_input_text      = "\n".join(lines)
    direct_instructions = parse_direct_instructions(raw_input_text)

    if direct_instructions:
        print(f"\n{colour}[{sub_mode.upper()}]{RESET} "
              f"Direct instructions captured ({len(direct_instructions)}):")
        for instr in direct_instructions:
            print(f"  → {instr}")
    else:
        print(f"\n{colour}[{sub_mode.upper()}]{RESET} "
              "No direct instructions found — using docs/coding/ guidelines only.")

    auto_load_input_files("coding")
    print()

    if sub_mode == "Builder":
        run_builder(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            verbose=True,
        )
    elif sub_mode == "Reviewer":
        run_reviewer(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            verbose=True,
        )
    elif sub_mode == "Tester":
        run_tester(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            verbose=True,
        )

    print(f"\n{colour}[{sub_mode.upper()}]{RESET} Done. "
          "Check output/coding/ and reports/coding/ for results.\n")

# ---------------------------------------------------------------------------
# Appraisal mode dispatcher
# ---------------------------------------------------------------------------
def handle_appraisal_mode(
    provider: str = "ollama",
    model: str | None = None,
    dry_run: bool = False,
) -> None:
    """Terminal entry point for Appraisal mode."""
    from src.modes.appraisal import run_appraisal, parse_direct_instructions as _pdi

    def _call_llm_fn(system_prompt: str, user_prompt: str) -> str:
        if dry_run:
            return "[DRY RUN] LLM would respond here."
        return call_ai(
            prompt=f"## System Instructions\n{system_prompt}\n\n## User Request\n{user_prompt}",
            provider=provider,
            model=model,
        )

    print("\n" + "=" * 60)
    print("  APPRAISAL MODE")
    print("=" * 60)
    print("  Place article(s) in input/appraisal/ before starting.")
    print("  Supported formats: PDF, DOCX, MD, TXT")
    print("  Output  → output/appraisal/")
    print("  Reports → reports/appraisal/")
    print("=" * 60)

    lines: list[str] = []
    print("\n  Enter optional direct instructions below.")
    print("  Lines starting with > are highest priority.")
    print("  Press ENTER on a blank line to start.\n")
    while True:
        line = input("  instruction> ").strip()
        if not line:
            print("\n  Instructions received -- starting appraisal...\n")
            break
        if not line.startswith(">"):
            line = "> " + line
        lines.append(line)

    direct = _pdi("\n".join(lines))
    run_appraisal(direct, _call_llm_fn, verbose=True)


# ---------------------------------------------------------------------------
# Writing mode dispatcher  (Writer pipeline / Editor / QA standalone)
# ---------------------------------------------------------------------------

def handle_search_mode(provider: str, model: str) -> None:
    model = model or OLLAMA_MODEL
    """Dispatcher for Search mode (Topic Search and Article Search)."""
    from src.modes.search import run_topic_search, run_article_search

    print("\n=== SEARCH MODE ===")
    print("  1. Topic Search  (web synopsis with reference links)")
    print("  2. Article Search (PubMed by article type + comparison)")
    sub = input("  Select sub-mode: ").strip()

    if sub not in ("1", "2"):
        print("  Invalid selection. Returning to menu.")
        return

    # Collect direct instructions  (same pattern as Coding / Writing modes)
    print("  Enter your search query below.")
    print("  Lines starting with > are instructions (auto-prefixed if omitted).")
    print("  Press ENTER on a blank line to start.\n")
    raw_lines: list[str] = []
    while True:
        try:
            line = input("  > ")
        except (EOFError, KeyboardInterrupt):
            break
        if line.strip() == "":
            print("\n  Instructions received — starting search...\n")
            break
        if not line.strip().startswith(">"):
            line = "> " + line.strip()
        raw_lines.append(line)
    direct_instructions: list[str] = [
        l[1:].strip() for l in raw_lines if l.strip().startswith(">") and l[1:].strip()
    ]

    def call_llm_fn(system_prompt: str = "", user_prompt: str = "") -> str:
        combined = f"{system_prompt}\n\n{user_prompt}" if system_prompt else user_prompt
        return call_ai(
            prompt=combined,
            provider=provider,
            model=model,
        )

    if sub == "1":
        run_topic_search(
            direct_instructions=direct_instructions,
            call_llm_fn=call_llm_fn,
            verbose=True,
            model=model,
        )
    else:
        run_article_search(
            direct_instructions=direct_instructions,
            call_llm_fn=call_llm_fn,
            verbose=True,
            model=model,
        )


def handle_writing_mode(
    provider: str = "ollama",
    model: str | None = None,
    dry_run: bool = False,
) -> None:
    """Terminal entry point for Writing mode (two-track: topic / article)."""
    try:
        from src.modes.writing import (
            run_writer, run_editor, run_qa,
            parse_direct_instructions as _pdi,
            TRACK_TOPIC, TRACK_ARTICLE, DEFAULT_WORDS,
        )
    except ModuleNotFoundError:
        from modes.writing import (
            run_writer, run_editor, run_qa,
            parse_direct_instructions as _pdi,
            TRACK_TOPIC, TRACK_ARTICLE, DEFAULT_WORDS,
        )

    def _call_llm_fn(system_prompt: str, user_prompt: str) -> str:
        if dry_run:
            return "[DRY RUN] LLM would respond here."
        combined = (
            f"## System Instructions\n{system_prompt}\n\n"
            f"## User Request\n{user_prompt}"
        )
        return call_ai(prompt=combined, provider=provider, model=model)

    # ── Track selection ──────────────────────────────────────────
    print("\n" + "=" * 60)
    print("  WRITING MODE")
    print("=" * 60)
    print("  1. Topic Track    — editorial / opinion (newspaper style)")
    print("  2. Article Track  — medical journal article")
    print("  0. Back to menu")
    print("=" * 60)
    try:
        track_choice = input("Select track [0-2]: ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        return
    if track_choice == "0" or not track_choice:
        print("Returning to menu.")
        return
    if track_choice == "1":
        track = TRACK_TOPIC
    elif track_choice == "2":
        track = TRACK_ARTICLE
    else:
        print("Invalid choice — returning to menu.")
        return

    # ── Sub-mode selection ───────────────────────────────────────
    print(f"\n  Track: {track.upper()}")
    print("  " + "-" * 40)
    print("  1. Writer  — full pipeline: Writer -> Editor -> QA")
    print("  2. Editor  — standalone: edit documents in input/writing/")
    print("  3. QA      — standalone: review documents in input/writing/")
    print("  0. Back")
    try:
        sub_choice = input("Select sub-mode [0-3]: ").strip()
    except (EOFError, KeyboardInterrupt):
        print()
        return
    if sub_choice == "0" or not sub_choice:
        print("Returning to menu.")
        return
    sub_mode_map = {"1": "Writer", "2": "Editor", "3": "QA"}
    if sub_choice not in sub_mode_map:
        print("Invalid choice — returning to menu.")
        return
    sub_mode = sub_mode_map[sub_choice]

    # ── Word limit ───────────────────────────────────────────────
    if sub_mode in ("Writer", "Editor"):
        default_wl = DEFAULT_WORDS[track]
        print(f"\n  Default word limit for {track} track: {default_wl}")
        try:
            wl_input = input(
                f"  Press Enter to keep {default_wl}, or type a number: "
            ).strip()
        except (EOFError, KeyboardInterrupt):
            wl_input = ""
        word_limit = int(wl_input) if wl_input.isdigit() and int(wl_input) > 0 \
            else default_wl
        print(f"  Word limit set to: {word_limit}\n")
    else:
        word_limit = DEFAULT_WORDS[track]

    # ── Instructions ─────────────────────────────────────────────
    print(f"  [{sub_mode.upper()} | {track.upper()}] Enter instructions below.")
    print("  Lines starting with > are DIRECT TASK INSTRUCTIONS (highest priority).")
    print("  Press ENTER on a blank line when done.\n")

    lines: list[str] = []
    while True:
        try:
            line = input("  instruction> ")
        except (EOFError, KeyboardInterrupt):
            break
        if line.strip() == "":
            print("\n  Instructions received -- starting processing...\n")
            break
        if not line.strip().startswith(">"):
            line = "> " + line.strip()
        lines.append(line)

    direct_instructions = _pdi("\n".join(lines))
    print()

    if sub_mode == "Writer":
        run_writer(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            track=track,
            word_limit=word_limit,
            verbose=True,
        )
    elif sub_mode == "Editor":
        run_editor(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            track=track,
            word_limit=word_limit,
            verbose=True,
        )
    elif sub_mode == "QA":
        run_qa(
            direct_instructions=direct_instructions,
            call_llm_fn=_call_llm_fn,
            track=track,
            verbose=True,
        )

    print(f"\n  [{sub_mode.upper()} | {track.upper()}] Done. "
          "Check output/writing/ and reports/writing/ for results.\n")


# ---------------------------------------------------------------------------
# Main interactive session loop
# ---------------------------------------------------------------------------
def main(
    model_override: str | None = None,
    dry_run: bool = False,
    mode: str = "coding",
    provider: str = "ollama",
) -> None:
    """Run an interactive AI session."""
    session_id  = uuid.uuid4().hex[:8]
    transcript  = start_session_transcript(REPORTS_DIR)
    step_count  = 0
    role_counts: dict[str, int] = {
        r: 0 for mode_roles in ALL_MODES.values() for r in mode_roles
    }

    print(f"\n{'='*55}")
    print(f"  AI Automation Tool  v{VERSION}")
    print(f"  Mode    : {mode}")
    print(f"  Provider: {provider}")
    print(f"  Session : {session_id}")
    if dry_run:
        print("  DRY RUN -- AI calls will be skipped")
    print(f"{'='*55}\n")

    input_files  = auto_load_input_files(mode)
    rag_enabled  = False
    input_folder = INPUT_DIR / mode
    input_folder.mkdir(parents=True, exist_ok=True)

    if not dry_run:
        try:
            from src import rag as rag_module
            n_chunks = rag_module.index_uploads(
                mode=mode,
                session_id=session_id,
                upload_base=str(INPUT_DIR),
            )
            if n_chunks > 0:
                rag_enabled = True
                print(f"[RAG] Indexed {n_chunks} chunk(s) from input/{mode}/\n")
        except Exception as exc:  # noqa: BLE001
            print(f"[RAG] Indexing skipped: {exc}\n")

    role_name, role_cfg = choose_role(mode)
    colour      = role_color(role_name)
    prompt_path = Path(role_cfg["prompt"])

    if not prompt_path.exists():
        print(f"Prompt file not found: {prompt_path}")
        sys.exit(1)

    role_prompt = prompt_path.read_text(encoding="utf-8", errors="replace")
    print(f"\n{colour}Starting session as {role_name}{RESET}")
    print("Type your task below. Press Ctrl+C to end the session.\n")

    previous_response = ""

    article_context = ""
    if mode == "appraisal":
        articles = _read_article_files(INPUT_APPRAISAL)
        if articles:
            parts_art = [
                f"### Article: {art['name']}\n{art['content']}"
                for art in articles
            ]
            article_context = "\n\n".join(parts_art)
            print(f"[Appraisal] {len(articles)} article(s) loaded from input/appraisal/.")
        else:
            print("[Appraisal] No articles in input/appraisal/ -- using RAG only.")

    try:
        while True:
            try:
                task = input(f"{colour}{role_name} >{RESET} ").strip()
            except (EOFError, KeyboardInterrupt):
                print()
                break
            if not task:
                continue

            context = build_project_context(
                role_name=role_name,
                query=task if rag_enabled else "",
                mode=mode,
                session_id=session_id,
            )

            parts = [role_prompt]
            if mode == "appraisal" and article_context:
                parts.append(f"## Articles for Appraisal\n{article_context}")
            if context:
                parts.append(f"## Project Context\n{context}")
            if previous_response:
                parts.append(
                    f"## Previous Response\n{truncate_context(previous_response)}"
                )
            parts.append(f"## Task\n{task}")
            full_prompt = "\n\n".join(parts)

            start = time.time()
            if dry_run:
                response = f"[DRY RUN] {role_name} would respond here."
            else:
                try:
                    response = call_ai(
                        prompt=full_prompt,
                        provider=provider,
                        model=model_override,
                    )
                except RuntimeError as exc:
                    response = f"[ERROR] {exc}"

            elapsed     = time.time() - start
            step_count += 1
            role_counts[role_name] = role_counts.get(role_name, 0) + 1

            print(f"\n{colour}--- {role_name} response ({elapsed:.1f}s) ---{RESET}")
            print(response)
            print()

            append_to_transcript(
                path=transcript,
                role_name=role_name,
                step=step_count,
                task=task,
                response=response,
            )
            previous_response = response

            if mode == "rct_search" and not dry_run:
                save_rct_search_links(response=response, reports_dir=REPORTS_DIR)

    except KeyboardInterrupt:
        print(f"\n\n{colour}Session ended.{RESET}")

    finally:
        print_session_summary(transcript, step_count, role_counts)
        if rag_enabled and not dry_run:
            try:
                from src import rag as rag_module
                rag_module.clear_session(mode=mode, session_id=session_id)
                print("[RAG] Session collection cleared.")
            except Exception as exc:  # noqa: BLE001
                print(f"[RAG] Clear warning: {exc}")


# ---------------------------------------------------------------------------
# API key validation
# ---------------------------------------------------------------------------
PROVIDER_ENV_VARS = {
    "anthropic": "ANTHROPIC_API_KEY",
    "openai":    "OPENAI_API_KEY",
    "deepseek":  "DEEPSEEK_API_KEY",
    "groq":      "GROQ_API_KEY",
    "qwen":      "DASHSCOPE_API_KEY",
    "ollama":    None,
}


def validate_api_keys(provider: str) -> None:
    """
    Validate the required API key is set for the selected provider.
    Raises EnvironmentError if missing. Ollama requires no key.
    """
    provider = provider.lower().strip()
    if provider not in PROVIDER_ENV_VARS:
        raise ValueError(
            f"Unknown provider '{provider}'. "
            f"Valid options: {', '.join(PROVIDER_ENV_VARS.keys())}"
        )
    env_var = PROVIDER_ENV_VARS[provider]
    if env_var is None:
        return
    value = os.environ.get(env_var, "").strip()
    if not value:
        raise EnvironmentError(
            f"\n"
            f"  [API KEY ERROR]\n"
            f"  Provider '{provider}' requires {env_var} to be set.\n"
            f"  Add it to your .env file in the project root:\n"
            f"  {env_var}=your-key-here\n"
        )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def open_help_guide() -> None:
    """Open the HTML flashcard help guide in the default browser."""
    import webbrowser
    guide_path = DOCS_DIR / "flashcard-help.html"
    if not guide_path.exists():
        print(f"Help guide not found: {guide_path}")
        print("Expected location: docs/flashcard-help.html")
        sys.exit(1)
    webbrowser.open(guide_path.as_uri())
    print(f"Opening help guide in browser: {guide_path.name}")


def parse_args(args: list[str] | None = None) -> argparse.Namespace:
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(
        description="AI Automation Tool",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python src/main.py                           # coding session\n"
            "  python src/main.py --mode writing            # writing mode\n"
            "  python src/main.py --mode rct_search         # RCT search mode\n"
            "  python src/main.py --mode appraisal          # appraisal mode\n"
            "  python src/main.py --mode search             # medical search\n"
            "  python src/main.py --mode sr                 # SR pipeline\n"
            "  python src/main.py --model llama3.2:3b       # different model\n"
            "  python src/main.py --provider qwen           # use Qwen\n"
            "  python src/main.py --list-sessions           # list transcripts\n"
            "  python src/main.py --list-roles              # show roles/docs\n"
            "  python src/main.py --list-roles --mode rct_search\n"
            "  python src/main.py --dry-run                 # simulate session\n"
            "  python src/main.py --version                 # show version\n"
        ),
    )
    parser.add_argument("--model",          type=str, default=None)
    parser.add_argument("--mode",           type=str, default="coding",
                        choices=["coding", "writing", "rct_search",
                                 "appraisal", "search", "sr"])
    parser.add_argument("--report",         action="store_true", default=False,
                        help="Generate a writing report from input/writing/ files")
    parser.add_argument("--revise",         action="store_true", default=False,
                        help="Run legacy code revision pipeline from docs/coding/")
    parser.add_argument("--role",           type=str, default="Builder",
                        choices=["Builder", "Reviewer", "Tester"],
                        help="Starting role for --revise pipeline (default: Builder)")
    parser.add_argument("--provider",       type=str, default="ollama",
                        choices=["ollama", "openai", "anthropic", "deepseek", "groq", "qwen"])
    parser.add_argument("--list-sessions",  action="store_true", default=False)
    parser.add_argument("--read-session",   type=str, default=None, metavar="FILENAME")
    parser.add_argument("--delete-session", type=str, default=None, metavar="FILENAME")
    parser.add_argument("--export-session", type=str, default=None, metavar="FILENAME")
    parser.add_argument("--rename-session", type=str, default=None, metavar="FILENAME")
    parser.add_argument("--stats",          action="store_true", default=False)
    parser.add_argument("--dry-run",        action="store_true", default=False)
    parser.add_argument("--version",        action="version",
                        version=f"AI Automation Tool v{VERSION}")
    parser.add_argument("--list-roles",     action="store_true", default=False)
    parser.add_argument("--help-guide",     action="store_true", default=False,
                        help="Open the interactive HTML help guide in your browser")
    parser.add_argument("--ui",             action="store_true", default=False,
                        help="Launch the main Streamlit UI")
    return parser.parse_args(args)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    try:
        args = parse_args()

        if args.ui:
            import subprocess as _sp
            proc = _sp.Popen(
                [sys.executable, "-m", "streamlit", "run",
                 str(Path(__file__).resolve().parent / "ui" / "app.py"),
                 "--server.runOnSave", "false"],
                cwd=str(Path(__file__).resolve().parent.parent),
            )
            print("UI launched -- open http://localhost:8501 in your browser.")
            print("Press Ctrl+C to stop the UI server.")
            try:
                proc.wait()
            except KeyboardInterrupt:
                proc.terminate()
            raise SystemExit(0)

        non_ai_flags = (
            args.list_sessions
            or args.read_session
            or args.delete_session
            or args.export_session
            or args.rename_session
            or args.stats
            or args.list_roles
            or args.help_guide
            or args.dry_run
        )
        if not non_ai_flags:
            try:
                validate_api_keys(args.provider)
            except (EnvironmentError, ValueError) as e:
                print(e)
                sys.exit(1)

        if args.list_sessions:
            list_sessions(reports_dir=str(REPORTS_DIR))
        elif args.read_session:
            read_session(filename=args.read_session, reports_dir=str(REPORTS_DIR))
        elif args.delete_session:
            delete_session(filename=args.delete_session, reports_dir=str(REPORTS_DIR))
        elif args.export_session:
            export_session(filename=args.export_session, reports_dir=str(REPORTS_DIR))
        elif args.rename_session:
            rename_session(filename=args.rename_session, reports_dir=str(REPORTS_DIR))
        elif args.stats:
            show_stats(reports_dir=str(REPORTS_DIR))
        elif args.list_roles:
            list_roles(mode=args.mode)
        elif args.help_guide:
            open_help_guide()
        elif args.revise:
            if args.mode != "coding":
                print("--revise is only available in coding mode.")
                print("Use: python src/main.py --mode coding --revise")
                sys.exit(1)
            generate_code_revision(
                start_role=args.role,
                provider=args.provider,
                model=args.model,
                dry_run=args.dry_run,
            )
        elif args.report:
            if args.mode != "writing":
                print("--report is only available in writing mode.")
                print("Use: python src/main.py --mode writing --report")
                sys.exit(1)
            generate_writing_report(provider=args.provider, model=args.model)
        elif args.mode == "search":
            handle_search_mode(
                provider=args.provider,
                model=args.model,
            )
        elif args.mode == "sr":
            run_sr_launcher()
        elif args.mode == "rct_search":
            run_rct_search_pipeline(
                provider=args.provider,
                model=args.model,
                dry_run=args.dry_run,
            )
        elif args.mode == "coding":
            handle_coding_mode(
                provider=args.provider,
                model=args.model,
                dry_run=args.dry_run,
            )
        elif args.mode == "writing":
            handle_writing_mode(
                provider=args.provider,
                model=args.model,
                dry_run=args.dry_run,
            )
        elif args.mode == "appraisal":
            handle_appraisal_mode(
                provider=args.provider,
                model=args.model,
                dry_run=args.dry_run,
            )
        else:
            main(
                model_override=args.model,
                dry_run=args.dry_run,
                mode=args.mode,
                provider=args.provider,
            )

    except KeyboardInterrupt:
        print("\n\nSession stopped. Returning to menu...\n")
        raise SystemExit(0)
