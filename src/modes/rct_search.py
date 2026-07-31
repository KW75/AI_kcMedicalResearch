"""
RCT Search Mode - PICO-driven PubMed search and AI ranking
Extracted from src/main.py for modularity
"""
import json
import re
import shutil
import sys
import urllib.parse
from datetime import datetime
from pathlib import Path
from typing import Optional

# Project paths
BASE = Path(__file__).resolve().parent.parent.parent
INPUT_DIR = BASE / "input"
OUTPUT_RCT_SEARCH = BASE / "output" / "rct_search"
REPORTS_DIR = BASE / "reports"
DOCS_RCT_SEARCH = BASE / "docs" / "rct_search"
AI_DIR = BASE / "ai"
INPUT_SR = INPUT_DIR / "sr"


# ANSI colors for console output
RESET = "\033[0m"
ACCENT = "\033[38;5;121m"
DIM = "\033[0;31;40m"


def role_color(role: str) -> str:
    """Return color for a specific role"""
    colors = {
        "Formulator": "\033[38;5;51m",
        "Searcher": "\033[38;5;121m",
        "Validator": "\033[38;5;215m",
    }
    return colors.get(role, ACCENT)


def _read_topic_file(topic_file: Path) -> Optional[str]:
    """Read topic from file if it exists"""
    if topic_file.exists():
        try:
            return topic_file.read_text(encoding="utf-8").strip()
        except Exception:
            return None
    return None


def _clean_pico_term(t: str) -> str:
    """Clean PICO terms for PubMed query"""
    t = re.sub(r"\([^)]*\)", "", t)           # remove parenthetical qualifiers
    t = re.sub(r"\b(at\s+)?(post[\s\-]intervention|post[\s\-]treatment|"
               r"follow[\s\-]up|immediately after|baseline|"
               r"at\s+\d+\s+weeks?|at\s+\d+\s+months?)\b", "", t, flags=re.IGNORECASE)
    t = re.sub(r"[^\w\s\-]", " ", t)          # remove special chars except hyphen
    t = re.sub(r"\s+", " ", t).strip()         # collapse whitespace
    # keep only first 4 words to avoid over-specific phrases
    words = t.split()
    return " ".join(words[:4]) if len(words) > 4 else t


def fetch_pubmed_articles(
    query: str,
    max_results: int = 100,
) -> list[dict]:
    """
    Search PubMed via NCBI E-utilities (no API key required).
    Returns list of dicts: pmid, title, abstract, url.
    Returns empty list on any error.
    """
    import urllib.request
    import xml.etree.ElementTree as ET

    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    search_url = (
        f"{base}/esearch.fcgi?db=pubmed&term={urllib.parse.quote(query)}"
        f"&retmax={max_results}&retmode=json"
    )

    try:
        with urllib.request.urlopen(search_url) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except Exception:
        return []

    ids = data.get("esearchresult", {}).get("idlist", [])
    if not ids:
        return []

    ids_str = ",".join(ids)
    fetch_url = (
        f"{base}/efetch.fcgi?db=pubmed&id={ids_str}"
        f"&rettype=abstract&retmode=xml"
    )

    try:
        with urllib.request.urlopen(fetch_url) as resp:
            xml_data = resp.read().decode("utf-8")
    except Exception:
        return []

    try:
        root = ET.fromstring(xml_data)
    except ET.ParseError:
        return []

    articles = []
    for article in root.findall(".//PubmedArticle"):
        pmid_el = article.find(".//PMID")
        title_el = article.find(".//ArticleTitle")
        abstract_el = article.find(".//AbstractText")

        pmid = pmid_el.text if pmid_el is not None and pmid_el.text else ""
        title = title_el.text if title_el is not None and title_el.text else ""
        abstract = abstract_el.text if abstract_el is not None and abstract_el.text else ""

        if pmid and title:
            articles.append({
                "pmid": pmid,
                "title": title,
                "abstract": abstract,
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
            })

    return articles


def call_ai(prompt: str, provider: str = "ollama", model: Optional[str] = None) -> str:
    """
    Call AI provider with prompt.
    This function should be imported from the main module,
    but is defined here as a fallback.
    """
    try:
        # Try to import from main
        from src.main import call_ai as _call_ai
        return _call_ai(prompt=prompt, provider=provider, model=model)
    except ImportError:
        # Fallback: use direct API call
        from openai import OpenAI
        import os
        from dotenv import load_dotenv

        load_dotenv(BASE / ".env")

        if provider == "ollama":
            client = OpenAI(base_url="http://localhost:11434/v1", api_key="ollama")
            model_name = model or "qwen2.5:7b"
        elif provider == "qwen":
            client = OpenAI(
                api_key=os.getenv("QWEN_API_KEY"),
                base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
            )
            model_name = model or "qwen-vl-plus"
        else:
            # Default to qwen
            client = OpenAI(
                api_key=os.getenv("QWEN_API_KEY"),
                base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
            )
            model_name = model or "qwen-vl-plus"

        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=4096,
        )
        return response.choices[0].message.content


def _md_to_docx(md_content: str, title: str, out_path: Path) -> None:
    """Convert markdown to DOCX using python-docx"""
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = docx.Document()
        # Title
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True

        doc.add_paragraph()

        # Process markdown
        for line in md_content.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("|") and "|" in line[1:]:
                # Skip table headers for now
                pass
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Pt(36)
            elif line.strip():
                doc.add_paragraph(line)

        doc.save(str(out_path))
    except ImportError:
        # Fallback: save as text
        out_path.with_suffix(".txt").write_text(md_content, encoding="utf-8")


def _ranked_articles_to_docx(
    ranked: list[dict],
    title: str,
    out_path: Path,
    topic: str = ""
) -> None:
    """Write ranked article list as DOCX with table"""
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = docx.Document()
        # Title
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True

        if topic:
            doc.add_paragraph(f"Topic: {topic}")

        doc.add_paragraph()

        if not ranked:
            doc.add_paragraph("No articles retrieved from PubMed.")
            doc.save(str(out_path))
            return

        # Caption
        cap = doc.add_paragraph()
        cap.add_run(
            f"All {len(ranked)} articles retrieved from PubMed, "
            "ordered by PICO relevance score (10 = most relevant)."
        ).italic = True

        doc.add_paragraph()

        # Table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header
        hdr = table.rows[0].cells
        hdr[0].text = "Rank"
        hdr[1].text = "Score"
        hdr[2].text = "Title"
        hdr[3].text = "PMID"
        hdr[4].text = "PubMed Link"

        # Data rows
        for r in ranked:
            row = table.add_row().cells
            row[0].text = str(r["rank"])
            row[1].text = f"{r['score']}/10"
            row[2].text = r["title"]
            row[3].text = r["pmid"]
            # Add link
            p = row[4].paragraphs[0]
            p.add_run().add_hyperlink(r["url"], "PubMed")

        doc.save(str(out_path))
    except ImportError:
        # Fallback: CSV
        import csv
        with open(out_path.with_suffix(".csv"), "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Rank", "Score", "Title", "PMID", "URL"])
            for r in ranked:
                writer.writerow([r["rank"], r["score"], r["title"], r["pmid"], r["url"]])


def run_rct_search_pipeline(
    provider: str = "ollama",
    model: str | None = None,
    reports_dir: Path = REPORTS_DIR,
    dry_run: bool = False,
) -> Path:
    """
    Single-pass RCT search pipeline:
      1. Formulator - structures user topic into PICO question
      2. Searcher - builds Boolean search strategy for all 7 databases
      3. Validator - validates alignment and approves or requests refinement
      4. PubMed Fetch + AI Ranking
    Saves output as output/rct_search/rct_search_{ts}.md and .docx.
    Returns path of the .md report.
    """
    print("\n" + "=" * 55)
    print("  RCT SEARCH PIPELINE")
    print("=" * 55)
    print("  This pipeline will:")
    print("  1. Structure your topic into a PICO question")
    print("  2. Build a search strategy for all 7 SR databases")
    print("  3. Validate the strategy before download")
    print("  4. Fetch and rank articles from PubMed")
    print("  No article appraisal - use --mode appraisal for that.")
    print("=" * 55 + "\n")

    topic_file = DOCS_RCT_SEARCH / "topic.md"
    file_topic = _read_topic_file(topic_file)
    if file_topic:
        topic = file_topic
        print(f"[RCT Search] Topic loaded from docs/rct_search/topic.md: {topic}")
    else:
        try:
            topic = input("Enter your research topic: ").strip()
        except (EOFError, KeyboardInterrupt):
            print()
            topic = ""
    if not topic:
        print("No topic entered. Exiting.")
        sys.exit(0)

    # -- PICO input: offer to import saved JSON or prompt manually --------------

    _pico_input_dir = INPUT_DIR / "rct_search"
    pico_json_files = sorted(
        _pico_input_dir.glob("pico_*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    ) if _pico_input_dir.exists() else []
    if not pico_json_files and OUTPUT_RCT_SEARCH.exists():
        pico_json_files = sorted(
            OUTPUT_RCT_SEARCH.glob("pico_*.json"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )
    _pico_source = "input/rct_search" if (_pico_input_dir.exists() and list(_pico_input_dir.glob("pico_*.json"))) else "output/rct_search"

    pico_population = ""
    pico_intervention = ""
    pico_comparator = ""
    pico_outcome = ""

    if pico_json_files:
        print()
        print(f"[RCT Search] Saved PICO files found in {_pico_source}:")
        for idx, pf in enumerate(pico_json_files[:5], 1):
            print(f"  {idx}. {pf.name}")
        print("  0. Enter new PICO manually")
        print()
        try:
            choice = input("Import a saved PICO? Enter number or 0 for new: ").strip()
        except (EOFError, KeyboardInterrupt):
            choice = "0"
        if choice.isdigit() and 1 <= int(choice) <= len(pico_json_files[:5]):
            chosen = pico_json_files[int(choice) - 1]
            try:
                pico_data = json.loads(chosen.read_text(encoding="utf-8"))
                pico_population = pico_data.get("population", "")
                pico_intervention = pico_data.get("intervention", "")
                pico_comparator = pico_data.get("comparator", "")
                pico_outcome = pico_data.get("outcome", "")
                if not topic:
                    topic = pico_data.get("topic", topic)
                print(f"[RCT Search] PICO imported from {chosen.name}")
            except Exception as _e:
                print(f"[RCT Search] Could not load PICO JSON: {_e}")

    if not any([pico_population, pico_intervention, pico_comparator, pico_outcome]):
        print()
        print("[RCT Search] Enter PICO components (press Enter to skip any field):")
        try:
            pico_population = input("  Population   (P): ").strip()
            pico_intervention = input("  Intervention (I): ").strip()
            pico_comparator = input("  Comparator   (C): ").strip()
            pico_outcome = input("  Outcome      (O): ").strip()
        except (EOFError, KeyboardInterrupt):
            print()

    pico_manual_context = ""
    if any([pico_population, pico_intervention, pico_comparator, pico_outcome]):
        pico_manual_context = (
            "\n\nUser-provided PICO components:\n"
            f"  Population:   {pico_population or '(not specified)'}\n"
            f"  Intervention: {pico_intervention or '(not specified)'}\n"
            f"  Comparator:   {pico_comparator or '(not specified)'}\n"
            f"  Outcome:      {pico_outcome or '(not specified)'}"
        )
    # -- end PICO input ---------------------------------------------------------

    pico_path = DOCS_RCT_SEARCH / "pico-framework.md"
    db_guide = DOCS_RCT_SEARCH / "database-guide.md"
    val_criteria = DOCS_RCT_SEARCH / "validation-criteria.md"
    pico_context = pico_path.read_text(encoding="utf-8", errors="replace") if pico_path.exists() else ""
    db_guide_text = db_guide.read_text(encoding="utf-8", errors="replace") if db_guide.exists() else ""
    val_criteria_text = val_criteria.read_text(encoding="utf-8", errors="replace") if val_criteria.exists() else ""

    stages = [
        {
            "role": "Formulator",
            "prompt_file": AI_DIR / "formulator-prompt.md",
            "extra_context": pico_context,
            "task": f"The user's research topic is: {topic}{pico_manual_context}\n\nStructure this into a formal PICO question.",
        },
        {
            "role": "Searcher",
            "prompt_file": AI_DIR / "searcher-prompt.md",
            "extra_context": db_guide_text,
            "task": "Build a comprehensive Boolean search strategy for all 7 SR databases based on the PICO question above.",
        },
        {
            "role": "Validator",
            "prompt_file": AI_DIR / "validator-prompt.md",
            "extra_context": val_criteria_text,
            "task": "Validate the search strategy above. Check PICO alignment, database coverage, syntax, and RCT filters. Return APPROVED FOR DOWNLOAD or REQUIRES REFINEMENT with specific justification.",
        },
    ]

    report_parts = [
        "# RCT Search Strategy Report",
        f"Topic: {topic}",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        "---",
        "",
        "> **Note:** This report contains the validated search strategy only.",
        "> For article appraisal, copy URLs into --mode appraisal.",
        "",
    ]

    previous_response = ""

    for stage in stages:
        role = stage["role"]
        colour = role_color(role)
        print(f"Running {role}...")

        prompt_file = Path(stage["prompt_file"])
        if not prompt_file.exists():
            print(f"  Prompt not found: {prompt_file} - skipping {role}.")
            continue

        role_prompt = prompt_file.read_text(encoding="utf-8", errors="replace")
        parts = [role_prompt]
        if pico_context:
            parts.append(f"## PICO Framework Reference\n{pico_context}")
        if stage["extra_context"]:
            parts.append(f"## Reference Document\n{stage['extra_context']}")
        if previous_response:
            parts.append(f"## Previous Stage Output\n{previous_response}")
        parts.append(f"## Task\n{stage['task']}")
        full_prompt = "\n\n".join(parts)

        if dry_run:
            response = f"[DRY RUN] {role} would respond here."
        else:
            try:
                response = call_ai(prompt=full_prompt, provider=provider, model=model)
            except RuntimeError as exc:
                response = f"[ERROR in {role}: {exc}]"

        previous_response = response
        report_parts.append(f"## {role} Output\n\n{response}\n")
        print(f"{colour}{role} complete.{RESET}\n")

        # -- save PICO JSON after Formulator completes -------------------------
        if role == "Formulator" and not dry_run:
            pico_json_dir = OUTPUT_RCT_SEARCH
            pico_json_dir.mkdir(parents=True, exist_ok=True)
            pico_json_path = pico_json_dir / f"pico_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            pico_data = {
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "topic": topic,
                "formulator_output": response,
                "population": "",
                "intervention": "",
                "comparator": "",
                "outcome": "",
                "study_design": "RCT",
                "effect_measure": "SMD",
                "source_mode": "rct_search",
                "pubmed_query_raw": "",
                "pubmed_query_cleaned": "",
            }

            def _extract_pico_field(text, labels):
                for lbl in labels:
                    m = re.search(r"(?i)" + lbl + r"[^:]*:\s*(.+)", text)
                    if m:
                        return m.group(1).strip()
                return ""

            pico_data["population"] = pico_population or _extract_pico_field(response, [r"P\s*\(Population\)", r"Population"])
            pico_data["intervention"] = pico_intervention or _extract_pico_field(response, [r"I\s*\(Intervention\)", r"Intervention"])
            pico_data["comparator"] = pico_comparator or _extract_pico_field(response, [r"C\s*\(Comp\w+\)", r"Compar"])
            pico_data["outcome"] = pico_outcome or _extract_pico_field(response, [r"O\s*\(Outcome\)", r"Outcome"])
            pico_json_path.write_text(
                json.dumps(pico_data, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
            print(f"[RCT Search] PICO saved to: {pico_json_path.name}")

        # ----------------------------------------------------------------------

    validator_output = previous_response.upper()
    if "APPROVED FOR DOWNLOAD" in validator_output:
        status = "APPROVED FOR DOWNLOAD"
        print("\n Search strategy APPROVED FOR DOWNLOAD.")
    else:
        status = "REQUIRES REFINEMENT"
        print("\n Search strategy REQUIRES REFINEMENT - see Validator output.")

    report_parts.append(f"## Final Status\n\n**{status}**\n")
    report_parts.append(
        "## Next Steps\n\n"
        "- If APPROVED: copy database search strings into each SR database platform\n"
        "- Download article lists from each database\n"
        "- Run python src/main.py --mode appraisal to appraise individual articles\n"
        "- Run python src/main.py --mode sr for full systematic review pipeline\n"
    )

    # -- Stage 4: PubMed Fetch + AI Ranking ----------------------------------

    # Fix 1: Broaden the RCT filter
    _pico_terms = [_clean_pico_term(t) for t in [pico_population, pico_intervention, pico_outcome] if t]
    _pico_terms = [t for t in _pico_terms if t]
    _pubmed_query = " AND ".join(_pico_terms) if _pico_terms else topic

    # KEY FIX: Broadened RCT filter to find more papers
    _pubmed_query = _pubmed_query + " AND (randomized controlled trial[pt] OR clinical trial[pt] OR random*[tiab] OR RCT[tiab])"
    print(f"[RCT Search] PubMed query (cleaned): {_pubmed_query}")
    print("[RCT Search] Searching PubMed for relevant articles...")

    # Fix 2: Increase max_results to 100
    articles = fetch_pubmed_articles(_pubmed_query, max_results=100) if not dry_run else [
        {"pmid": "00000001", "title": "[DRY RUN] Test Article",
         "abstract": "Dry run abstract.",
         "url": "https://pubmed.ncbi.nlm.nih.gov/00000001/"}
    ]

    if articles:
        print(f"[RCT Search] Found {len(articles)} article(s). Ranking by PICO relevance...")
        abstracts_block = ""
        for idx, art in enumerate(articles, 1):
            abstracts_block += (
                "\n### Article " + str(idx) + "\n"
                + "Title: " + art["title"] + "\n"
                + "PMID: " + art["pmid"] + "\n"
                + "URL: " + art["url"] + "\n"
                + "Abstract: " + art["abstract"] + "\n"
            )
        rank_prompt = (
            "You are a systematic review expert. Your ONLY task is to rate articles.\n\n"
            "PICO Question:\n" + previous_response + "\n\n"
            "INSTRUCTIONS:\n"
            "- Rate each article for relevance to the PICO question on a scale of 1-10.\n"
            "- 10 = highly relevant RCT directly matching the PICO.\n"
            "- 1 = not relevant.\n"
            "- Output ONLY the rating lines. No summaries, no headers, no extra text.\n"
            "- Each line MUST follow this EXACT format:\n"
            "ARTICLE_RANK: 1 | SCORE: 9 | PMID: 12345678 | TITLE: Example title here | URL: https://pubmed.ncbi.nlm.nih.gov/12345678/\n\n"
            "EXAMPLE OUTPUT (do not copy - use real values):\n"
            "ARTICLE_RANK: 1 | SCORE: 9 | PMID: 11111111 | TITLE: CBT for fibromyalgia RCT | URL: https://pubmed.ncbi.nlm.nih.gov/11111111/\n"
            "ARTICLE_RANK: 2 | SCORE: 7 | PMID: 22222222 | TITLE: Mindfulness for chronic pain | URL: https://pubmed.ncbi.nlm.nih.gov/22222222/\n"
            "ARTICLE_RANK: 3 | SCORE: 3 | PMID: 33333333 | TITLE: Ketamine infusion study | URL: https://pubmed.ncbi.nlm.nih.gov/33333333/\n\n"
            "NOW RATE THESE ARTICLES - output rating lines ONLY:\n\n"
            + abstracts_block
        )
        if dry_run:
            rank_response = "\n".join(
                "ARTICLE_RANK: " + str(i) + " | SCORE: " + str(10 - i) + " | "
                + "PMID: " + art["pmid"] + " | TITLE: " + art["title"]
                + " | URL: " + art["url"]
                for i, art in enumerate(articles, 1)
            )
        else:
            try:
                rank_response = call_ai(
                    prompt=rank_prompt, provider=provider, model=model
                )
            except RuntimeError as exc:
                rank_response = "[ERROR ranking articles: " + str(exc) + "]"
        ranked = []
        if rank_response.startswith("[ERROR"):
            print(f"[RCT Search] Ranking error: {rank_response}")
        for ln in rank_response.splitlines():
            m = re.match(
                r"ARTICLE_RANK:\s*(\d+)\s*\|\s*SCORE:\s*(\d+)\s*\|"
                r"\s*PMID:\s*(\S+)\s*\|\s*TITLE:\s*(.+?)\s*\|\s*URL:\s*(\S+)",
                ln.strip()
            )
            if m:
                ranked.append({
                    "rank": int(m.group(1)),
                    "score": int(m.group(2)),
                    "pmid": m.group(3),
                    "title": m.group(4),
                    "url": m.group(5),
                })
        ranked.sort(key=lambda x: x["score"], reverse=True)
        for new_rank, r in enumerate(ranked, 1):
            r["rank"] = new_rank
        if not ranked:
            print("[RCT Search] Warning: AI ranking returned no parseable results.")
            print("[RCT Search] Raw ranking response (first 500 chars):")
            print(rank_response[:500])
            # fallback: use fetch order, score descending by position
            ranked = [
                {"rank": i, "score": len(articles) - i + 1,
                 "pmid": a["pmid"], "title": a["title"], "url": a["url"]}
                for i, a in enumerate(articles, 1)
            ]
        table_lines = [
            "## Ranked Article List\n",
            f"_All {len(ranked)} articles retrieved from PubMed, ordered by PICO relevance score (10 = most relevant)._\n",
            "| Rank | Score | Title | PMID | Link |",
            "|------|-------|-------|------|------|",
        ]
        for r in ranked:
            table_lines.append(
                "| " + str(r["rank"]) + " | " + str(r["score"]) + "/10 | "
                + r["title"] + " | " + r["pmid"]
                + " | [PubMed](" + r["url"] + ") |"
            )
        table_lines.append(
            "\n> Select your top articles, download PDFs "
            "and place them in input/sr/ to run the SR pipeline."
        )
        table_lines.append(
            "\n> For explanation on ranking, please refer to the full report "
            "in the reports folder."
        )

        report_parts.append("\n".join(table_lines))
        top = ranked[0]["title"] if ranked else "N/A"
        print("[RCT Search] Ranking complete. Top article: " + top)
        pico_files = sorted(OUTPUT_RCT_SEARCH.glob("pico_*.json"), reverse=True)
        if pico_files:
            pico_data = json.loads(pico_files[0].read_text(encoding="utf-8"))
            pico_data["ranked_articles"] = ranked
            pico_data["pubmed_query_raw"] = " AND ".join(
                [t for t in [pico_population, pico_intervention, pico_outcome] if t]
            ) if any([pico_population, pico_intervention, pico_outcome]) else topic
            pico_data["pubmed_query_cleaned"] = _pubmed_query
            pico_files[0].write_text(
                json.dumps(pico_data, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
            print("[RCT Search] PICO JSON updated with "
                  + str(len(ranked)) + " ranked article(s).")
    else:
        print("[RCT Search] No articles found on PubMed for this topic.")
        report_parts.append(
            "## Ranked Article List\n\n_No articles retrieved from PubMed._\n"
        )

    # -- End Stage 4 ----------------------------------------------------------

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if reports_dir != REPORTS_DIR:
        out_dir = reports_dir
        final_dir = reports_dir
    else:
        out_dir = REPORTS_DIR / "rct_search"
        final_dir = OUTPUT_RCT_SEARCH
    out_dir.mkdir(parents=True, exist_ok=True)
    final_dir.mkdir(parents=True, exist_ok=True)
    md_path = out_dir / f"rct_search_{timestamp}.md"
    docx_path = out_dir / f"rct_search_{timestamp}.docx"

    md_content = "\n".join(report_parts)
    md_path.write_text(md_content, encoding="utf-8")
    print(f"Intermediate report   : reports/rct_search/{md_path.name}")

    final_parts = [p for p in report_parts if "Ranked Article List" in p or "Final Status" in p]
    if final_parts and final_dir != out_dir:
        final_md_path = final_dir / f"rct_search_{timestamp}.md"
        final_md_content = "\n".join(final_parts)
        final_md_path.write_text(final_md_content, encoding="utf-8")
        print(f"Final report saved    : output/rct_search/{final_md_path.name}")
        final_docx_path = final_dir / f"rct_search_{timestamp}.docx"
        try:
            _ranked_articles_to_docx(ranked, f"RCT Search Results - {timestamp}", final_docx_path, topic=topic)
            print(f"Results DOCX saved    : output/rct_search/{final_docx_path.name}")
        except Exception as exc:
            print(f"  Warning: could not generate results .docx - {exc}")

    try:
        _md_to_docx(md_content, f"RCT Search Strategy - {timestamp}", docx_path)
        print(f"Word document saved   : {out_dir.name}\\{docx_path.name}")
    except Exception as exc:
        print(f"  Warning: could not generate .docx - {exc}")

    # -- Auto-copy PICO JSON to input/sr/ -------------------------------------
    if not dry_run:
        pico_files = sorted(OUTPUT_RCT_SEARCH.glob("pico_*.json"), reverse=True)
        if pico_files:
            latest_pico = pico_files[0]
            print()
            print(f"[RCT Search] Latest PICO JSON: {latest_pico.name}")
            try:
                copy_choice = input(
                    "[RCT Search] Copy pico_*.json to input/sr/ for SR pipeline? [Y/N]: "
                ).strip().upper()
            except (EOFError, KeyboardInterrupt):
                copy_choice = "N"
            if copy_choice == "Y":
                INPUT_SR.mkdir(parents=True, exist_ok=True)
                dest = INPUT_SR / latest_pico.name
                shutil.copy2(str(latest_pico), str(dest))
                print(f"[RCT Search] PICO JSON copied to input/sr/{latest_pico.name}")
            else:
                print("[RCT Search] PICO JSON not copied - you can copy it manually later.")
    # -------------------------------------------------------------------------

    return md_path