"""
writing.py — Writing Mode engine (v1.2 two-track system)
Tracks:
  topic   — newspaper/editorial style, default 800 words
  article — medical journal style, default 3500 words
Pipeline: Writer → Editor → QA (linear)
Standalone: Editor, QA (require input/writing/ files)
"""

import sys
import threading
import datetime
from pathlib import Path
from typing import Callable

try:
    from docx import Document
    from docx.shared import Pt
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
TRACK_TOPIC   = "topic"
TRACK_ARTICLE = "article"

DEFAULT_WORDS = {
    TRACK_TOPIC:   800,
    TRACK_ARTICLE: 3500,
}

DISCLOSURE = (
    "AI Involvement Disclosure: This document was produced with the assistance "
    "of an AI language model. All factual claims should be independently verified "
    "before clinical or public use."
)

# ---------------------------------------------------------------------------
# Path helpers
# ---------------------------------------------------------------------------
def _project_root() -> Path:
    return Path(__file__).resolve().parent.parent.parent


def _ts() -> str:
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")


def _paths(root: Path) -> dict:
    return {
        "doc_root": root / "docs"    / "writing",
        "input":    root / "input"   / "writing",
        "output":   root / "output"  / "writing",
        "reports":  root / "reports" / "writing",
    }


def _track_doc_path(root: Path, track: str) -> Path:
    return root / "docs" / "writing" / track


# ---------------------------------------------------------------------------
# Spinner
# ---------------------------------------------------------------------------
class _Spinner:
    def __init__(self, message: str = "Processing"):
        self._message = message
        self._stop_event = threading.Event()
        self._thread = threading.Thread(target=self._spin, daemon=True)

    def _spin(self) -> None:
        frames = ["|", "/", "-", "\\"]
        idx = 0
        while not self._stop_event.is_set():
            sys.stdout.write(f"\r  {frames[idx % len(frames)]}  {self._message}... ")
            sys.stdout.flush()
            self._stop_event.wait(0.15)
            idx += 1
        sys.stdout.write(f"\r  OK  {self._message} -- done.          \n")
        sys.stdout.flush()

    def __enter__(self):
        self._thread.start()
        return self

    def __exit__(self, *_):
        self._stop_event.set()
        self._thread.join()


# ---------------------------------------------------------------------------
# LLM wrapper
# ---------------------------------------------------------------------------
def _call_llm(
    system_prompt: str,
    user_prompt: str,
    call_llm_fn: Callable,
    spinner_message: str = "LLM processing",
) -> str:
    result_holder: list = []
    error_holder: list  = []

    def _run() -> None:
        try:
            result_holder.append(
                call_llm_fn(system_prompt=system_prompt, user_prompt=user_prompt)
            )
        except Exception as exc:
            error_holder.append(exc)

    with _Spinner(spinner_message):
        t = threading.Thread(target=_run, daemon=True)
        t.start()
        t.join(timeout=300)

    if t.is_alive():
        return "[ERROR] LLM call timed out after 5 minutes."
    if error_holder:
        return f"[ERROR] LLM call failed: {error_holder[0]}"
    if not result_holder:
        return "[ERROR] LLM returned no response."
    return result_holder[0]


# ---------------------------------------------------------------------------
# I/O helpers
# ---------------------------------------------------------------------------
def _load_guidelines(track_doc_path: Path, shared_doc_path: Path) -> str:
    """
    Load all .md files from the track subfolder plus project-brief.md
    from the shared root if present.
    """
    parts = []

    # Shared root: project-brief.md only (if still at root)
    brief = shared_doc_path / "project-brief.md"
    if brief.exists():
        parts.append(f"### project-brief.md\n{brief.read_text(encoding='utf-8')}")

    # Track-specific docs
    if track_doc_path.exists():
        for f in sorted(track_doc_path.glob("*.md")):
            parts.append(f"### {f.name}\n{f.read_text(encoding='utf-8')}")

    return "\n\n".join(parts)


def _load_input_files(input_path: Path) -> list:
    extensions = {".md", ".txt", ".docx", ".html", ".tex", ".rst"}
    if not input_path.exists():
        return []
    results = []
    for f in sorted(input_path.iterdir()):
        if f.suffix.lower() not in extensions:
            continue
        if f.suffix.lower() == ".docx":
            try:
                doc = Document(str(f))
                content = "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                content = f"[Could not read {f.name}]"
        else:
            content = f.read_text(encoding="utf-8", errors="replace")
        results.append((f.stem, content, f.suffix.lower()))
    return results


def _write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def _add_inline_runs(paragraph, text: str) -> None:
    import re
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))")
    for match in pattern.finditer(text):
        if match.group(2):
            paragraph.add_run(match.group(2)).bold = True
        elif match.group(3):
            paragraph.add_run(match.group(3)).italic = True
        elif match.group(4):
            paragraph.add_run(match.group(4))


def _write_docx(path: Path, content: str, title: str = "") -> None:
    if not _DOCX_AVAILABLE:
        print("  [WARN] python-docx not available -- skipping .docx export.")
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    for line in content.splitlines():
        s = line.rstrip()
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif len(s) > 2 and s[0].isdigit() and s[1] in ".)":
            doc.add_paragraph(s[2:].strip(), style="List Number")
        elif s.startswith("---"):
            doc.add_paragraph("")
        else:
            _add_inline_runs(doc.add_paragraph(), s)
    disc = doc.add_paragraph()
    run  = disc.add_run(DISCLOSURE)
    run.italic = True
    run.font.size = Pt(9)
    doc.save(str(path))
    print(f"  [DOCX] {path.name}")


def _strip_fences(text: str) -> str:
    import re
    text = re.sub(r"^```[a-zA-Z]*\n", "", text.strip())
    text = re.sub(r"\n```$", "", text.strip())
    return text.strip()


# ---------------------------------------------------------------------------
# Prompt builders
# ---------------------------------------------------------------------------
_WRITER_PERSONA = {
    TRACK_TOPIC: (
        "You are a senior newspaper columnist and editorial writer with 20 years of "
        "experience writing for national broadsheets. You write with authority, clarity, "
        "and wit. You take clear positions, use plain English, and never hide behind "
        "jargon. Your pieces have a strong hook, a clear argument, and a memorable close."
    ),
    TRACK_ARTICLE: (
        "You are an experienced medical academic writer and researcher. You produce "
        "rigorous, clearly structured journal articles following IMRAD format. You cite "
        "evidence precisely, report effect sizes and confidence intervals, name study "
        "designs correctly, and write in formal academic prose suitable for peer review."
    ),
}

_EDITOR_PERSONA = {
    TRACK_TOPIC: (
        "You are a senior newspaper and magazine editor. You sharpen arguments, improve "
        "sentence rhythm, cut padding, strengthen hooks, and ensure every paragraph "
        "earns its place. You preserve the writer's voice while making the piece tighter, "
        "clearer, and more engaging. You apply newspaper editorial standards rigorously."
    ),
    TRACK_ARTICLE: (
        "You are a medical journal editor and peer reviewer with expertise across clinical "
        "medicine and research methodology. You check structure (IMRAD), verify that "
        "claims are supported by cited evidence, ensure statistical reporting is complete "
        "(effect sizes, CIs, p-values), flag unsupported assertions, and improve academic "
        "prose without altering the scientific content."
    ),
}

_QA_PERSONA = {
    TRACK_TOPIC: (
        "You are a quality assurance reviewer for editorial and opinion writing. You "
        "review documents against the editorial QA checklist: hook quality, argument "
        "clarity, logical flow, plain language, word count, accuracy, and ethics. "
        "Produce a structured QA report with PASS items, FAIL items with specific "
        "references, and a RECOMMENDATION: APPROVE / REVISE MINOR / REVISE MAJOR."
    ),
    TRACK_ARTICLE: (
        "You are a quality assurance reviewer for medical journal articles. You "
        "systematically review the document against the medical article QA checklist: "
        "IMRAD structure, citation completeness, statistical reporting, word count, "
        "ethical compliance, and formatting. Produce a structured QA report with "
        "PASS items, FAIL items with specific section references, and a "
        "RECOMMENDATION: APPROVE / REVISE MINOR / REVISE MAJOR."
    ),
}


def _system_prompt(guidelines: str, role: str, track: str) -> str:
    personas = {
        "Writer": _WRITER_PERSONA,
        "Editor": _EDITOR_PERSONA,
        "QA":     _QA_PERSONA,
    }
    base = personas[role][track]
    if guidelines:
        return (
            f"{base}\n\n"
            "## Guidelines\n"
            "Apply the following guidelines throughout your work.\n\n"
            f"{guidelines}"
        )
    return base


def _writer_user_prompt(
    direct_instructions: list,
    original_content: str,
    is_scratch: bool,
    track: str,
    word_limit: int,
) -> str:
    track_label = "newspaper editorial / opinion piece" if track == TRACK_TOPIC \
        else "medical journal article"
    parts = []
    if direct_instructions:
        parts.append("## DIRECT TASK INSTRUCTIONS (highest priority)")
        for inst in direct_instructions:
            parts.append(f"- {inst.lstrip('>').strip()}")
        parts.append("")
    if is_scratch:
        parts.append(
            f"## Task\n"
            f"No input document was provided. Using the instructions above, produce a "
            f"complete {track_label}. Target word count: {word_limit} words. "
            f"Follow all style and editorial guidelines provided."
        )
    else:
        parts.append("## Source Document")
        parts.append(original_content)
        parts.append(
            f"\n## Task\n"
            f"Using the source document above and the instructions, produce a complete "
            f"{track_label}. Target word count: {word_limit} words. "
            f"Follow all style and editorial guidelines."
        )
    parts.append(
        "\n## Output Format\n"
        "Return the complete document in Markdown. Do not truncate. "
        "Do not add meta-commentary -- output only the document itself."
    )
    return "\n".join(parts)


def _editor_user_prompt(
    direct_instructions: list,
    writer_output: str,
    original_content: str,
    track: str,
    word_limit: int,
) -> str:
    track_label = "editorial/opinion piece" if track == TRACK_TOPIC \
        else "medical journal article"
    parts = []
    if direct_instructions:
        parts.append("## DIRECT TASK INSTRUCTIONS (highest priority)")
        for inst in direct_instructions:
            parts.append(f"- {inst.lstrip('>').strip()}")
        parts.append("")
    parts.append("## Original Source (for reference)")
    parts.append(original_content if original_content else "(No original source.)")
    parts.append("")
    parts.append("## Writer Draft (edit this)")
    parts.append(writer_output)
    parts.append(
        f"\n## Task\n"
        f"Edit the Writer Draft above as a {track_label} editor. "
        f"Target word count: {word_limit} words. "
        f"Return the complete edited document in Markdown. Do not truncate."
    )
    return "\n".join(parts)


def _qa_user_prompt(editor_output: str, track: str) -> str:
    track_label = "editorial/opinion" if track == TRACK_TOPIC else "medical journal article"
    return (
        f"## Document to Review ({track_label})\n\n"
        f"{editor_output}\n\n"
        "## Task\n"
        "Review the document against the QA checklist in your guidelines. "
        "Produce a structured report:\n\n"
        "### PASS\nList every checklist item satisfied.\n\n"
        "### FAIL\nList every unsatisfied item with specific section or line reference.\n\n"
        "### RECOMMENDATION\nState: APPROVE / REVISE MINOR / REVISE MAJOR "
        "with a one-paragraph justification."
    )


# ---------------------------------------------------------------------------
# Report helpers
# ---------------------------------------------------------------------------
def _write_report(
    reports_path: Path,
    role: str,
    track: str,
    stem: str,
    timestamp: str,
    content: str,
    extra_meta: str = "",
) -> Path:
    track_upper = track.upper()
    filename    = f"{role.upper()}_{track_upper}_{stem}_{timestamp}.md"
    header = (
        f"# {role} Report ({track_upper}) -- {stem}\n"
        f"**Timestamp:** {timestamp}\n"
        f"**Track:** {track_upper}\n"
    )
    if extra_meta:
        header += f"{extra_meta}\n"
    header += "\n---\n\n"
    path = reports_path / filename
    _write_text(path, header + content)
    print(f"  [REPORT] {filename}")
    return path


def _write_session_summary(
    reports_path: Path,
    timestamp: str,
    track: str,
    subprojects: list,
) -> None:
    filename = f"WRITER_SESSION_{track.upper()}_{timestamp}_SUMMARY.md"
    lines = [
        f"# Writing Session Summary ({track.upper()})\n",
        f"**Timestamp:** {timestamp}  ",
        f"**Track:** {track.upper()}  ",
        f"**Sub-projects:** {len(subprojects)}\n",
        "---\n",
    ]
    for sp in subprojects:
        lines.append(f"## {sp['stem']}")
        lines.append(f"- Writer:  `{sp.get('writer_docx', 'N/A')}`")
        lines.append(f"- Editor:  `{sp.get('editor_docx', 'N/A')}`")
        lines.append(f"- QA:      `{sp.get('qa_report',   'N/A')}`")
        lines.append(f"- Status:  {sp.get('status', 'unknown')}\n")
    _write_text(reports_path / filename, "\n".join(lines))
    print(f"\n  [SESSION] {filename}")


# ---------------------------------------------------------------------------
# Utility
# ---------------------------------------------------------------------------
def parse_direct_instructions(raw_text: str) -> list:
    return [
        line.strip()
        for line in raw_text.splitlines()
        if line.strip().startswith(">")
    ]


def _prompt_word_limit(track: str) -> int:
    default = DEFAULT_WORDS[track]
    try:
        raw = input(
            f"  Word limit [{default}]: "
        ).strip()
    except (EOFError, KeyboardInterrupt):
        return default
    if raw.isdigit() and int(raw) > 0:
        return int(raw)
    return default


# ---------------------------------------------------------------------------
# Pipeline: Writer → Editor → QA
# ---------------------------------------------------------------------------
def run_writer(
    direct_instructions: list,
    call_llm_fn: Callable,
    track: str = TRACK_TOPIC,
    word_limit: int | None = None,
    verbose: bool = True,
) -> None:
    root      = _project_root()
    paths     = _paths(root)
    timestamp = _ts()

    for p in [paths["input"], paths["output"], paths["reports"]]:
        p.mkdir(parents=True, exist_ok=True)

    if word_limit is None:
        word_limit = DEFAULT_WORDS[track]

    guidelines   = _load_guidelines(_track_doc_path(root, track), paths["doc_root"])
    input_files  = _load_input_files(paths["input"])
    is_scratch   = len(input_files) == 0
    subprojects  = [("new_doc", "", ".md")] if is_scratch else input_files
    session_log  = []

    for stem, original_content, original_suffix in subprojects:
        if verbose:
            print(f"\n{'='*60}")
            print(f"  [WRITER PIPELINE | {track.upper()}] {stem}")
            print(f"  Word limit: {word_limit}")
            print(f"{'='*60}")

        sp = {"stem": stem, "status": "started"}

        # ── WRITER ──────────────────────────────────────────────
        if verbose:
            print(f"\n  [WRITER] Generating draft...")
        writer_out = _call_llm(
            _system_prompt(guidelines, "Writer", track),
            _writer_user_prompt(direct_instructions, original_content,
                                is_scratch, track, word_limit),
            call_llm_fn, "Writer generating",
        )
        writer_out = _strip_fences(writer_out)

        _write_report(paths["reports"], "WRITER", track, stem, timestamp,
                      writer_out,
                      extra_meta=f"**Source:** {'scratch' if is_scratch else stem}")

        w_md   = paths["output"] / f"WRITER_{track.upper()}_{stem}_{timestamp}.md"
        w_docx = paths["output"] / f"WRITER_{track.upper()}_{stem}_{timestamp}.docx"
        _write_text(w_md, writer_out)
        print(f"  [MD]   {w_md.name}")
        _write_docx(w_docx, writer_out, title=f"Writer Draft ({track.upper()}) -- {stem}")
        sp["writer_docx"] = w_docx.name

        # ── EDITOR ──────────────────────────────────────────────
        if verbose:
            print(f"\n  [EDITOR] Editing draft...")
        editor_out = _call_llm(
            _system_prompt(guidelines, "Editor", track),
            _editor_user_prompt(direct_instructions, writer_out,
                                original_content, track, word_limit),
            call_llm_fn, "Editor revising",
        )
        editor_out = _strip_fences(editor_out)

        _write_report(paths["reports"], "EDITOR", track, stem, timestamp, editor_out)

        e_md   = paths["output"] / f"EDITOR_{track.upper()}_{stem}_{timestamp}.md"
        e_docx = paths["output"] / f"EDITOR_{track.upper()}_{stem}_{timestamp}.docx"
        _write_text(e_md, editor_out)
        print(f"  [MD]   {e_md.name}")
        _write_docx(e_docx, editor_out,
                    title=f"Edited Document ({track.upper()}) -- {stem}")
        sp["editor_docx"] = e_docx.name

        # ── QA ──────────────────────────────────────────────────
        if verbose:
            print(f"\n  [QA] Running quality check...")
        qa_out = _call_llm(
            _system_prompt(guidelines, "QA", track),
            _qa_user_prompt(editor_out, track),
            call_llm_fn, "QA reviewing",
        )
        qa_out = _strip_fences(qa_out)

        qa_path = _write_report(paths["reports"], "QA", track, stem, timestamp, qa_out)
        sp["qa_report"] = qa_path.name
        sp["status"]    = "complete"

        if verbose:
            print(f"\n  Done: {stem}")
            print(f"    Writer:  output/writing/{w_docx.name}")
            print(f"    Editor:  output/writing/{e_docx.name}")
            print(f"    QA:      reports/writing/{qa_path.name}")

        session_log.append(sp)

    _write_session_summary(paths["reports"], timestamp, track, session_log)


# ---------------------------------------------------------------------------
# Standalone: Editor
# ---------------------------------------------------------------------------
def run_editor(
    direct_instructions: list,
    call_llm_fn: Callable,
    track: str = TRACK_TOPIC,
    word_limit: int | None = None,
    verbose: bool = True,
) -> None:
    root  = _project_root()
    paths = _paths(root)
    timestamp = _ts()

    for p in [paths["input"], paths["output"], paths["reports"]]:
        p.mkdir(parents=True, exist_ok=True)

    if word_limit is None:
        word_limit = DEFAULT_WORDS[track]

    input_files = _load_input_files(paths["input"])
    if not input_files:
        print("  [EDITOR] No files in input/writing/ -- place documents there first.")
        return

    guidelines = _load_guidelines(_track_doc_path(root, track), paths["doc_root"])

    for stem, content, suffix in input_files:
        if verbose:
            print(f"\n  [EDITOR STANDALONE | {track.upper()}] {stem}{suffix}")
        editor_out = _call_llm(
            _system_prompt(guidelines, "Editor", track),
            _editor_user_prompt(direct_instructions, content, content,
                                track, word_limit),
            call_llm_fn, "Editor processing",
        )
        editor_out = _strip_fences(editor_out)

        _write_report(paths["reports"], "EDITOR", track, stem, timestamp, editor_out)

        e_md   = paths["output"] / f"EDITOR_{track.upper()}_{stem}_{timestamp}.md"
        e_docx = paths["output"] / f"EDITOR_{track.upper()}_{stem}_{timestamp}.docx"
        _write_text(e_md, editor_out)
        _write_docx(e_docx, editor_out,
                    title=f"Edited ({track.upper()}) -- {stem}")


# ---------------------------------------------------------------------------
# Standalone: QA
# ---------------------------------------------------------------------------
def run_qa(
    direct_instructions: list,
    call_llm_fn: Callable,
    track: str = TRACK_TOPIC,
    verbose: bool = True,
) -> None:
    root  = _project_root()
    paths = _paths(root)
    timestamp = _ts()

    for p in [paths["input"], paths["reports"]]:
        p.mkdir(parents=True, exist_ok=True)

    input_files = _load_input_files(paths["input"])
    if not input_files:
        print("  [QA] No files in input/writing/ -- place documents there first.")
        return

    guidelines = _load_guidelines(_track_doc_path(root, track), paths["doc_root"])

    for stem, content, suffix in input_files:
        if verbose:
            print(f"\n  [QA STANDALONE | {track.upper()}] {stem}{suffix}")
        qa_out = _call_llm(
            _system_prompt(guidelines, "QA", track),
            _qa_user_prompt(content, track),
            call_llm_fn, "QA reviewing",
        )
        qa_out = _strip_fences(qa_out)
        _write_report(paths["reports"], "QA", track, stem, timestamp, qa_out)
